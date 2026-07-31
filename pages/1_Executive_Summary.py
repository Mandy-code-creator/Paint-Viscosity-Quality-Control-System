import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import matplotlib.pyplot as plt

from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# =========================================================
# GLOBAL CONFIGURATION
# =========================================================
MIN_REFERENCE_RECORDS = 5          # Minimum adjustment records to use zone-specific reference
FIRST_ADD_WARNING_FRACTION = 0.50  # First addition = 50% of warning ratio
NEXT_ADD_FRACTION = 0.80           # Second addition = 80% of theoretical remaining demand
MIXING_TIME_MINUTES = 5
K_FACTOR = 0.60                    # Damping factor for closed-loop control (60% of theoretical)
STEP_MAX_RATIO_NORMAL = 3.0        # Maximum % per step in safe zone
STEP_MAX_RATIO_MICRO = 1.0         # Maximum % per step in warning zone

# =========================================================
# EXPORT HISTORICAL CHART TO WORD
# =========================================================
def export_chart_to_word(
    selected_resin,
    selected_pos,
    selected_vendor,
    selected_solvent,
    system_df
):
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.40)
    section.right_margin = Inches(0.40)

    # REPORT TITLE
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)

    title_run = title.add_run("Historical Viscosity Transition Analysis")
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)

    subtitle_run = subtitle.add_run(
        f"Resin: {selected_resin} | Position: {selected_pos} | "
        f"Vendor: {selected_vendor} | Solvent Type: {selected_solvent}"
    )
    subtitle_run.font.size = Pt(9)

    # KPI TABLE
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"

    headers = [
        "Valid Paint Batches",
        "Valid Paint Buckets",
        "Median Sensitivity",
        "P10-P90 Ratio Range",
        "Maximum Viscosity Drop"
    ]

    values = [
        f"{system_df['塗料批號'].nunique():,}",
        f"{len(system_df):,}",
        f"{system_df['Sensitivity'].median():.2f} s/%",
        (
            f"{system_df['Solvent_Ratio_Percent'].quantile(0.10):.1f}%"
            f" - {system_df['Solvent_Ratio_Percent'].quantile(0.90):.1f}%"
        ),
        f"{system_df['Delta_V'].max():.1f} s"
    ]

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    for i, value in enumerate(values):
        cell = table.cell(1, i)
        cell.text = value
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(8)

    # CHART
    try:
        fig = plt.figure(figsize=(9.7, 4.45), facecolor="white")
        ax = fig.add_axes([0.10, 0.18, 0.86, 0.55])

        fig.text(
            0.5, 0.95,
            "Viscosity Transition by Solvent Ratio",
            ha="center", va="center",
            fontsize=14, fontweight="bold"
        )
        fig.text(
            0.5, 0.89,
            f"Resin: {selected_resin} | Position: {selected_pos} | "
            f"Vendor: {selected_vendor} | Solvent: {selected_solvent}",
            ha="center", va="center",
            fontsize=11, fontweight="bold"
        )

        for _, row in system_df.iterrows():
            ax.plot(
                [row["Solvent_Ratio_Percent"], row["Solvent_Ratio_Percent"]],
                [row["黏度(秒)"], row["黏度(秒)_1"]],
                linestyle=":", linewidth=0.7, color="lightgray", zorder=1
            )

        before_points = ax.scatter(
            system_df["Solvent_Ratio_Percent"], system_df["黏度(秒)"],
            s=30, color="#ED7D31", edgecolors="white", linewidths=0.4,
            label="Initial Viscosity (Before)", zorder=3
        )

        after_points = ax.scatter(
            system_df["Solvent_Ratio_Percent"], system_df["黏度(秒)_1"],
            s=30, color="#4472C4", edgecolors="white", linewidths=0.4,
            label="Final Viscosity (After)", zorder=3
        )

        ax.set_xlabel("Solvent Blending Ratio (%)", fontsize=10)
        ax.set_ylabel("Viscosity (seconds)", fontsize=10)
        ax.tick_params(axis="both", labelsize=9)
        ax.grid(True, linestyle="--", linewidth=0.5, alpha=0.45)

        fig.legend(
            handles=[before_points, after_points],
            labels=["Initial Viscosity (Before)", "Final Viscosity (After)"],
            loc="upper center", bbox_to_anchor=(0.5, 0.855),
            ncol=2, frameon=False, fontsize=9
        )

        chart_stream = BytesIO()
        fig.savefig(chart_stream, format="png", dpi=220, facecolor="white")
        chart_stream.seek(0)
        plt.close(fig)

        chart_paragraph = doc.add_paragraph()
        chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_paragraph.paragraph_format.space_before = Pt(5)
        chart_paragraph.paragraph_format.space_after = Pt(2)
        chart_paragraph.add_run().add_picture(chart_stream, width=Inches(9.55))

    except Exception as e:
        error_paragraph = doc.add_paragraph()
        error_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        error_run = error_paragraph.add_run(f"[CHART EXPORT FAILED] {str(e)}")
        error_run.bold = True

    # NOTE
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(0)
    note_run = note.add_run(
        "Note: Orange points represent viscosity before solvent addition. "
        "Blue points represent viscosity after solvent addition. "
        "The dotted line connects the initial and final viscosity of the same paint bucket."
    )
    note_run.italic = True
    note_run.font.size = Pt(8)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output.getvalue()


# =========================================================
# PAGE CONFIGURATION
# =========================================================
st.set_page_config(
    page_title="Intelligent SOP System",
    page_icon="⚙️",
    layout="wide"
)

if not st.session_state.get("raw_data_loaded", False):
    st.warning("⚠️ Please upload data on the main page first.")
    st.stop()


# =========================================================
# HELPER FUNCTIONS
# =========================================================
def get_viscosity_zone(value):
    if pd.isna(value):
        return "Unknown"
    if value <= 70:
        return "<=70 s"
    elif value <= 90:
        return "71-90 s"
    elif value <= 110:
        return "91-110 s"
    elif value <= 130:
        return "111-130 s"
    else:
        return ">130 s"

def get_zone_order(zone):
    zone = str(zone)
    if zone.startswith("<="): return 1
    elif zone.startswith("71"): return 2
    elif zone.startswith("91"): return 3
    elif zone.startswith("111"): return 4
    elif zone.startswith("130-") or zone.startswith(">130"): return 5
    return 99

def get_temperature_zone(value):
    if pd.isna(value): return "Unknown"
    if value < 20: return "<20°C"
    elif value < 25: return "20-24.9°C"
    elif value < 30: return "25-29.9°C"
    elif value < 35: return "30-34.9°C"
    else: return ">=35°C"

def format_range(lower_val, upper_val, decimals=1):
    if pd.isna(lower_val) or pd.isna(upper_val):
        return "-"
    l = round(float(lower_val), decimals)
    u = round(float(upper_val), decimals)
    if abs(l - u) < 0.05:
        return f"{l:.{decimals}f}"
    return f"{l:.{decimals}f} - {u:.{decimals}f}"

def get_spc_lower_bound(series):
    """Applies 1-Sigma filtering to eliminate outliers and return a tighter lower limit."""
    clean_s = series.dropna()
    if len(clean_s) < 4:
        return clean_s.quantile(0.25) if not clean_s.empty else np.nan
    mean_val = clean_s.mean()
    std_val = clean_s.std()
    core = clean_s[(clean_s >= mean_val - std_val) & (clean_s <= mean_val + std_val)]
    return core.min() if not core.empty else clean_s.quantile(0.25)

def get_spc_upper_bound(series):
    """Applies 1-Sigma filtering to eliminate outliers and return a tighter upper limit."""
    clean_s = series.dropna()
    if len(clean_s) < 4:
        return clean_s.quantile(0.75) if not clean_s.empty else np.nan
    mean_val = clean_s.mean()
    std_val = clean_s.std()
    core = clean_s[(clean_s >= mean_val - std_val) & (clean_s <= mean_val + std_val)]
    return core.max() if not core.empty else clean_s.quantile(0.75)

def reset_execution_states():
    keys_to_reset = [
        "sop_calculated", "sop_result", "step2_result",
        "step1_added_kg", "step1_after_visc"
    ]
    for key in keys_to_reset:
        st.session_state[key] = None


# =========================================================
# DATA CLEANSING
# =========================================================
@st.cache_data(show_spinner=False)
def prepare_valid_records(df):
    data = df.copy()
    if data.empty: return data

    SPECIAL_INDEPENDENT_PAINT_CODES = {"PS30213X8"}
    required_columns = [
        "添加重量", "塗料重量", "黏度(秒)", "黏度(秒)_1",
        "Resin", "Vendor", "Solvent_Type", "塗料批號", "塗料桶號", "塗料編號"
    ]
    
    missing_columns = [col for col in required_columns if col not in data.columns]
    if missing_columns: return pd.DataFrame()

    numeric_columns = ["添加重量", "塗料重量", "黏度(秒)", "黏度(秒)_1"]
    if "溫度" in data.columns:
        numeric_columns.append("溫度")
    for col in numeric_columns:
        data[col] = pd.to_numeric(data[col], errors="coerce")

    if "塗裝位置" not in data.columns:
        data["塗裝位置"] = "Unknown"

    position_mapping = {
        "TP": "Primer", "正底漆": "Primer", "BP": "Primer", "背底漆": "Primer",
        "TF": "Top Finish", "正面漆": "Top Finish", "BF": "Back Finish", "背面漆": "Back Finish"
    }

    data["Position_UI"] = (
        data["塗裝位置"].fillna("Unknown").astype(str).str.strip()
        .map(position_mapping).fillna(data["塗裝位置"])
    )
    data["塗料編號"] = data["塗料編號"].astype(str).str.strip()

    data = data[
        (data["添加重量"] > 0) & (data["塗料重量"] > 0) &
        (data["黏度(秒)"].notna()) & (data["黏度(秒)_1"].notna()) &
        (data["Resin"].notna()) & (data["Vendor"].notna()) &
        (data["Solvent_Type"].notna()) & (data["塗料批號"].notna()) &
        (data["塗料桶號"].notna()) & (data["塗料編號"].notna())
    ].copy()

    if data.empty: return data
    data["_Original_Row_Order"] = np.arange(len(data))

    # 1. SPECIAL CODE
    special_df = data[data["塗料編號"].isin(SPECIAL_INDEPENDENT_PAINT_CODES)].copy()
    if not special_df.empty:
        special_df["Raw_Adjustment_Rows"] = 1
        special_df["Cumulative_Add_Decreased"] = False
        special_df["Cumulative_Solvent_Added"] = special_df["添加重量"]
        special_df["Final_Mixture_Weight"] = special_df["塗料重量"]
        special_df["Original_Paint_Weight"] = special_df["Final_Mixture_Weight"] - special_df["Cumulative_Solvent_Added"]
        special_df["Dilution_Base"] = special_df["Original_Paint_Weight"]

    # 2. NORMAL CODES
    normal_df = data[~data["塗料編號"].isin(SPECIAL_INDEPENDENT_PAINT_CODES)].copy()
    if not normal_df.empty:
        sort_cols = []
        if "攪拌日期" in normal_df.columns:
            normal_df["_Sort_Date"] = pd.to_datetime(normal_df["攪拌日期"], errors="coerce")
            sort_cols.append("_Sort_Date")
        if "攪拌時間(起)" in normal_df.columns:
            time_text = (
                normal_df["攪拌時間(起)"].astype(str)
                .str.replace(r"\.0$", "", regex=True)
                .str.replace(":", "", regex=False).str.strip().str.zfill(4)
            )
            normal_df["_Sort_Start_Time"] = pd.to_numeric(time_text, errors="coerce")
            sort_cols.append("_Sort_Start_Time")
        sort_cols.append("_Original_Row_Order")

        bucket_group_cols = ["塗料批號", "塗料桶號", "塗料編號"]
        if "攪拌日期" in normal_df.columns: bucket_group_cols.insert(0, "攪拌日期")
        for optional_key in ["Solvent_Type", "Position_UI"]:
            if optional_key in normal_df.columns: bucket_group_cols.append(optional_key)

        normal_df = normal_df.sort_values(bucket_group_cols + sort_cols, kind="stable")
        normal_df["_Previous_Cumulative_Add"] = normal_df.groupby(bucket_group_cols, dropna=False)["添加重量"].shift(1)
        normal_df["_Cumulative_Add_Decreased"] = normal_df["_Previous_Cumulative_Add"].notna() & (normal_df["添加重量"] < normal_df["_Previous_Cumulative_Add"])

        agg_map = {
            "黏度(秒)": "first", "黏度(秒)_1": "last", "添加重量": "last", "塗料重量": "last",
            "Resin": "first", "Vendor": "first", "塗裝位置": "first",
            "_Original_Row_Order": "size", "_Cumulative_Add_Decreased": "max"
        }
        if "溫度" in normal_df.columns: agg_map["溫度"] = "median"
        for optional_col in ["稀釋劑", "稀釋劑批號", "稀釋劑桶號", "攪拌時間(起)", "攪拌時間(迄)"]:
            if optional_col in normal_df.columns and optional_col not in bucket_group_cols:
                agg_map[optional_col] = "last"

        normal_bucket_df = normal_df.groupby(bucket_group_cols, dropna=False, as_index=False).agg(agg_map)
        normal_bucket_df = normal_bucket_df.rename(columns={
            "_Original_Row_Order": "Raw_Adjustment_Rows",
            "_Cumulative_Add_Decreased": "Cumulative_Add_Decreased"
        })
        normal_bucket_df["Cumulative_Solvent_Added"] = normal_bucket_df["添加重量"]
        normal_bucket_df["Final_Mixture_Weight"] = normal_bucket_df["塗料重量"]
        normal_bucket_df["Original_Paint_Weight"] = normal_bucket_df["Final_Mixture_Weight"] - normal_bucket_df["Cumulative_Solvent_Added"]
        normal_bucket_df["Dilution_Base"] = normal_bucket_df["Original_Paint_Weight"]
    else:
        normal_bucket_df = pd.DataFrame()

    # 3. COMBINE
    frames = [df for df in [normal_bucket_df, special_df] if not df.empty]
    if not frames: return pd.DataFrame()
    bucket_df = pd.concat(frames, ignore_index=True, sort=False)

    # 4. COMMON CALCULATIONS
    bucket_df["Delta_V"] = bucket_df["黏度(秒)"] - bucket_df["黏度(秒)_1"]
    bucket_df["Solvent_Ratio_Percent"] = bucket_df["Cumulative_Solvent_Added"] / bucket_df["Dilution_Base"].replace(0, np.nan) * 100
    bucket_df["Sensitivity"] = bucket_df["Delta_V"] / bucket_df["Solvent_Ratio_Percent"].replace(0, np.nan)

    bucket_df = bucket_df[
        (bucket_df["Cumulative_Add_Decreased"] == False) &
        (bucket_df["Original_Paint_Weight"] > 0) &
        (bucket_df["黏度(秒)"] > bucket_df["黏度(秒)_1"]) &
        (bucket_df["Delta_V"] > 0) &
        (bucket_df["Solvent_Ratio_Percent"] > 0) &
        bucket_df["Sensitivity"].notna() &
        np.isfinite(bucket_df["Sensitivity"]) &
        (bucket_df["Sensitivity"] > 0)
    ].copy()

    if bucket_df.empty: return bucket_df
    bucket_df["Initial_Viscosity_Zone"] = bucket_df["黏度(秒)"].apply(get_viscosity_zone)
    
    if "溫度" in bucket_df.columns:
        bucket_df["Temperature_Zone"] = bucket_df["溫度"].apply(get_temperature_zone)
    else:
        bucket_df["溫度"] = np.nan
        bucket_df["Temperature_Zone"] = "Unknown"

    bucket_df["Record_Logic"] = np.where(
        bucket_df["塗料編號"].isin(SPECIAL_INDEPENDENT_PAINT_CODES),
        "Independent row - large source tank",
        "Merged by paint batch and paint bucket"
    )
    return bucket_df.copy()

@st.cache_data(show_spinner=False)
def build_master_system_data(valid_df):
    return valid_df.copy()

# =========================================================
# SATURATION / DIMINISHING RETURNS ANALYSIS
# =========================================================
def build_saturation_profile(df):
    if df.empty:
        return {
            "profile": pd.DataFrame(), "baseline_sensitivity": np.nan,
            "warning_ratio": np.nan, "saturation_ratio": np.nan
        }

    ratio_bins = [0, 3, 5, 7, 9, 11, np.inf]
    ratio_labels = ["0-3%", "3-5%", "5-7%", "7-9%", "9-11%", ">11%"]

    sat_df = df.copy()
    sat_df["Ratio_Zone"] = pd.cut(
        sat_df["Solvent_Ratio_Percent"], bins=ratio_bins, labels=ratio_labels, include_lowest=True, right=False
    )

    profile = (
        sat_df.groupby("Ratio_Zone", observed=False)
        .agg(
            Adjustment_Records=("塗料批號", "size"),
            Paint_Batches=("塗料批號", "nunique"),
            Ratio_Median=("Solvent_Ratio_Percent", "median"),
            Ratio_Min=("Solvent_Ratio_Percent", "min"),
            Ratio_Max=("Solvent_Ratio_Percent", "max"),
            DeltaV_Median=("Delta_V", "median"),
            Sensitivity_Median=("Sensitivity", "median"),
            Sensitivity_P25=("Sensitivity", lambda x: x.quantile(0.25)),
            Sensitivity_P75=("Sensitivity", lambda x: x.quantile(0.75))
        ).reset_index()
    )

    profile["Efficiency_vs_Baseline_%"] = np.nan
    profile["Saturation_Status"] = "Insufficient Data"

    valid_profile = profile[
        (profile["Adjustment_Records"] >= MIN_REFERENCE_RECORDS) &
        (profile["Sensitivity_Median"] > 0) &
        (profile["Sensitivity_Median"].notna())
    ].copy()

    baseline_sensitivity = np.nan
    warning_ratio = np.nan
    saturation_ratio = np.nan

    if not valid_profile.empty:
        valid_profile = valid_profile.sort_values("Ratio_Min")
        baseline_sensitivity = valid_profile.iloc[0]["Sensitivity_Median"]
        
        if baseline_sensitivity > 0:
            profile["Efficiency_vs_Baseline_%"] = profile["Sensitivity_Median"] / baseline_sensitivity * 100

        for idx, row in profile.iterrows():
            if row["Adjustment_Records"] < MIN_REFERENCE_RECORDS or pd.isna(row["Efficiency_vs_Baseline_%"]):
                continue
            efficiency = row["Efficiency_vs_Baseline_%"]
            if efficiency <= 50:
                profile.loc[idx, "Saturation_Status"] = "🔴 Saturation Zone"
                if pd.isna(saturation_ratio): saturation_ratio = row["Ratio_Min"]
            elif efficiency <= 70:
                profile.loc[idx, "Saturation_Status"] = "🟠 Diminishing Returns"
                if pd.isna(warning_ratio): warning_ratio = row["Ratio_Min"]
            else:
                profile.loc[idx, "Saturation_Status"] = "🟢 Normal Efficiency"

    return {
        "profile": profile, "baseline_sensitivity": baseline_sensitivity,
        "warning_ratio": warning_ratio, "saturation_ratio": saturation_ratio
    }

def get_reference_data(system_df, current_viscosity):
    current_zone = get_viscosity_zone(current_viscosity)
    zone_df = system_df[system_df["Initial_Viscosity_Zone"] == current_zone].copy()

    if len(zone_df) >= MIN_REFERENCE_RECORDS:
        return {
            "reference_df": zone_df, "reference_source": f"Zone-Specific ({current_zone})",
            "current_zone": current_zone, "record_count": len(zone_df), "batch_count": zone_df["塗料批號"].nunique()
        }
    return {
        "reference_df": system_df.copy(), "reference_source": "Overall System Fallback",
        "current_zone": current_zone, "record_count": len(system_df), "batch_count": system_df["塗料批號"].nunique()
    }

def get_safety_limits(reference_df, saturation_result):
    if reference_df.empty:
        return {
            "warning_ratio": np.nan, "stop_ratio": np.nan, "ratio_p90": np.nan,
            "ratio_p95": np.nan, "drop_p90": np.nan, "drop_max": np.nan
        }

    ratio_p90 = reference_df["Solvent_Ratio_Percent"].quantile(0.90)
    ratio_p95 = reference_df["Solvent_Ratio_Percent"].quantile(0.95)
    drop_p90 = reference_df["Delta_V"].quantile(0.90)
    drop_max = reference_df["Delta_V"].max()

    warning_ratio = ratio_p90
    stop_ratio = ratio_p95
    sat_warning = saturation_result.get("warning_ratio", np.nan)
    sat_stop = saturation_result.get("saturation_ratio", np.nan)

    if not pd.isna(sat_warning): warning_ratio = min(warning_ratio, sat_warning)
    if not pd.isna(sat_stop): stop_ratio = min(stop_ratio, sat_stop)
    stop_ratio = max(stop_ratio, warning_ratio)

    return {
        "warning_ratio": warning_ratio, "stop_ratio": stop_ratio,
        "ratio_p90": ratio_p90, "ratio_p95": ratio_p95,
        "drop_p90": drop_p90, "drop_max": drop_max
    }

def get_temperature_check(reference_df, current_temperature):
    if reference_df.empty or pd.isna(current_temperature):
        return {"available": False, "median": np.nan, "p25": np.nan, "p75": np.nan, "warning": False}
    temp_data = reference_df["溫度"].dropna()
    if len(temp_data) < MIN_REFERENCE_RECORDS:
        return {"available": False, "median": np.nan, "p25": np.nan, "p75": np.nan, "warning": False}
    
    temp_median = temp_data.median()
    temp_p25 = temp_data.quantile(0.25)
    temp_p75 = temp_data.quantile(0.75)
    warning = current_temperature < (temp_p25 - 3) or current_temperature > (temp_p75 + 3)
    return {"available": True, "median": temp_median, "p25": temp_p25, "p75": temp_p75, "warning": warning}


# =========================================================
# LOAD DATA
# =========================================================
group_a_data = st.session_state.get("group_a_data")
if group_a_data is None or group_a_data.empty:
    st.warning("⚠️ No valid Group A data found. Please upload the source file again from the main page.")
    st.stop()

valid_df = prepare_valid_records(group_a_data)
master_df = build_master_system_data(valid_df)

if master_df.empty:
    st.warning("⚠️ 無可用歷史資料，請確認是否具有有效加料前後黏度資料。")
    st.stop()


# =========================================================
# PAGE TITLE
# =========================================================
st.title("⚙️ AI-Assisted Viscosity Optimization System")
st.markdown(
    "Closed-loop solvent recommendation system. "
    "The system guarantees safety by applying a strict first-addition limit and calculating "
    "subsequent steps based on actual measured efficiency."
)
st.markdown("---")


# =========================================================
# GLOBAL FILTERS
# =========================================================
col_f1, col_f2, col_f3, col_f4 = st.columns(4)
with col_f1:
    selected_resin = st.selectbox("Select Resin:", sorted(master_df["Resin"].dropna().unique()), on_change=reset_execution_states)
with col_f2:
    available_positions = sorted(master_df.loc[master_df["Resin"] == selected_resin, "Position_UI"].dropna().unique())
    selected_pos = st.selectbox("Select Position:", available_positions, on_change=reset_execution_states)
with col_f3:
    available_vendors = sorted(master_df.loc[(master_df["Resin"] == selected_resin) & (master_df["Position_UI"] == selected_pos), "Vendor"].dropna().unique())
    selected_vendor = st.selectbox("Select Vendor:", available_vendors, on_change=reset_execution_states)
with col_f4:
    available_solvents = sorted(master_df.loc[(master_df["Resin"] == selected_resin) & (master_df["Position_UI"] == selected_pos) & (master_df["Vendor"] == selected_vendor), "Solvent_Type"].dropna().unique())
    selected_solvent = st.selectbox("Select Solvent Type:", available_solvents, on_change=reset_execution_states)

system_df = master_df[
    (master_df["Resin"] == selected_resin) & (master_df["Position_UI"] == selected_pos) &
    (master_df["Vendor"] == selected_vendor) & (master_df["Solvent_Type"] == selected_solvent)
].copy()

if system_df.empty:
    st.error("No valid historical data available for this configuration.")
    st.stop()


# =========================================================
# TABS
# =========================================================
tab1, tab2, tab3, tab4 = st.tabs([
    "📊 Tab 1: Historical Analysis",
    "🎯 Tab 2: SOP Recommendation",
    "🔬 Tab 3: Engineering Matrix",
    "🖨️ Tab 4: Master Shop Floor SOP (V2.0)"
])

# =========================================================
# TAB 1: HISTORICAL ANALYSIS
# =========================================================
with tab1:
    st.markdown("### Historical Performance Review")
    st.markdown("Historical records are shown only for the selected Resin × Position × Vendor × Solvent Type system.")

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Valid Paint Batches", f"{system_df['塗料批號'].nunique():,}")
    c2.metric("Valid Paint Buckets", f"{len(system_df):,}")
    c3.metric("Median Sensitivity", f"{system_df['Sensitivity'].median():.2f} s/%")
    c4.metric("P10 - P90 Ratio Range", f"{system_df['Solvent_Ratio_Percent'].quantile(0.10):.1f}% - {system_df['Solvent_Ratio_Percent'].quantile(0.90):.1f}%")
    c5.metric("Maximum Viscosity Drop", f"{system_df['Delta_V'].max():.1f} s")

    fig_scatter = go.Figure()
    plot_df = system_df.reset_index(drop=True).copy()

    for _, row in plot_df.iterrows():
        fig_scatter.add_trace(go.Scatter(
            x=[row["Solvent_Ratio_Percent"], row["Solvent_Ratio_Percent"]],
            y=[row["黏度(秒)"], row["黏度(秒)_1"]],
            mode="lines",
            line=dict(color="rgba(120,120,120,0.35)", width=1.4, dash="dot"),
            customdata=[[row["塗料批號"], row["Delta_V"]]] * 2,
            hovertemplate="<b>Batch: %{customdata[0]}</b><br>Viscosity Drop: %{customdata[1]:.1f}s<extra></extra>",
            showlegend=False
        ))

    fig_scatter.add_trace(go.Scatter(
        x=plot_df["Solvent_Ratio_Percent"], y=plot_df["黏度(秒)"], mode="markers",
        name="Initial Viscosity (Before)",
        marker=dict(color="#ED7D31", size=8, opacity=0.85, line=dict(width=0.8, color="white")),
        customdata=plot_df[["黏度(秒)_1", "Delta_V", "Initial_Viscosity_Zone", "塗料批號", "溫度"]].values,
        hovertemplate=(
            "<b>Batch: %{customdata[3]}</b><br><b>Initial Zone: %{customdata[2]}</b><br>"
            "Temperature: %{customdata[4]:.1f}°C<br>Solvent Ratio: %{x:.2f}%<br>"
            "Initial Viscosity: %{y:.1f}s<br>Final Viscosity: %{customdata[0]:.1f}s<br>"
            "Viscosity Drop: %{customdata[1]:.1f}s<extra></extra>"
        )
    ))

    fig_scatter.add_trace(go.Scatter(
        x=plot_df["Solvent_Ratio_Percent"], y=plot_df["黏度(秒)_1"], mode="markers",
        name="Final Viscosity (After)",
        marker=dict(color="#4472C4", size=8, opacity=0.85, line=dict(width=0.8, color="white")),
        customdata=plot_df[["黏度(秒)", "Delta_V", "Initial_Viscosity_Zone", "塗料批號", "溫度"]].values,
        hovertemplate=(
            "<b>Batch: %{customdata[3]}</b><br><b>Initial Zone: %{customdata[2]}</b><br>"
            "Temperature: %{customdata[4]:.1f}°C<br>Solvent Ratio: %{x:.2f}%<br>"
            "Initial Viscosity: %{customdata[0]:.1f}s<br>Final Viscosity: %{y:.1f}s<br>"
            "Viscosity Drop: %{customdata[1]:.1f}s<extra></extra>"
        )
    ))

    fig_scatter.update_layout(
        title=dict(text=(
            "Viscosity Transition by Solvent Ratio<br>"
            f"<sup>Resin: {selected_resin} | Position: {selected_pos} | Vendor: {selected_vendor} | Solvent: {selected_solvent}</sup>"
        ), x=0.5, xanchor="center", y=0.97, yanchor="top", font=dict(size=18, color="#1F3855")),
        plot_bgcolor="white", paper_bgcolor="white", height=620, margin=dict(l=70, r=50, t=95, b=70),
        xaxis=dict(title="Solvent Blending Ratio (%)", showgrid=True, gridcolor="#EAEAEA", linecolor="black", linewidth=1.5, showline=True, mirror=True, ticks="outside"),
        yaxis=dict(title="Viscosity (seconds)", showgrid=True, gridcolor="#EAEAEA", linecolor="black", linewidth=1.5, showline=True, mirror=True, ticks="outside"),
        legend=dict(orientation="h", yanchor="bottom", y=1.07, xanchor="center", x=0.5), hovermode="closest"
    )

    st.plotly_chart(fig_scatter, use_container_width=True)

    word_data = export_chart_to_word(selected_resin, selected_pos, selected_vendor, selected_solvent, system_df)
    st.download_button(
        label="📄 Export Historical Chart to Word", data=word_data,
        file_name=f"Viscosity_Transition_{selected_resin}_{selected_pos}_{selected_vendor}_{selected_solvent}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# =========================================================
# TAB 2: SOP RECOMMENDATION (CLOSED-LOOP CONTROL)
# =========================================================
with tab2:
    st.markdown("### 現場稀釋劑添加 SOP 計算")
    st.info(
        "操作原則：系統已升級為閉迴路控制 (Closed-loop Control)。\n\n"
        "1. 第一次添加將被嚴格限制在警戒比例的 50% 或安全界限內。\n"
        "2. 攪拌後重新量測，系統將根據該桶塗料的真實反應，計算後續安全添加量。"
    )

    input_col1, input_col2, input_col3, input_col4, input_col5 = st.columns(5)
    with input_col1: current_visc = st.number_input("Current Viscosity (s)", min_value=0.0, value=0.0, step=1.0, key="current_viscosity_input")
    with input_col2: target_lsl = st.number_input("Approved Target LSL (s)", min_value=0.0, value=0.0, step=1.0, key="target_lsl_input")
    with input_col3: target_usl = st.number_input("Approved Target USL (s)", min_value=0.0, value=0.0, step=1.0, key="target_usl_input")
    with input_col4: actual_paint_weight = st.number_input("Actual Paint Weight (kg)", min_value=0.0, value=0.0, step=1.0, key="actual_paint_weight_input")
    with input_col5: current_temperature = st.number_input("Current Temperature (°C)", min_value=0.0, value=0.0, step=0.5, key="current_temperature_input")

    st.markdown("---")

    if st.button("🚀 Calculate First Addition", type="primary", use_container_width=True):
        st.session_state["sop_calculated"] = False
        st.session_state["sop_result"] = None
        st.session_state["step2_result"] = None

        if current_visc <= 0 or target_lsl <= 0 or target_usl <= 0 or actual_paint_weight <= 0:
            st.error("⚠️ 請輸入完整且大於 0 的黏度、規格及實際塗料重量。")
        elif target_lsl >= target_usl:
            st.error("⚠️ LSL 必須小於 USL。")
        elif current_visc < target_lsl:
            st.error("🚨 目前黏度低於規格下限（LSL）。不可添加稀釋劑，請通知製程工程師確認。")
        elif current_visc <= target_usl:
            st.success("✅ 目前黏度已在規格範圍內，不需添加稀釋劑。")
        else:
            target_center = (target_lsl + target_usl) / 2
            required_drop = current_visc - target_center
            
            reference_info = get_reference_data(system_df, current_visc)
            reference_df = reference_info["reference_df"]

            if len(reference_df) < MIN_REFERENCE_RECORDS:
                st.error("🚨 歷史有效紀錄不足，系統無法提供安全建議。請通知製程工程師。")
            else:
                saturation_result = build_saturation_profile(reference_df)
                safety_limits = get_safety_limits(reference_df, saturation_result)

                ref_sensitivity = reference_df["Sensitivity"].median()
                sensitivity_p25 = reference_df["Sensitivity"].quantile(0.25)
                ref_sensitivity_p75 = reference_df["Sensitivity"].quantile(0.75)
                
                required_ratio = required_drop / ref_sensitivity
                recommended_total_kg = (actual_paint_weight * required_ratio) / 100

                warning_ratio = safety_limits["warning_ratio"]
                stop_ratio = safety_limits["stop_ratio"]

                # CLOSED LOOP STEP 1: Dual-Safety Lock Algorithm
                safe_visc_target = target_usl 
                max_safe_ratio = max(0.0, (current_visc - safe_visc_target) / ref_sensitivity_p75) if ref_sensitivity_p75 > 0 else 0.0
                warning_lock_ratio = warning_ratio * FIRST_ADD_WARNING_FRACTION
                
                first_add_ratio = min(required_ratio, warning_lock_ratio, max_safe_ratio)
                first_add_kg = (actual_paint_weight * first_add_ratio) / 100

                blocked = False
                risk_status = ""
                risk_color = "green"

                if required_ratio > stop_ratio or required_drop > safety_limits["drop_max"]:
                    blocked = True
                    risk_status = "🚨 超過歷史停止比例或最大降黏範圍，不可執行自動加料。"
                    risk_color = "red"
                elif required_ratio > warning_ratio or required_drop > safety_limits["drop_p90"]:
                    risk_status = "⚠️ 已進入警戒區，系統啟動雙重安全鎖定，請執行首次安全添加。"
                    risk_color = "orange"
                else:
                    risk_status = "✅ 正常作業區：建議第一次添加後重新量測。"
                    risk_color = "green"
                    
                if first_add_ratio <= 0 and not blocked:
                    blocked = True
                    risk_status = "🚨 安全鎖定觸發：系統計算首次安全添加量為 0，不可執行自動加料。"
                    risk_color = "red"

                temperature_check = get_temperature_check(reference_df, current_temperature)
                st.markdown(f"### 評估結果：<span style='color:{risk_color}'>{risk_status}</span>", unsafe_allow_html=True)
                
                if temperature_check["available"] and temperature_check["warning"]:
                    st.warning("⚠️ 目前溫度與歷史參考溫度差異較大。系統已自動套用安全機制，請嚴格遵守首次添加量。")

                if not blocked:
                    result_col1, result_col2, result_col3 = st.columns(3)
                    result_col1.metric("Required Viscosity Drop", f"{required_drop:.1f} s")
                    result_col2.metric("Warning / Stop Ratio", f"{warning_ratio:.2f}% / {stop_ratio:.2f}%")
                    result_col3.metric("Data Confidence", "🟢 Reliable" if reference_info["record_count"] >= 20 else "🟡 Limited")

                    st.success(f"### 第一次安全添加建議：{first_add_kg:.2f} kg ({first_add_ratio:.2f}%)")
                    st.markdown(
                        f"**現場操作：**\n\n"
                        f"1. 實際塗料重量：`{actual_paint_weight:.1f} kg`\n"
                        f"2. 理論總需求比例：`{required_ratio:.2f}%`\n"
                        f"3. 第一次安全添加比例：`{first_add_ratio:.2f}%`\n"
                        f"4. 第一次安全添加量：`{first_add_kg:.2f} kg`\n"
                        f"5. 添加後攪拌至少：`{MIXING_TIME_MINUTES} 分鐘`\n"
                        f"6. 攪拌後重新量測黏度，再使用下方功能計算下一步。"
                    )

                    st.session_state["sop_calculated"] = True
                    st.session_state["sop_result"] = {
                        "initial_visc": current_visc, "target_center": target_center,
                        "target_lsl": target_lsl, "target_usl": target_usl,
                        "actual_paint_weight": actual_paint_weight,
                        "reference_source": reference_info["reference_source"],
                        "reference_record_count": reference_info["record_count"],
                        "reference_batch_count": reference_info["batch_count"],
                        "reference_sensitivity": ref_sensitivity,
                        "sensitivity_p25": sensitivity_p25,
                        "sensitivity_p75": ref_sensitivity_p75,
                        "warning_ratio": warning_ratio, "stop_ratio": stop_ratio,
                        "estimated_total_ratio": required_ratio,
                        "estimated_total_kg": recommended_total_kg,
                        "warning_lock_ratio": warning_lock_ratio,
                        "first_add_ratio": first_add_ratio, "first_add_kg": first_add_kg
                    }

    # STEP 2: RECALCULATE AFTER FIRST ADDITION (Feedback Loop)
    if st.session_state.get("sop_calculated", False):
        sop_result = st.session_state.get("sop_result")
        if sop_result:
            st.markdown("---")
            st.markdown("### 🔁 第二次計算：回饋控制 (Feedback Loop)")
            st.caption("請輸入實際第一次添加量與攪拌後的實測黏度。系統將計算真實稀釋效率，並應用三層防護計算後續添加量。")

            step_col1, step_col2 = st.columns(2)
            with step_col1: actual_step1_kg = st.number_input("Actual First Addition (kg)", min_value=0.0, value=float(sop_result["first_add_kg"]), step=0.1, key="actual_step1_kg_input")
            with step_col2: measured_after_step1 = st.number_input("Measured Viscosity After First Addition (s)", min_value=0.0, value=0.0, step=1.0, key="measured_after_step1_input")

            if st.button("Calculate Additional Solvent", type="secondary", use_container_width=True):
                if actual_step1_kg <= 0 or measured_after_step1 <= 0:
                    st.error("⚠️ 請輸入第一次實際添加量與添加後實測黏度。")
                elif measured_after_step1 < sop_result["target_lsl"]:
                    st.error("🚨 添加後黏度已低於 LSL。不可再添加稀釋劑，請通知製程工程師。")
                elif measured_after_step1 <= sop_result["target_usl"]:
                    st.success("✅ 添加後黏度已落入規格範圍，不需再添加稀釋劑。")
                elif measured_after_step1 >= sop_result["initial_visc"]:
                    st.error("🚨 第一次添加後黏度未下降。請確認稀釋劑種類、攪拌時間、量測方法及原料狀態。")
                else:
                    actual_step1_ratio = (actual_step1_kg / sop_result["actual_paint_weight"]) * 100
                    
                    # CLOSED LOOP STEP 2: Calculate Actual Efficiency
                    observed_drop = sop_result["initial_visc"] - measured_after_step1
                    observed_sensitivity = observed_drop / actual_step1_ratio
                    
                    safe_sensitivity = np.clip(
                        observed_sensitivity,
                        sop_result["sensitivity_p25"],
                        sop_result["sensitivity_p75"]
                    )
                    
                    remaining_drop = measured_after_step1 - sop_result["target_center"]
                    theoretical_additional_ratio = remaining_drop / safe_sensitivity

                    # CLOSED LOOP STEP 3: Three-Layer Protection
                    protected_additional_ratio = theoretical_additional_ratio * NEXT_ADD_FRACTION
                    
                    remaining_ratio_to_warning = max(sop_result["warning_ratio"] - actual_step1_ratio, 0.0)
                    remaining_ratio_to_stop = sop_result["stop_ratio"] - actual_step1_ratio

                    if remaining_ratio_to_warning > 0:
                        limit_2 = STEP_MAX_RATIO_NORMAL
                        limit_3 = remaining_ratio_to_warning
                        warning_triggered = False
                    else:
                        limit_2 = STEP_MAX_RATIO_MICRO
                        limit_3 = max(remaining_ratio_to_stop, 0.0)
                        warning_triggered = True

                    additional_ratio = min(protected_additional_ratio * K_FACTOR, limit_2, limit_3)
                    additional_kg = (sop_result["actual_paint_weight"] * additional_ratio) / 100
                    total_ratio_after_step2 = actual_step1_ratio + additional_ratio

                    second_step_blocked = False

                    if remaining_ratio_to_stop <= 0:
                        second_step_blocked = True
                        st.error("🚨 第一次添加後已達停止比例，不可繼續添加稀釋劑。")
                    elif actual_step1_ratio >= sop_result["warning_ratio"]:
                        second_step_blocked = True
                        st.error("⚠️ 目前累積添加比例已達警戒值。系統停止自動建議，請由製程工程師確認後續處理。")
                    elif additional_ratio <= 0:
                        second_step_blocked = True
                        st.error("🚨 已無警戒比例以下的安全追加空間。請通知製程工程師。")
                    elif total_ratio_after_step2 >= sop_result["warning_ratio"]:
                        st.warning("⚠️ 本次建議最多追加至警戒比例。添加並攪拌後請重新量測；若仍高於規格，請由工程師確認。")
                        
                    if not second_step_blocked:
                        step2_col1, step2_col2, step2_col3, step2_col4 = st.columns(4)
                        step2_col1.metric("Observed Sensitivity", f"{observed_sensitivity:.2f} s/%")
                        step2_col2.metric("Safe Sensitivity Used", f"{safe_sensitivity:.2f} s/%")
                        step2_col3.metric("Theoretical / Allowed Ratio", f"{theoretical_additional_ratio:.2f}% / {additional_ratio:.2f}%")
                        step2_col4.metric("Additional Solvent", f"{additional_kg:.2f} kg")

                        st.success(f"### 建議下一次添加：{additional_kg:.2f} kg")
                        st.markdown(
                            "**重要：** 本次建議已應用衰減係數與步進限制，並限制不得超過累積添加警戒比例。\n"
                            "添加後請重新攪拌並確認黏度。"
                        )

                        st.session_state["step2_result"] = {
                            "observed_sensitivity": observed_sensitivity,
                            "safe_sensitivity": safe_sensitivity,
                            "theoretical_additional_ratio": theoretical_additional_ratio,
                            "additional_ratio": additional_ratio,
                            "additional_kg": additional_kg,
                            "total_ratio_after_step2": total_ratio_after_step2
                        }


# =========================================================
# TAB 3: ENGINEERING MATRIX
# =========================================================
with tab3:
    st.markdown("### 🔬 Engineering Matrix")
    st.markdown("This page is for engineering review. Tab 2 should be used for actual shop-floor execution.")

    eng_matrix = (
        system_df.groupby("Initial_Viscosity_Zone", observed=False)
        .agg(
            Adjustment_Records=("塗料批號", "size"),
            Paint_Batches=("塗料批號", "nunique"),
            Sensitivity_Median=("Sensitivity", "median"),
            Sensitivity_P25=("Sensitivity", lambda x: x.quantile(0.25)),
            Sensitivity_P75=("Sensitivity", lambda x: x.quantile(0.75)),
            Ratio_Median=("Solvent_Ratio_Percent", "median"),
            Ratio_P90=("Solvent_Ratio_Percent", lambda x: x.quantile(0.90)),
            Ratio_P95=("Solvent_Ratio_Percent", lambda x: x.quantile(0.95)),
            Drop_Median=("Delta_V", "median"),
            Drop_P90=("Delta_V", lambda x: x.quantile(0.90)),
            Drop_Max=("Delta_V", "max"),
            Temp_Median=("溫度", "median"),
            Temp_P25=("溫度", lambda x: x.quantile(0.25)),
            Temp_P75=("溫度", lambda x: x.quantile(0.75))
        ).reset_index()
    )

    eng_matrix["_zone_order"] = eng_matrix["Initial_Viscosity_Zone"].apply(get_zone_order)
    eng_matrix = eng_matrix.sort_values("_zone_order").drop(columns="_zone_order")

    st.dataframe(
        eng_matrix,
        column_config={
            "Initial_Viscosity_Zone": st.column_config.TextColumn("Initial Viscosity Zone"),
            "Adjustment_Records": st.column_config.NumberColumn("有效調整紀錄數", format="%d"),
            "Paint_Batches": st.column_config.NumberColumn("涉及塗料批號數", format="%d"),
            "Sensitivity_Median": st.column_config.NumberColumn("Median Sensitivity (s/%)", format="%.2f"),
            "Sensitivity_P25": st.column_config.NumberColumn("Sensitivity P25", format="%.2f"),
            "Sensitivity_P75": st.column_config.NumberColumn("Sensitivity P75", format="%.2f"),
            "Ratio_Median": st.column_config.NumberColumn("Median Total Ratio (%)", format="%.2f"),
            "Ratio_P90": st.column_config.NumberColumn("Warning Ratio P90 (%)", format="%.2f"),
            "Ratio_P95": st.column_config.NumberColumn("Stop Ratio P95 (%)", format="%.2f"),
            "Drop_Median": st.column_config.NumberColumn("Median Drop (s)", format="%.1f"),
            "Drop_P90": st.column_config.NumberColumn("Drop P90 (s)", format="%.1f"),
            "Drop_Max": st.column_config.NumberColumn("Maximum Drop (s)", format="%.1f"),
            "Temp_Median": st.column_config.NumberColumn("Median Temperature (°C)", format="%.1f"),
            "Temp_P25": st.column_config.NumberColumn("Temperature P25 (°C)", format="%.1f"),
            "Temp_P75": st.column_config.NumberColumn("Temperature P75 (°C)", format="%.1f")
        },
        use_container_width=True, hide_index=True
    )

    st.markdown("---")
    st.markdown("### 📉 Saturation / Diminishing Returns by Initial Viscosity Zone")

    available_zones = sorted(system_df["Initial_Viscosity_Zone"].dropna().unique(), key=get_zone_order)
    selected_sat_zone = st.selectbox("Select Initial Viscosity Zone for Saturation Analysis:", available_zones, key="saturation_zone_selector")
    
    saturation_df = system_df[system_df["Initial_Viscosity_Zone"] == selected_sat_zone].copy()
    saturation_result = build_saturation_profile(saturation_df)
    saturation_profile = saturation_result["profile"]

    sat_col1, sat_col2, sat_col3, sat_col4 = st.columns(4)
    sat_col1.metric("Zone Records", f"{len(saturation_df):,}")
    sat_col2.metric("Baseline Sensitivity", f"{saturation_result['baseline_sensitivity']:.2f} s/%" if not pd.isna(saturation_result['baseline_sensitivity']) else "Not Detected")
    sat_col3.metric("Diminishing Return Threshold", f"{saturation_result['warning_ratio']:.2f}%" if not pd.isna(saturation_result['warning_ratio']) else "Not Detected")
    sat_col4.metric("Saturation Stop Threshold", f"{saturation_result['saturation_ratio']:.2f}%" if not pd.isna(saturation_result['saturation_ratio']) else "Not Detected")

    if saturation_profile.empty:
        st.warning("No saturation profile available for this viscosity zone.")
    else:
        st.dataframe(
            saturation_profile,
            column_config={
                "Ratio_Zone": st.column_config.TextColumn("Solvent Ratio Zone"),
                "Adjustment_Records": st.column_config.NumberColumn("有效調整紀錄數", format="%d"),
                "Paint_Batches": st.column_config.NumberColumn("涉及塗料批號數", format="%d"),
                "Ratio_Median": st.column_config.NumberColumn("Median Ratio (%)", format="%.2f"),
                "Ratio_Min": st.column_config.NumberColumn("Minimum Ratio (%)", format="%.2f"),
                "Ratio_Max": st.column_config.NumberColumn("Maximum Ratio (%)", format="%.2f"),
                "DeltaV_Median": st.column_config.NumberColumn("Median Drop (s)", format="%.2f"),
                "Sensitivity_Median": st.column_config.NumberColumn("Median Sensitivity (s/%)", format="%.2f"),
                "Sensitivity_P25": st.column_config.NumberColumn("Sensitivity P25", format="%.2f"),
                "Sensitivity_P75": st.column_config.NumberColumn("Sensitivity P75", format="%.2f"),
                "Efficiency_vs_Baseline_%": st.column_config.NumberColumn("Efficiency vs Baseline (%)", format="%.1f%%"),
                "Saturation_Status": st.column_config.TextColumn("Efficiency Status")
            },
            use_container_width=True, hide_index=True
        )


# =========================================================
# TAB 4: MASTER SHOP FLOOR SOP (V2.0 DELTA VISCOSITY MODEL)
# =========================================================
with tab4:
    st.markdown("### 🖨️ 現場 SOP (V2.0 物理降幅模型)")

    st.warning(
        "操作方式：先確認目前黏度區間，再選擇欲達到的目標黏度區間，"
        "依建議添加比例一次加料，攪拌5分鐘後重新量測。"
        "建議添加比例不得超過警戒比例，累積添加不得達到停止比例。"
    )

    st.caption(
        "註：系統已升級至 V2.0 (Delta Viscosity Model)。\n"
        "系統先透過「(初始黏度 - 目標黏度) ÷ 稀釋效率」計算預估累積總需求，"
        "再另外套用單次步進與警戒限制，產生第一次安全添加比例。兩者用途不同，不可混為同一數值。"
    )

    matrix_df = valid_df.copy()

    if matrix_df.empty:
        st.warning("無有效 Group A 調整資料可建立現場 SOP。")
        st.stop()

    # -----------------------------------------------------
    # 1. 建立初始黏度區間 (Initial Viscosity Zones)
    # -----------------------------------------------------
    def create_worker_viscosity_zone(df):
        temp_df = df.copy()
        group_cols = ["Resin", "Position_UI", "Vendor", "Solvent_Type"]
        system_max_visc = temp_df.groupby(group_cols)["黏度(秒)"].transform("max")
        
        def worker_zone(viscosity):
            if pd.isna(viscosity): return "Unknown"
            if viscosity <= 70: return "<=70"
            elif viscosity <= 90: return "71-90"
            elif viscosity <= 110: return "91-110"
            elif viscosity <= 130: return "111-130"
            return ">130"

        temp_df["Worker_Viscosity_Zone"] = temp_df["黏度(秒)"].apply(worker_zone)
        high_visc_mask = temp_df["黏度(秒)"] > 130
        temp_df.loc[high_visc_mask, "Worker_Viscosity_Zone"] = "130-" + system_max_visc.loc[high_visc_mask].round(1).astype(str)
        return temp_df

    # -----------------------------------------------------
    # 2. 建立目標黏度區間 (Target Viscosity Zones)
    # -----------------------------------------------------
    def create_final_viscosity_zone(value):
        if pd.isna(value) or value <= 0: return "Unknown"
        upper = int(np.ceil(float(value) / 10.0) * 10)
        lower = max(1, upper - 9)
        return f"{lower}~{upper}"

    def get_final_zone_order(zone):
        try: return int(str(zone).split("~")[0])
        except Exception: return 9999

    matrix_df = create_worker_viscosity_zone(matrix_df)
    matrix_df["Target_Viscosity_Zone"] = matrix_df["黏度(秒)_1"].apply(create_final_viscosity_zone)

    matrix_df = matrix_df[
        (matrix_df["Worker_Viscosity_Zone"] != "Unknown") & 
        (matrix_df["Target_Viscosity_Zone"] != "Unknown")
    ].copy()

    # -----------------------------------------------------
    # 3. 物理飽和分析 (Physical Saturation Analysis by Initial Zone)
    # -----------------------------------------------------
    saturation_summary = []
    group_cols_initial = ["Resin", "Position_UI", "Vendor", "Solvent_Type", "Worker_Viscosity_Zone"]
    initial_keys = matrix_df[group_cols_initial].drop_duplicates()

    for _, key_row in initial_keys.iterrows():
        zone_df = matrix_df[
            (matrix_df["Resin"] == key_row["Resin"]) & 
            (matrix_df["Position_UI"] == key_row["Position_UI"]) &
            (matrix_df["Vendor"] == key_row["Vendor"]) & 
            (matrix_df["Solvent_Type"] == key_row["Solvent_Type"]) &
            (matrix_df["Worker_Viscosity_Zone"] == key_row["Worker_Viscosity_Zone"])
        ].copy()
        
        sat_result = build_saturation_profile(zone_df)
        safe_limits = get_safety_limits(zone_df, sat_result)
        
        saturation_summary.append({
            "Resin": key_row["Resin"], "Position_UI": key_row["Position_UI"],
            "Vendor": key_row["Vendor"], "Solvent_Type": key_row["Solvent_Type"],
            "Worker_Viscosity_Zone": key_row["Worker_Viscosity_Zone"],
            "Phys_Sat_Warning": safe_limits["warning_ratio"],
            "Phys_Sat_Stop": safe_limits["stop_ratio"]
        })

    sat_df = pd.DataFrame(saturation_summary)

    # -----------------------------------------------------
    # 4. 依據「初始 + 目標」雙重分組計算基礎統計值
    # -----------------------------------------------------
    group_cols_target = [
        "Resin", "Position_UI", "Vendor", "Solvent_Type", 
        "Worker_Viscosity_Zone", "Target_Viscosity_Zone"
    ]

    worker_sop = (
        matrix_df.groupby(group_cols_target, observed=False)
        .agg(
            Adjustment_Records=("塗料批號", "size"),
            Ref_Sensitivity=("Sensitivity", "median"), # Critical factor for V2.0
            Final_Visc_Lower=("黏度(秒)_1", get_spc_lower_bound), 
            Final_Visc_Upper=("黏度(秒)_1", get_spc_upper_bound), 
            Temperature_P25=("溫度", lambda x: x.quantile(0.25)),
            Temperature_P75=("溫度", lambda x: x.quantile(0.75))
        ).reset_index()
    )

    worker_sop = worker_sop[worker_sop["Adjustment_Records"] >= MIN_REFERENCE_RECORDS].copy()

    # -----------------------------------------------------
    # 5. V2.0 DELTA VISCOSITY ALGORITHM (物理降幅核心計算)
    # -----------------------------------------------------
    def get_zone_center(zone_str, zone_type="initial"):
        if zone_type == "initial":
            z = str(zone_str)
            if z == "<=70": return 60.0
            elif z == "71-90": return 80.5
            elif z == "91-110": return 100.5
            elif z == "111-130": return 120.5
            elif z.startswith("130-"):
                try: return (130 + float(z.split("-")[1])) / 2
                except: return 140.0
            elif z == ">130": return 140.0
        else: # target (e.g. 31~40)
            try:
                parts = str(zone_str).split("~")
                return (float(parts[0]) + float(parts[1])) / 2
            except: return np.nan
        return np.nan

    # 擷取區間中心值
    worker_sop["Initial_Center"] = worker_sop["Worker_Viscosity_Zone"].apply(lambda x: get_zone_center(x, "initial"))
    worker_sop["Target_Center"] = worker_sop["Target_Viscosity_Zone"].apply(lambda x: get_zone_center(x, "target"))
    
    # 計算預期降幅 Delta V
    worker_sop["Expected_Delta_V"] = worker_sop["Initial_Center"] - worker_sop["Target_Center"]

    # 過濾掉不合理的降幅 (例如目標比初始還高)
    worker_sop = worker_sop[worker_sop["Expected_Delta_V"] > 0].copy()

    # 計算物理理論比例 = 降幅 / 稀釋效率
    worker_sop["Theoretical_Ratio"] = worker_sop["Expected_Delta_V"] / worker_sop["Ref_Sensitivity"]

    # 合併物理飽和極限 (來自初始塗料屬性，而非目標經驗)
    worker_sop = worker_sop.merge(sat_df, on=group_cols_initial, how="left")

    # 警戒極限 = 物理飽和警戒 (若無資料則用 1.3 倍理論值)
    worker_sop["Saturation_Warning_Ratio"] = worker_sop["Phys_Sat_Warning"].fillna(worker_sop["Theoretical_Ratio"] * 1.3)

    # 停止極限 = 物理飽和停止 (若無資料則用 1.5 倍理論值)
    worker_sop["Saturation_Stop_Ratio"] = worker_sop["Phys_Sat_Stop"].fillna(worker_sop["Theoretical_Ratio"] * 1.5)
    
    # 確保停止極限 >= 警戒極限
    worker_sop["Saturation_Stop_Ratio"] = np.maximum(worker_sop["Saturation_Stop_Ratio"], worker_sop["Saturation_Warning_Ratio"])

    # -----------------------------------------------------
    # 5.1 區分「預估總需求」與「第一次安全添加」
    # -----------------------------------------------------
    # 預估累積總需求：反映不同目標黏度所需的完整降黏需求。
    # 目標越低、需要下降的黏度越大，預估累積需求比例應越高。
    worker_sop["Estimated_Total_Ratio"] = worker_sop["Theoretical_Ratio"].clip(lower=0)

    # 第一次添加只是安全測試步驟，不代表全部需求。
    # 三重限制：
    # 1. 不超過預估總需求的 50%
    # 2. 不超過正常區單次最大 3%
    # 3. 不超過警戒比例的 50%
    worker_sop["First_Add_Ratio"] = np.minimum.reduce([
        (worker_sop["Estimated_Total_Ratio"] * FIRST_ADD_WARNING_FRACTION).to_numpy(dtype=float),
        np.full(len(worker_sop), STEP_MAX_RATIO_NORMAL, dtype=float),
        (worker_sop["Saturation_Warning_Ratio"] * FIRST_ADD_WARNING_FRACTION).to_numpy(dtype=float)
    ])

    # 第一次添加後仍可能需要的預估比例。
    # 實際第二次添加仍必須回到 Tab 2，依第一次量測結果重新計算。
    worker_sop["Estimated_Remaining_Ratio"] = np.maximum(
        worker_sop["Estimated_Total_Ratio"] - worker_sop["First_Add_Ratio"],
        0.0
    )

    # -----------------------------------------------------
    # 6. Format Data & UI Output 
    # -----------------------------------------------------
    worker_sop["Historical_Final_Visc_Range"] = worker_sop.apply(
        lambda row: format_range(row["Final_Visc_Lower"], row["Final_Visc_Upper"], decimals=1), axis=1
    )

    worker_sop["Historical_Temp_Range"] = worker_sop.apply(
        lambda row: format_range(row["Temperature_P25"], row["Temperature_P75"], decimals=1), axis=1
    )

    worker_sop["塗裝位置"] = worker_sop["Position_UI"].map({
        "Primer": "底漆 (P)", "Top Finish": "正面漆 (TF)", "Back Finish": "背面漆 (BF)"
    }).fillna(worker_sop["Position_UI"])

    worker_output = worker_sop[[
        "Resin", 
        "塗裝位置", 
        "Vendor", 
        "Solvent_Type", 
        "Worker_Viscosity_Zone",
        "Target_Viscosity_Zone",
        "Ref_Sensitivity",
        "Expected_Delta_V",
        "Estimated_Total_Ratio",
        "First_Add_Ratio",
        "Estimated_Remaining_Ratio",
        "Historical_Final_Visc_Range",
        "Historical_Temp_Range",
        "Saturation_Warning_Ratio",
        "Saturation_Stop_Ratio"
    ]].copy()

    worker_output = worker_output.rename(columns={
        "Resin": "樹脂種類", 
        "Vendor": "塗料供應商", 
        "Solvent_Type": "稀釋劑種類",
        "Worker_Viscosity_Zone": "初始黏度區間", 
        "Target_Viscosity_Zone": "目標黏度區間",
        "Ref_Sensitivity": "參考稀釋效率",
        "Expected_Delta_V": "預計降黏幅度(秒)",
        "Estimated_Total_Ratio": "預估累積需求比例(%)",
        "First_Add_Ratio": "建議首次添加比例(%)",
        "Estimated_Remaining_Ratio": "首次後預估剩餘比例(%)",
        "Historical_Final_Visc_Range": "歷史最終黏度範圍",
        "Historical_Temp_Range": "歷史參考溫度範圍",
        "Saturation_Warning_Ratio": "累積添加警戒比例(%)",
        "Saturation_Stop_Ratio": "累積添加停止比例(%)"
    })

    worker_output["_initial_order"] = worker_output["初始黏度區間"].apply(get_zone_order)
    worker_output["_target_order"] = worker_output["目標黏度區間"].apply(get_final_zone_order)

    worker_output = worker_output.sort_values(
        by=["樹脂種類", "塗裝位置", "塗料供應商", "稀釋劑種類", "_initial_order", "_target_order"]
    ).drop(columns=["_initial_order", "_target_order"]).reset_index(drop=True)

    st.dataframe(
        worker_output,
        column_config={
            "初始黏度區間": st.column_config.TextColumn("初始黏度區間 (秒)"),
            "目標黏度區間": st.column_config.TextColumn("目標黏度區間 (秒)"),
            "參考稀釋效率": st.column_config.NumberColumn("參考稀釋效率 (s/%)", format="%.2f"),
            "預計降黏幅度(秒)": st.column_config.NumberColumn("預計降黏幅度 (秒)", format="%.1f"),
            "預估累積需求比例(%)": st.column_config.NumberColumn("預估累積需求比例 (%)", format="%.2f"),
            "建議首次添加比例(%)": st.column_config.NumberColumn("建議首次添加比例 (%)", format="%.2f"),
            "首次後預估剩餘比例(%)": st.column_config.NumberColumn("首次後預估剩餘比例 (%)", format="%.2f"),
            "歷史最終黏度範圍": st.column_config.TextColumn("歷史最終黏度範圍"),
            "歷史參考溫度範圍": st.column_config.TextColumn("歷史參考溫度範圍 (°C)"),
            "累積添加警戒比例(%)": st.column_config.NumberColumn("累積添加警戒比例(%)", format="%.2f"),
            "累積添加停止比例(%)": st.column_config.NumberColumn("累積添加停止比例(%)", format="%.2f")
        },
        use_container_width=True, hide_index=True
    )

    st.markdown("---")
    st.markdown("### 現場 SOP 使用方式")
    st.markdown(
        """
        1. 先確認樹脂、塗裝位置、供應商及稀釋劑種類。  
        2. 量測目前黏度，查詢相對應的「初始黏度區間」。  
        3. 依製程需求選擇「目標黏度區間」。 
        4. 「預估累積需求比例」代表依目前區間降至目標區間的完整理論需求，不可直接一次全部加入。  
        5. 「建議首次添加比例」僅為第一次安全添加量；不同目標可能因單次 3% 上限而顯示相同數值。  
        6. 添加量 = 原始塗料重量 × 建議首次添加比例(%) ÷ 100。  
        7. 一次添加後攪拌至少 5 分鐘，再重新量測確認。  
        8. 若仍高於目標上限，回到 Tab 2 依實測結果與閉迴路控制邏輯計算下一次添加量。  
        9. 累積添加比例不得超過「累積添加停止比例」。  
        """
    )

    # -----------------------------------------------------
    # 7. CSV Export
    # -----------------------------------------------------
    export_output = worker_output.copy()
    numeric_columns = [
        "預計降黏幅度(秒)",
        "預估累積需求比例(%)",
        "建議首次添加比例(%)",
        "首次後預估剩餘比例(%)",
        "累積添加警戒比例(%)",
        "累積添加停止比例(%)"
    ]

    for col in numeric_columns:
        export_output[col] = pd.to_numeric(export_output[col], errors="coerce").round(1)

    csv_export = export_output.to_csv(index=False, float_format="%.1f").encode("utf-8-sig")

    st.download_button(
        label="下載現場歷史加料參考表 CSV",
        data=csv_export,
        file_name="現場歷史加料參考表_V2_Delta_Model.csv",
        mime="text/csv"
    )
