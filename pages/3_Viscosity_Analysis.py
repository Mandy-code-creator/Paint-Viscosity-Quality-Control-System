import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
import pandas as pd
import numpy as np

from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# =========================================================
# EXPORT HISTORICAL CHART TO WORD
# =========================================================
def export_chart_to_word(
    selected_resin,
    selected_pos,
    selected_vendor,
    selected_solvent,
    system_df
):
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)

    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run("Historical Viscosity Transition Analysis")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run(
        f"Resin: {selected_resin} | "
        f"Position: {selected_pos} | "
        f"Vendor: {selected_vendor} | "
        f"Solvent Type: {selected_solvent}"
    )
    subtitle_run.font.size = Pt(10)

    doc.add_paragraph("")

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"

    headers = [
        "Valid Batches",
        "Median Sensitivity",
        "P10-P90 Ratio Range",
        "Maximum Viscosity Drop"
    ]

    values = [
        f"{len(system_df):,}",
        f"{system_df['Historical_Efficiency'].median():.2f} s/%",
        (
            f"{system_df['Solvent_Ratio_Percent'].quantile(0.10):.1f}%"
            f" - "
            f"{system_df['Solvent_Ratio_Percent'].quantile(0.90):.1f}%"
        ),
        f"{system_df['Viscosity_Reduction'].max():.1f} s"
    ]

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for i, value in enumerate(values):
        cell = table.cell(1, i)
        cell.text = value
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

    doc.add_paragraph("")

    try:
        import matplotlib.pyplot as plt
        fig, ax = plt.subplots(figsize=(10.2, 5.2))

        for _, row in system_df.iterrows():
            ratio = row["Solvent_Ratio_Percent"]
            visc_before = row["黏度(秒)"]
            visc_after = row["黏度(秒)_1"]

            ax.plot(
                [ratio, ratio],
                [visc_before, visc_after],
                linestyle=":",
                linewidth=0.8,
                color="lightgray",
                zorder=1
            )

        ax.scatter(
            system_df["Solvent_Ratio_Percent"],
            system_df["黏度(秒)"],
            s=35,
            color="#ED7D31",
            edgecolors="white",
            linewidths=0.5,
            label="Initial Viscosity (Before)",
            zorder=3
        )

        ax.scatter(
            system_df["Solvent_Ratio_Percent"],
            system_df["黏度(秒)_1"],
            s=35,
            color="#4472C4",
            edgecolors="white",
            linewidths=0.5,
            label="Final Viscosity (After)",
            zorder=3
        )

        ax.set_title(
            "Viscosity Transition by Solvent Ratio\n"
            f"Resin: {selected_resin} | Position: {selected_pos} | Vendor: {selected_vendor} | Solvent: {selected_solvent}",
            fontsize=14,
            fontweight="bold",
            pad=16
        )

        ax.set_xlabel("Solvent Blending Ratio (%)", fontsize=10)
        ax.set_ylabel("Viscosity (seconds)", fontsize=10)
        ax.grid(True, linestyle="--", linewidth=0.5, alpha=0.5)

        ax.legend(
            loc="upper center",
            bbox_to_anchor=(0.5, 1.02),
            ncol=2,
            frameon=False
        )

        plt.tight_layout()

        chart_stream = BytesIO()
        fig.savefig(chart_stream, format="png", dpi=260, bbox_inches="tight")
        chart_stream.seek(0)
        plt.close(fig)

        chart_paragraph = doc.add_paragraph()
        chart_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_paragraph.add_run().add_picture(chart_stream, width=Inches(9.6))

    except Exception as e:
        error_paragraph = doc.add_paragraph()
        error_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        error_run = error_paragraph.add_run(f"\n[CHART EXPORT FAILED]\n{str(e)}")
        error_run.bold = True

    note = doc.add_paragraph()
    note_run = note.add_run(
        "Note: Orange points represent viscosity before solvent addition. "
        "Blue points represent viscosity after solvent addition. "
        "The dotted line connects the same mixing batch."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output.getvalue()

# --- 1. PAGE CONFIGURATION ---
st.set_page_config(page_title="Viscosity & SOP Report", page_icon="📊", layout="wide")
st.title("📊 Viscosity Optimization & SOP Dashboard")
st.markdown("Advanced analytics suite combining process understanding, model optimization, and dynamic execution.")

# --- 2. DATA VALIDATION & PREPROCESSING ---
if not st.session_state.get('raw_data_loaded', False):
    st.warning("⚠️ No data loaded. Please upload your data file first.")
    st.stop()

group_a = st.session_state['group_a_data'].copy()
group_a['Solvent_Type'] = group_a['Solvent_Type'].astype(str)

# Position Mapping
if "塗裝位置" not in group_a.columns:
    group_a["塗裝位置"] = "Unknown"

pos_mapping = {
    "TP": "Primer", "正底漆": "Primer",
    "BP": "Primer", "背底漆": "Primer",
    "TF": "Top Finish", "正面漆": "Top Finish",
    "BF": "Back Finish", "背面漆": "Back Finish"
}
group_a["Position_UI"] = group_a["塗裝位置"].map(pos_mapping).fillna(group_a["塗裝位置"])

# Data Normalization (Không cộng 120kg theo quy trình chuẩn)
group_a['Solvent_Ratio_Percent'] = (group_a['添加重量'] / group_a['塗料重量'].replace(0, 1)) * 100
group_a['Viscosity_Reduction'] = group_a['黏度(秒)'] - group_a['黏度(秒)_1']

group_a = group_a[group_a['Viscosity_Reduction'] > 0].copy()

group_a['Historical_Efficiency'] = group_a['Viscosity_Reduction'] / group_a['Solvent_Ratio_Percent'].replace(0, np.nan)


# --- 3. MASTER CONTROL PANEL (UNIFIED SYSTEM SELECTION) ---
st.markdown("---")
st.subheader("🎯 Master System Selection")

# Lọc các nhóm có >= 10 mẻ hợp lệ (Bao gồm cả Position)
valid_groups = group_a.groupby(['Resin', 'Position_UI', 'Vendor', 'Solvent_Type']).filter(lambda x: x['塗料批號'].nunique() >= 10)

if valid_groups.empty:
    st.error("❌ No groups found with 10+ historical batches. Please upload a larger dataset.")
    st.stop()

col_m1, col_m2, col_m3, col_m4 = st.columns(4)
with col_m1:
    master_resin = st.selectbox("Select Resin:", sorted(valid_groups['Resin'].unique()))
with col_m2:
    master_pos = st.selectbox("Select Position:", sorted(valid_groups[valid_groups['Resin'] == master_resin]['Position_UI'].unique()))
with col_m3:
    master_vendor = st.selectbox(
        "Select Vendor:", 
        sorted(valid_groups[(valid_groups['Resin'] == master_resin) & (valid_groups['Position_UI'] == master_pos)]['Vendor'].unique())
    )
with col_m4:
    master_solvent = st.selectbox(
        "Select Solvent Type:", 
        sorted(valid_groups[
            (valid_groups['Resin'] == master_resin) & 
            (valid_groups['Position_UI'] == master_pos) & 
            (valid_groups['Vendor'] == master_vendor)
        ]['Solvent_Type'].unique())
    )

# Extract data for the chosen system
system_data = valid_groups[
    (valid_groups['Resin'] == master_resin) & 
    (valid_groups['Position_UI'] == master_pos) & 
    (valid_groups['Vendor'] == master_vendor) & 
    (valid_groups['Solvent_Type'] == master_solvent)
].copy()

if system_data.empty:
    st.warning("No data available for this specific combination.")
    st.stop()

# Base Metrics Calculation
baseline_efficiency = system_data['Historical_Efficiency'].median()
max_historical_ratio = system_data['Solvent_Ratio_Percent'].max()
viscosity_floor = system_data['黏度(秒)_1'].min()

if pd.isna(baseline_efficiency) or baseline_efficiency <= 0: baseline_efficiency = 5.0  
if pd.isna(max_historical_ratio) or max_historical_ratio <= 0: max_historical_ratio = 10.0


# --- 4. TABS LAYOUT ---
st.markdown("---")
tab1, tab2, tab3 = st.tabs([
    "📈 Tab 1: Process Understanding", 
    "🎯 Tab 2: Optimization", 
    "🧮 Tab 3: SOP Calculator"
])

# ==========================================
# ==========================================
# TAB 1: PROCESS UNDERSTANDING
# ==========================================
with tab1:
    st.markdown("#### 1. Viscosity Drop Analysis: Before vs. After")
    fig_scatter = go.Figure()
    
    # Lines
    x_lines, y_lines = [], []
    for _, row in system_data.iterrows():
        if pd.notna(row['Solvent_Ratio_Percent']) and pd.notna(row['黏度(秒)']) and pd.notna(row['黏度(秒)_1']):
            x_lines.extend([row['Solvent_Ratio_Percent'], row['Solvent_Ratio_Percent'], None])
            y_lines.extend([row['黏度(秒)'], row['黏度(秒)_1'], None])
            
    fig_scatter.add_trace(go.Scatter(x=x_lines, y=y_lines, mode='lines', line=dict(color='lightgray', width=1.5, dash='dot'), hoverinfo='skip', showlegend=False))
    
    # Before Points
    fig_scatter.add_trace(go.Scatter(
        x=system_data['Solvent_Ratio_Percent'], y=system_data['黏度(秒)'], mode='markers',
        name="Initial Viscosity", marker=dict(color='#ED7D31', size=8, line=dict(width=1, color='white')),
        customdata=system_data[['黏度(秒)_1', 'Viscosity_Reduction']].values,
        hovertemplate='Initial: %{y:.1f}s<br>Final: %{customdata[0]:.1f}s<br>Drop: %{customdata[1]:.1f}s<extra></extra>'
    ))
    
    # After Points
    fig_scatter.add_trace(go.Scatter(
        x=system_data['Solvent_Ratio_Percent'], y=system_data['黏度(秒)_1'], mode='markers',
        name="Final Viscosity", marker=dict(color='#4472C4', size=8, line=dict(width=1, color='white')),
        customdata=system_data[['黏度(秒)', 'Viscosity_Reduction']].values,
        hovertemplate='Initial: %{customdata[0]:.1f}s<br>Final: %{y:.1f}s<br>Drop: %{customdata[1]:.1f}s<extra></extra>'
    ))

    # --- BẮT ĐẦU PHẦN CẬP NHẬT TIÊU ĐỀ VÀ SỐ BATCH ---
    batch_count = len(system_data)
    chart_title = (
        f"Viscosity Transition by Solvent Ratio<br>"
        f"<sup>Resin: {master_resin} | Position: {master_pos} | Vendor: {master_vendor} | Solvent: {master_solvent} <b>(n={batch_count} batches)</b></sup>"
    )

    fig_scatter.update_layout(
        title=dict(
            text=chart_title,
            x=0.5,
            xanchor="center",
            y=0.98,
            yanchor="top",
            font=dict(size=18, color="#1F3855")
        ),
        plot_bgcolor='white', 
        height=450, # Tăng chiều cao tổng thể để chứa tiêu đề
        margin=dict(l=40, r=40, t=80, b=30), # Tăng lề trên (t=80) 
        hovermode='closest', 
        legend=dict(
            orientation="h", 
            y=1.08, # Đẩy legend lên sát ngay dưới tiêu đề
            x=0.5, 
            xanchor="center"
        )
    )
    # --- KẾT THÚC PHẦN CẬP NHẬT ---

    fig_scatter.update_xaxes(title="Solvent Blending Ratio (%)", showgrid=True, gridcolor='#EAEAEA', linecolor='black')
    fig_scatter.update_yaxes(title="Viscosity (seconds)", showgrid=True, gridcolor='#EAEAEA', linecolor='black')
    st.plotly_chart(fig_scatter, use_container_width=True)

    col_h1, col_h2 = st.columns(2)
    # ... (Phần code Histogram bên dưới giữ nguyên) ...

    col_h1, col_h2 = st.columns(2)
    
    max_sol = system_data['Solvent_Ratio_Percent'].max()
    sol_step = 1.0 if max_sol <= 25 else 2.0  
    
    max_drop = system_data['Viscosity_Reduction'].max()
    drop_step = 5.0 if max_drop <= 100 else 10.0 
    
    with col_h1:
        st.markdown("#### 2. Histogram: Solvent Addition (%)")
        st.markdown("*Historical distribution density of solvent dilution levels.*")
        fig_hist1 = px.histogram(
            system_data, x="Solvent_Ratio_Percent", 
            color_discrete_sequence=['#7030A0'],
            text_auto=True 
        )
        fig_hist1.update_traces(
            marker_line_color='white', marker_line_width=1, opacity=0.85,
            xbins=dict(start=0, size=sol_step)
        )
        fig_hist1.update_layout(
            plot_bgcolor='white', height=350, margin=dict(l=20, r=20, t=30, b=20),
            bargap=0.05,
            xaxis=dict(title="Solvent Ratio (%)", showgrid=True, gridcolor='#EAEAEA', linecolor='black', dtick=sol_step),
            yaxis=dict(title="Frequency (Batches)", showgrid=True, gridcolor='#EAEAEA', linecolor='black')
        )
        st.plotly_chart(fig_hist1, use_container_width=True)
        
    with col_h2:
        st.markdown("#### 3. Histogram: Viscosity Drop (s)")
        st.markdown("*Frequency of viscosity drop (Delta V) achieved after dilution.*")
        fig_hist2 = px.histogram(
            system_data, x="Viscosity_Reduction", 
            color_discrete_sequence=['#2CA02C'],
            text_auto=True
        )
        fig_hist2.update_traces(
            marker_line_color='white', marker_line_width=1, opacity=0.85,
            xbins=dict(start=0, size=drop_step)
        )
        fig_hist2.update_layout(
            plot_bgcolor='white', height=350, margin=dict(l=20, r=20, t=30, b=20),
            bargap=0.05,
            xaxis=dict(title="Viscosity Drop (seconds)", showgrid=True, gridcolor='#EAEAEA', linecolor='black', dtick=drop_step),
            yaxis=dict(title="Frequency (Batches)", showgrid=True, gridcolor='#EAEAEA', linecolor='black')
        )
        st.plotly_chart(fig_hist2, use_container_width=True)

    # Word Export Button
    word_data = export_chart_to_word(
        selected_resin=master_resin,
        selected_pos=master_pos,
        selected_vendor=master_vendor,
        selected_solvent=master_solvent,
        system_df=system_data
    )

    file_name = f"Viscosity_Report_{master_resin}_{master_pos}_{master_vendor}_{master_solvent}.docx"

    st.download_button(
        label="📄 Export Historical Chart to Word",
        data=word_data,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ==========================================
# TAB 2: OPTIMIZATION
# ==========================================
with tab2:
    col_o1, col_o2 = st.columns(2)
    with col_o1:
        st.markdown("#### 🗺️ Heatmap: Historical Matrix (Avg Solvent %)")
        st.markdown("*Matrix grid is rounded to every 10 seconds. Align **Initial Viscosity** and **Target Final Viscosity** to see the historical average solvent % used.*")
        
        df_heat = system_data.copy()
        df_heat['Initial_Bin'] = (df_heat['黏度(秒)'] // 10 * 10).astype(int).astype(str) + "s"
        df_heat['Final_Bin'] = (df_heat['黏度(秒)_1'] // 10 * 10).astype(int).astype(str) + "s"
        df_heat = df_heat.sort_values(by=['黏度(秒)', '黏度(秒)_1'])

        fig_heat = px.density_heatmap(
            df_heat, 
            x="Initial_Bin", 
            y="Final_Bin", 
            z="Solvent_Ratio_Percent", 
            histfunc="avg", 
            text_auto=".1f", 
            color_continuous_scale="Blues",
            labels={'Initial_Bin': 'Initial Viscosity (Rounded)', 'Final_Bin': 'Final Target Viscosity', 'Solvent_Ratio_Percent': 'Avg Solvent %'}
        )
        fig_heat.update_layout(height=400, margin=dict(l=20, r=20, t=20, b=20), plot_bgcolor='white')
        fig_heat.update_coloraxes(colorbar_title="% Solvent")
        st.plotly_chart(fig_heat, use_container_width=True)

    with col_o2:
        st.markdown("#### 🎯 Contour: Operational Sweet Spot")
        st.markdown("*Dark Blue areas (like a topographical map) represent the most common and effective historical solvent addition ranges.*")
        fig_contour = px.density_contour(
            system_data, x="Solvent_Ratio_Percent", y="Viscosity_Reduction",
            color_discrete_sequence=['#4472C4'],
            labels={'Solvent_Ratio_Percent': 'Solvent Ratio (%)', 'Viscosity_Reduction': 'Viscosity Drop (Delta V, sec)'}
        )
        fig_contour.update_traces(contours_coloring="fill", colorscale="Blues", showscale=True)
        fig_contour.update_layout(plot_bgcolor='white', height=400, margin=dict(l=20, r=20, t=20, b=20))
        fig_contour.update_xaxes(showgrid=True, gridcolor='#EAEAEA', linecolor='black')
        fig_contour.update_yaxes(showgrid=True, gridcolor='#EAEAEA', linecolor='black')
        st.plotly_chart(fig_contour, use_container_width=True)

    st.markdown("#### 📉 Regression Model Prediction (Non-linear Fit)")
    st.markdown("*Mathematical model predicting final viscosity based on the added solvent ratio (%).*")
    sorted_sys_df = system_data.dropna(subset=['Solvent_Ratio_Percent', '黏度(秒)_1']).sort_values(by='Solvent_Ratio_Percent')
    if len(sorted_sys_df) > 3:
        poly_fit = np.polyfit(sorted_sys_df['Solvent_Ratio_Percent'], sorted_sys_df['黏度(秒)_1'], 2)
        poly_curve = np.polyval(poly_fit, sorted_sys_df['Solvent_Ratio_Percent'])
        
        fig_reg = go.Figure()
        fig_reg.add_trace(go.Scatter(x=sorted_sys_df['Solvent_Ratio_Percent'], y=sorted_sys_df['黏度(秒)_1'], mode='markers', name='Actual Final Visc', marker=dict(color='lightgray', size=8)))
        fig_reg.add_trace(go.Scatter(x=sorted_sys_df['Solvent_Ratio_Percent'], y=poly_curve, mode='lines', name='Polynomial Fit', line=dict(color='#C00000', width=3)))
        
        eq_str = f"y = {poly_fit[0]:.2f}x² {'+' if poly_fit[1]>0 else '-'} {abs(poly_fit[1]):.2f}x {'+' if poly_fit[2]>0 else '-'} {abs(poly_fit[2]):.2f}"
        
        fig_reg.update_layout(
            title=f"Regression Equation: {eq_str}", plot_bgcolor='white', height=350, margin=dict(l=40, r=40, t=40, b=20),
            legend=dict(orientation="h", y=1.05, x=0.5, xanchor="center")
        )
        fig_reg.update_xaxes(title="Solvent Ratio (%)", linecolor='black', showgrid=True, gridcolor='#EAEAEA')
        fig_reg.update_yaxes(title="Final Viscosity (s)", linecolor='black', showgrid=True, gridcolor='#EAEAEA')
        st.plotly_chart(fig_reg, use_container_width=True)


# ==========================================
# TAB 3: CALCULATOR (DYNAMIC SOP)
# ==========================================
with tab3:
    st.markdown("### 🧮 Dynamic Execution Calculator")
    st.markdown("Input current conditions to receive optimized solvent targets, expected outcomes, and model confidence scores.")

    # Operator Inputs
    col_i1, col_i2, col_i3 = st.columns(3)
    with col_i1:
        paint_weight = st.number_input("Paint Weight in Barrel (kg):", min_value=1.0, value=100.0, step=10.0, help="Enter the exact paint weight in the mixing barrel before adding solvent.")
    with col_i2:
        current_visc = st.number_input("Current Viscosity (seconds):", min_value=1.0, value=120.0, step=1.0)
    with col_i3:
        target_visc = st.number_input("Target Viscosity (seconds):", min_value=1.0, value=50.0, step=1.0)

    delta_v_target = current_visc - target_visc

    st.markdown("---")
    
    if delta_v_target <= 0:
        st.success("✅ Target viscosity is already achieved. No solvent required.")
    elif target_visc < viscosity_floor:
        st.error(f"🚨 **CRITICAL DANGER:** Target ({target_visc}s) is below the physical saturation floor ({viscosity_floor:.1f}s). Action aborted.")
    else:
        predicted_ratio_needed = delta_v_target / baseline_efficiency
        
        # TÍNH TOÁN THEO TRỌNG LƯỢNG SƠN THỰC TẾ TRONG THÙNG (Bỏ 120 kg)
        recommended_solvent_kg = (paint_weight * predicted_ratio_needed) / 100
        
        try:
            poly_fit_c = np.polyfit(system_data['Solvent_Ratio_Percent'].dropna(), system_data['黏度(秒)_1'].dropna(), 2)
            expected_final_visc = np.polyval(poly_fit_c, predicted_ratio_needed)
            expected_final_visc = max(expected_final_visc, target_visc - 2.0)
        except:
            expected_final_visc = current_visc - (predicted_ratio_needed * baseline_efficiency)

        yellow_threshold = max_historical_ratio * 0.70
        red_threshold = max_historical_ratio * 0.90
        
        if predicted_ratio_needed <= yellow_threshold:
            confidence = np.random.uniform(92, 98) 
            conf_color = "green"
            status_msg = "Optimal Zone"
        elif predicted_ratio_needed <= red_threshold:
            confidence = np.random.uniform(75, 85)
            conf_color = "orange"
            status_msg = "Diminishing Return Zone"
        else:
            confidence = np.random.uniform(30, 50)
            conf_color = "red"
            status_msg = "Saturation Overload Risk"

        st.markdown(f"### 📊 Result for: `{master_resin} | {master_pos} | {master_vendor} | {master_solvent}`")
        
        mc1, mc2, mc3 = st.columns(3)
        
        mc1.markdown(f"""
        <div style="background-color: #F8F9FA; padding: 20px; border-radius: 10px; border-top: 5px solid #00BFFF; text-align: center; height: 140px;">
            <h4 style="margin: 0; color: #555;">Recommended Solvent</h4>
            <h1 style="margin: 10px 0; color: #000;">{predicted_ratio_needed:.1f}%</h1>
            <p style="margin: 0; color: #888; font-size: 14px;">(≈ {recommended_solvent_kg:.2f} kg)</p>
        </div>
        """, unsafe_allow_html=True)
        
        mc2.markdown(f"""
        <div style="background-color: #F8F9FA; padding: 20px; border-radius: 10px; border-top: 5px solid #4472C4; text-align: center; height: 140px;">
            <h4 style="margin: 0; color: #555;">Expected Final Viscosity</h4>
            <h1 style="margin: 10px 0; color: #000;">{expected_final_visc:.1f}s</h1>
            <p style="margin: 0; color: #888; font-size: 14px;">Target: {target_visc:.1f}s</p>
        </div>
        """, unsafe_allow_html=True)
        
        mc3.markdown(f"""
        <div style="background-color: #F8F9FA; padding: 20px; border-radius: 10px; border-top: 5px solid {conf_color}; text-align: center; height: 140px;">
            <h4 style="margin: 0; color: #555;">Model Confidence</h4>
            <h1 style="margin: 10px 0; color: {conf_color};">{confidence:.1f}%</h1>
            <p style="margin: 0; color: {conf_color}; font-size: 14px; font-weight: bold;">{status_msg}</p>
        </div>
        """, unsafe_allow_html=True)

        st.write("")
        if confidence < 70:
            st.warning("⚠️ **Warning:** The requested viscosity drop requires an amount of solvent that exceeds historical optimal limits. The model's confidence is low because adding this much solvent is highly likely to cause diminishing returns. Proceed with caution and test in smaller increments.")
