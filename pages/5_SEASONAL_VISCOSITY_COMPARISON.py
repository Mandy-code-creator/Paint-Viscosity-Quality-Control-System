import io

import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
import matplotlib.pyplot as plt

# Word report export
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# =========================================================
# 1. PAGE CONFIGURATION
# =========================================================
st.set_page_config(
    page_title="Seasonal Viscosity Comparison",
    page_icon="🌡️",
    layout="wide",
)

st.title("🌡️ 季節別黏度比較分析")
st.caption(
    "依色號與膜厚條件比較四季之添加前／後黏度、稀釋劑添加比例及溫度。"
)


# =========================================================
# 2. LOAD DATA
# =========================================================
if (
    not st.session_state.get("raw_data_loaded", False)
    or st.session_state.get("group_a_data") is None
):
    st.warning(
        "⚠️ 尚未載入資料，請先返回首頁上傳原始資料。"
    )
    st.stop()

df = st.session_state["group_a_data"].copy()


# =========================================================
# 3. HELPERS
# =========================================================
@st.cache_data(show_spinner=False)
def dataframe_to_csv_bytes(dataframe: pd.DataFrame) -> bytes:
    return dataframe.to_csv(index=False).encode("utf-8-sig")


def set_cell_shading(cell, fill):
    """Apply background color to a Word table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text_color(cell, color="FFFFFF", bold=False, size=8):
    """Format all runs inside a Word table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = None
            run.font.bold = bold
            run.font.size = Pt(size)
            r_pr = run._element.get_or_add_rPr()
            color_el = r_pr.find(qn("w:color"))
            if color_el is None:
                color_el = OxmlElement("w:color")
                r_pr.append(color_el)
            color_el.set(qn("w:val"), color)


def set_run_font(run, size=10, bold=False, color=None):
    """Use a Word-safe Traditional Chinese font."""
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_fonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    if color:
        color_el = r_pr.find(qn("w:color"))
        if color_el is None:
            color_el = OxmlElement("w:color")
            r_pr.append(color_el)
        color_el.set(qn("w:val"), color)


def add_report_paragraph(doc, text="", size=10, bold=False, align=None):
    paragraph = doc.add_paragraph()

    if align is not None:
        paragraph.alignment = align

    run = paragraph.add_run(str(text))
    set_run_font(
        run,
        size=size,
        bold=bold,
    )

    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.08
    return paragraph


def add_report_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()

    if level == 1:
        size = 16
        color = "1F4E78"
        before = 8
        after = 6
    else:
        size = 12
        color = "334155"
        before = 6
        after = 4

    run = paragraph.add_run(text)
    set_run_font(
        run,
        size=size,
        bold=True,
        color=color,
    )

    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    return paragraph


def add_key_value_table(doc, items, col_widths=(4.2, 18.2)):
    table = doc.add_table(
        rows=len(items),
        cols=2,
    )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for row_idx, (key, value) in enumerate(items):
        key_cell = table.rows[row_idx].cells[0]
        value_cell = table.rows[row_idx].cells[1]

        key_cell.text = str(key)
        value_cell.text = str(value)

        set_cell_shading(
            key_cell,
            "EAF2F8",
        )

        for cell, width_cm in zip(
            [key_cell, value_cell],
            col_widths,
        ):
            cell.width = Cm(width_cm)
            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=9,
                        bold=(cell is key_cell),
                    )

    doc.add_paragraph()
    return table


def add_dataframe_table(
    doc,
    dataframe,
    max_rows=None,
    font_size=7,
):
    """Add a compact management-style Word table."""
    if dataframe is None or dataframe.empty:
        add_report_paragraph(
            doc,
            "No data available.",
            size=9,
        )
        return None

    work_df = dataframe.copy()

    if max_rows is not None:
        work_df = work_df.head(max_rows)

    # Convert NaN to dash and limit long decimals.
    for col in work_df.columns:
        if pd.api.types.is_numeric_dtype(work_df[col]):
            work_df[col] = work_df[col].map(
                lambda value: (
                    "—"
                    if pd.isna(value)
                    else (
                        f"{value:,.2f}"
                        if isinstance(value, (float, np.floating))
                        else f"{value:,}"
                    )
                )
            )
        else:
            work_df[col] = (
                work_df[col]
                .fillna("—")
                .astype(str)
            )

    table = doc.add_table(
        rows=1,
        cols=len(work_df.columns),
    )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells

    for col_idx, column in enumerate(work_df.columns):
        header_cells[col_idx].text = str(column)
        set_cell_shading(
            header_cells[col_idx],
            "1F4E78",
        )
        header_cells[
            col_idx
        ].vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        for paragraph in header_cells[
            col_idx
        ].paragraphs:
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            for run in paragraph.runs:
                set_run_font(
                    run,
                    size=font_size,
                    bold=True,
                    color="FFFFFF",
                )

    for _, row in work_df.iterrows():
        cells = table.add_row().cells

        for col_idx, value in enumerate(row.tolist()):
            cells[col_idx].text = str(value)
            cells[
                col_idx
            ].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            for paragraph in cells[
                col_idx
            ].paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )

                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=font_size,
                    )

    doc.add_paragraph()
    return table



def save_matplotlib_to_buffer(fig):
    """Save matplotlib figure to an in-memory PNG buffer."""
    buffer = io.BytesIO()
    fig.savefig(
        buffer,
        format="png",
        dpi=220,
        bbox_inches="tight",
        facecolor="white",
        pad_inches=0.12,
    )
    plt.close(fig)
    buffer.seek(0)
    return buffer


def create_word_seasonal_overview_png():
    """Compact seasonal overview for Word export — no Kaleido required."""
    chart_df = season_summary.copy().reset_index(drop=True)

    fig, ax = plt.subplots(figsize=(10.8, 3.4), dpi=180)

    x = np.arange(len(chart_df))

    before = pd.to_numeric(
        chart_df["Median_Before_Viscosity"],
        errors="coerce",
    ).to_numpy(dtype=float)

    after = pd.to_numeric(
        chart_df["Median_After_Viscosity"],
        errors="coerce",
    ).to_numpy(dtype=float)

    ratio = pd.to_numeric(
        chart_df["Median_Solvent_Ratio"],
        errors="coerce",
    ).to_numpy(dtype=float)

    temperature = pd.to_numeric(
        chart_df["Median_Temperature"],
        errors="coerce",
    ).to_numpy(dtype=float)

    for i in range(len(chart_df)):
        ax.text(
            x[i],
            0.5,
            (
                f"{before[i]:.0f} → {after[i]:.0f} s\n"
                f"{ratio[i]:.1f}%\n"
                + (
                    f"{temperature[i]:.1f} °C"
                    if np.isfinite(temperature[i])
                    else "—"
                )
            ),
            ha="center",
            va="center",
            fontsize=11,
            fontweight="bold",
            bbox=dict(
                boxstyle="round,pad=0.55",
                facecolor="white",
                edgecolor="#94A3B8",
                linewidth=1.1,
            ),
        )

    ax.set_xlim(-0.5, len(chart_df) - 0.5)
    ax.set_ylim(0, 1)
    ax.set_xticks(x)
    ax.set_xticklabels(
        chart_df[period_col].astype(str),
        fontsize=10,
    )
    ax.set_yticks([])
    ax.set_xlabel("Season", fontsize=10)
    ax.set_title(
        f"{selected_paint_code} — Seasonal Viscosity Overview",
        fontsize=14,
        fontweight="bold",
        pad=18,
    )

    for spine in ax.spines.values():
        spine.set_visible(False)

    ax.set_facecolor("white")
    fig.patch.set_facecolor("white")
    fig.subplots_adjust(left=0.04, right=0.98, top=0.80, bottom=0.22)

    return save_matplotlib_to_buffer(fig)


def create_word_before_after_png():
    """Before vs After seasonal viscosity chart for Word export."""
    chart_df = season_summary.copy().reset_index(drop=True)

    fig, ax = plt.subplots(figsize=(10.8, 5.5), dpi=180)

    y = np.arange(len(chart_df))

    before = pd.to_numeric(
        chart_df["Median_Before_Viscosity"],
        errors="coerce",
    ).to_numpy(dtype=float)

    after = pd.to_numeric(
        chart_df["Median_After_Viscosity"],
        errors="coerce",
    ).to_numpy(dtype=float)

    for yi, b, a in zip(y, before, after):
        ax.plot(
            [a, b],
            [yi, yi],
            linewidth=4,
            color="#94A3B8",
            zorder=1,
        )

    ax.scatter(
        before,
        y,
        s=90,
        label="Before Viscosity",
        zorder=3,
    )

    ax.scatter(
        after,
        y,
        s=90,
        label="After Viscosity",
        zorder=3,
    )

    for yi, b, a in zip(y, before, after):
        ax.text(
            b,
            yi + 0.16,
            f"{b:.1f}",
            ha="center",
            va="bottom",
            fontsize=9,
        )
        ax.text(
            a,
            yi - 0.16,
            f"{a:.1f}",
            ha="center",
            va="top",
            fontsize=9,
        )

    ax.set_yticks(y)
    ax.set_yticklabels(
        chart_df[period_col].astype(str),
        fontsize=10,
    )
    ax.set_xlabel("Viscosity (s)", fontsize=11)
    ax.grid(axis="x", linewidth=0.8, alpha=0.35)
    ax.legend(
        loc="upper center",
        bbox_to_anchor=(0.5, 1.13),
        ncol=2,
        frameon=False,
    )
    ax.set_title(
        f"{selected_paint_code} — Seasonal Before vs After Viscosity",
        fontsize=14,
        fontweight="bold",
        pad=28,
    )

    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    fig.subplots_adjust(left=0.19, right=0.98, top=0.82, bottom=0.12)

    return save_matplotlib_to_buffer(fig)


def create_word_condition_png():
    """
    Seasonal viscosity + solvent ratio + temperature chart for Word export.
    Recreated with matplotlib; does not require Kaleido.
    """
    chart_df = season_summary.copy().reset_index(drop=True)

    fig, ax1 = plt.subplots(figsize=(10.8, 5.8), dpi=180)

    x = np.arange(len(chart_df))

    before = pd.to_numeric(
        chart_df["Median_Before_Viscosity"],
        errors="coerce",
    ).to_numpy(dtype=float)

    after = pd.to_numeric(
        chart_df["Median_After_Viscosity"],
        errors="coerce",
    ).to_numpy(dtype=float)

    ratio = pd.to_numeric(
        chart_df["Median_Solvent_Ratio"],
        errors="coerce",
    ).to_numpy(dtype=float)

    temp = pd.to_numeric(
        chart_df["Median_Temperature"],
        errors="coerce",
    ).to_numpy(dtype=float)

    ax1.plot(
        x,
        before,
        marker="o",
        linewidth=2.4,
        label="Before Viscosity (s)",
    )

    ax1.plot(
        x,
        after,
        marker="o",
        linewidth=2.4,
        label="After Viscosity (s)",
    )

    ax2 = ax1.twinx()

    ax2.plot(
        x,
        ratio,
        marker="D",
        linewidth=2.1,
        linestyle="--",
        label="Solvent Ratio (%)",
    )

    # Temperature is normalized to a separate line visually but shown
    # against the second axis to avoid adding a crowded third Word axis.
    if np.isfinite(temp).any():
        temp_min = np.nanmin(temp)
        temp_max = np.nanmax(temp)

        ratio_min = np.nanmin(ratio)
        ratio_max = np.nanmax(ratio)

        if temp_max > temp_min and ratio_max > ratio_min:
            temp_scaled = (
                (temp - temp_min)
                / (temp_max - temp_min)
                * (ratio_max - ratio_min)
                + ratio_min
            )
        else:
            temp_scaled = np.full_like(
                temp,
                np.nanmean(ratio),
            )

        ax2.plot(
            x,
            temp_scaled,
            marker="o",
            linewidth=2.0,
            linestyle=":",
            label="Temperature (°C)",
        )

        for xi, scaled, original in zip(
            x,
            temp_scaled,
            temp,
        ):
            if np.isfinite(original):
                ax2.text(
                    xi,
                    scaled,
                    f"{original:.1f}°",
                    fontsize=8,
                    ha="center",
                    va="bottom",
                )

    for xi, value in zip(x, ratio):
        if np.isfinite(value):
            ax2.text(
                xi,
                value,
                f"{value:.1f}%",
                fontsize=8,
                ha="center",
                va="bottom",
            )

    ax1.set_xticks(x)
    ax1.set_xticklabels(
        chart_df[period_col].astype(str),
        fontsize=9,
    )

    ax1.set_ylabel("Viscosity (s)", fontsize=10)
    ax2.set_ylabel(
        "Solvent Ratio (%) / Temperature trend",
        fontsize=10,
    )

    ax1.grid(axis="y", linewidth=0.8, alpha=0.35)

    handles1, labels1 = ax1.get_legend_handles_labels()
    handles2, labels2 = ax2.get_legend_handles_labels()

    ax1.legend(
        handles1 + handles2,
        labels1 + labels2,
        loc="upper center",
        bbox_to_anchor=(0.5, 1.16),
        ncol=4,
        frameon=False,
        fontsize=9,
    )

    ax1.set_title(
        f"{selected_paint_code} — Seasonal Viscosity, Solvent Ratio & Temperature",
        fontsize=14,
        fontweight="bold",
        pad=32,
    )

    fig.patch.set_facecolor("white")
    ax1.set_facecolor("white")
    fig.subplots_adjust(left=0.08, right=0.90, top=0.80, bottom=0.15)

    return save_matplotlib_to_buffer(fig)


def create_word_recommendation_png():
    """Current incoming vs recommended viscosity interval for Word."""
    if "final_p25" not in locals():
        return None

    fig, ax = plt.subplots(figsize=(10.8, 4.6), dpi=180)

    y_current = 1
    y_recommend = 0

    ax.plot(
        [current_before_p25, current_before_p75],
        [y_current, y_current],
        linewidth=14,
        solid_capstyle="round",
        label="Current Incoming P25–P75",
    )

    ax.scatter(
        [current_before_median],
        [y_current],
        s=100,
        marker="D",
        zorder=3,
    )

    ax.text(
        current_before_median,
        y_current + 0.13,
        f"{current_before_median:.1f} s",
        ha="center",
        va="bottom",
        fontsize=9,
    )

    ax.plot(
        [final_p25, final_p75],
        [y_recommend, y_recommend],
        linewidth=14,
        solid_capstyle="round",
        label="Recommended P25–P75",
    )

    ax.scatter(
        [final_median],
        [y_recommend],
        s=100,
        marker="D",
        zorder=3,
    )

    ax.text(
        final_median,
        y_recommend - 0.13,
        f"Target {final_median:.1f} s",
        ha="center",
        va="top",
        fontsize=9,
    )

    ax.text(
        final_p25,
        y_recommend + 0.13,
        f"Lower {final_p25:.1f}",
        ha="center",
        va="bottom",
        fontsize=8,
    )

    ax.text(
        final_p75,
        y_recommend + 0.13,
        f"Upper {final_p75:.1f}",
        ha="center",
        va="bottom",
        fontsize=8,
    )

    ax.set_yticks(
        [y_recommend, y_current]
    )
    ax.set_yticklabels(
        [
            "Recommended Incoming",
            "Current Incoming",
        ],
        fontsize=10,
    )

    ax.set_xlabel("Viscosity (s)", fontsize=10)
    ax.grid(axis="x", linewidth=0.8, alpha=0.35)
    ax.set_title(
        f"{selected_paint_code} — Current vs Recommended Incoming Viscosity",
        fontsize=14,
        fontweight="bold",
        pad=20,
    )

    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    fig.subplots_adjust(left=0.22, right=0.98, top=0.82, bottom=0.18)

    return save_matplotlib_to_buffer(fig)


def add_image_buffer_to_word(
    doc,
    image_buffer,
    caption,
    width_inches=9.6,
):
    add_report_paragraph(
        doc,
        caption,
        size=10,
        bold=True,
    )

    if image_buffer is None:
        add_report_paragraph(
            doc,
            "Chart data are unavailable for this section.",
            size=8,
        )
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    run.add_picture(
        image_buffer,
        width=Inches(width_inches),
    )

    paragraph.paragraph_format.space_after = Pt(6)


def create_management_word_report():
    if not HAS_DOCX:
        raise ImportError(
            "python-docx is not installed. "
            "Please add python-docx to requirements.txt."
        )

    doc = Document()

    # -----------------------------------------------------
    # Landscape A4 for management tables and charts
    # -----------------------------------------------------
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    # -----------------------------------------------------
    # Title
    # -----------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run(
        "季節別黏度與進料黏度最佳化分析報告"
    )
    set_run_font(
        title_run,
        size=18,
        bold=True,
        color="1F4E78",
    )

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run(
        f"Paint Code: {selected_paint_code}"
    )
    set_run_font(
        subtitle_run,
        size=11,
        color="64748B",
    )

    # -----------------------------------------------------
    # Basic information
    # -----------------------------------------------------
    add_key_value_table(
        doc,
        [
            ("供應商", selected_vendor),
            ("色號", selected_paint_code),
            ("塗裝位置", selected_position),
            ("膜厚組合", structure_display),
            ("樹脂", selected_resin),
            ("稀釋劑", selected_solvent),
            ("資料期間", f"{min_date} ～ {max_date}"),
            ("有效紀錄", f"{len(analysis_df):,} 筆"),
        ],
    )

    # -----------------------------------------------------
    # Section 1
    # -----------------------------------------------------
    add_report_heading(
        doc,
        "一、分析目的",
        level=1,
    )
    add_report_paragraph(
        doc,
        "本分析針對指定色號，依供應商、塗裝位置、膜厚組合、樹脂種類及稀釋劑條件，"
        "比較不同季節之黏度變化、稀釋劑使用及塗料使用量，並進一步評估免加稀釋劑"
        "直接上線之試驗進料黏度範圍。",
        size=10,
    )

    # -----------------------------------------------------
    # Section 2
    # -----------------------------------------------------
    add_report_heading(
        doc,
        "二、分析條件與方法",
        level=1,
    )
    add_report_paragraph(
        doc,
        "分析條件：Vendor → Paint Code → Position → Coating Structure → "
        "Resin → Solvent → Production Line。",
        size=10,
    )
    add_report_paragraph(
        doc,
        "季節分類：冬季 12–02 月、春季 03–05 月、夏季 06–08 月、秋季 09–11 月。",
        size=10,
    )
    add_report_paragraph(
        doc,
        "膜厚條件以完整塗層組合表示，例如 5 µm + 20 µm，避免不同訂單膜厚條件混合分析。",
        size=10,
    )

    # -----------------------------------------------------
    # Section 3 — Seasonal viscosity
    # -----------------------------------------------------
    add_report_heading(
        doc,
        "三、季節別黏度分析",
        level=1,
    )

    seasonal_kpi_df = pd.DataFrame(
        [
            {
                "指標": "最高添加前黏度",
                "期間": str(
                    highest_before_row[
                        period_col
                    ]
                ),
                "數值": (
                    f"{highest_before_row['Median_Before_Viscosity']:.1f} s"
                ),
            },
            {
                "指標": "最低添加前黏度",
                "期間": str(
                    lowest_before_row[
                        period_col
                    ]
                ),
                "數值": (
                    f"{lowest_before_row['Median_Before_Viscosity']:.1f} s"
                ),
            },
            {
                "指標": "最高添加比例",
                "期間": str(
                    highest_ratio_row[
                        period_col
                    ]
                ),
                "數值": (
                    f"{highest_ratio_row['Median_Solvent_Ratio']:.1f}%"
                ),
            },
            {
                "指標": "最大降黏幅度",
                "期間": str(
                    largest_drop_row[
                        period_col
                    ]
                ),
                "數值": (
                    f"{largest_drop_row['Median_Viscosity_Drop']:.1f} s"
                ),
            },
        ]
    )

    add_dataframe_table(
        doc,
        seasonal_kpi_df,
        font_size=8,
    )

    add_image_buffer_to_word(
        doc,
        create_word_seasonal_overview_png(),
        "圖1　Seasonal Viscosity Overview",
        width_inches=9.6,
    )

    add_image_buffer_to_word(
        doc,
        create_word_before_after_png(),
        "圖2　各季節添加前後黏度比較",
        width_inches=9.6,
    )

    add_report_paragraph(
        doc,
        f"季節黏度判讀：{seasonal_before_comment} {seasonal_after_comment}",
        size=10,
        bold=True,
    )

    # -----------------------------------------------------
    # Section 4 — Solvent + production volume
    # -----------------------------------------------------
    add_report_heading(
        doc,
        "四、稀釋劑、溫度與生產量分析",
        level=1,
    )

    add_report_paragraph(
        doc,
        "為避免將添加比例與稀釋劑總用量直接比較而產生誤判，"
        "本分析同時納入季節塗料使用量與加權添加比例。",
        size=10,
    )

    add_report_paragraph(
        doc,
        "加權添加比例 (%) = 季節稀釋劑總用量 ÷ 季節塗料使用量 × 100",
        size=10,
        bold=True,
    )

    add_image_buffer_to_word(
        doc,
        create_word_condition_png(),
        "圖3　各季節黏度、稀釋劑添加比例與溫度趨勢",
        width_inches=9.6,
    )

    add_dataframe_table(
        doc,
        season_display,
        font_size=6,
    )

    add_report_paragraph(
        doc,
        f"生產使用量最高期間為 {highest_production_period}，"
        f"塗料使用量約 {highest_production_kg:,.1f} kg。"
        "因此，稀釋劑總用量較高不一定代表添加比例較高，"
        "亦可能主要受到該季節生產量增加影響。",
        size=10,
        bold=True,
    )

    # -----------------------------------------------------
    # Section 5 — Optimization method
    # -----------------------------------------------------
    add_report_heading(
        doc,
        "五、進料黏度最佳化方法",
        level=1,
    )

    add_report_paragraph(
        doc,
        "以歷史添加後黏度作為現場實際可生產黏度之參考：",
        size=10,
    )
    add_report_paragraph(
        doc,
        "建議進料下限 = P25(Final Viscosity)",
        size=10,
        bold=True,
    )
    add_report_paragraph(
        doc,
        "建議進料目標 = Median(Final Viscosity)",
        size=10,
        bold=True,
    )
    add_report_paragraph(
        doc,
        "建議進料上限 = P75(Final Viscosity)",
        size=10,
        bold=True,
    )
    add_report_paragraph(
        doc,
        "安全性同時使用歷史資料量、Final Viscosity IQR 及 Seasonal Gap 評估。",
        size=10,
    )

    # -----------------------------------------------------
    # Section 6 — Recommendation
    # -----------------------------------------------------
    add_report_heading(
        doc,
        "六、建議進料黏度結果",
        level=1,
    )

    if "recommendation_table" in locals():
        compact_rec_df = recommendation_table[
            [
                "Paint Code",
                "Coating Structure",
                "Records",
                "Batches",
                "Recommended Lower (s)",
                "Recommended Target (s)",
                "Recommended Upper (s)",
                "Final Viscosity IQR (s)",
                "Seasonal Final Viscosity Gap (s)",
                "Recommendation",
            ]
        ].copy()

        add_dataframe_table(
            doc,
            compact_rec_df,
            font_size=6,
        )

    if "fig_recommend" in locals():
        add_image_buffer_to_word(
            doc,
            create_word_recommendation_png(),
            "圖4　目前進料黏度與建議試驗進料範圍比較",
            width_inches=9.6,
        )

    if "seasonal_final_display" in locals():
        add_report_heading(
            doc,
            "各季節添加後黏度驗證",
            level=2,
        )
        add_dataframe_table(
            doc,
            seasonal_final_display,
            font_size=7,
        )

    if "recommendation_status" in locals():
        add_report_paragraph(
            doc,
            f"系統判定：{recommendation_icon} {recommendation_status}",
            size=11,
            bold=True,
        )

        add_report_paragraph(
            doc,
            f"建議試驗範圍：{recommended_range_text}",
            size=11,
            bold=True,
        )

        add_report_paragraph(
            doc,
            recommendation_detail,
            size=10,
        )
    else:
        add_report_paragraph(
            doc,
            "目前資料不足，無法建立建議進料黏度範圍。",
            size=10,
        )

    # -----------------------------------------------------
    # Section 7 — Management recommendation
    # -----------------------------------------------------
    add_report_heading(
        doc,
        "七、管理建議",
        level=1,
    )

    add_report_paragraph(
        doc,
        "若判定為 Ready for No-Solvent Pilot，建議優先要求供應商依建議目標值"
        "進行小批量試配，並於現場確認是否可免加稀釋劑直接生產。",
        size=10,
    )

    add_report_paragraph(
        doc,
        "試驗時應同步確認膜厚、光澤、色差、表面品質及成品品質。"
        "驗證穩定後，再評估是否轉為正式進料黏度管制規格。",
        size=10,
    )

    add_report_paragraph(
        doc,
        "注意：本分析之建議進料範圍為試驗參考值，並非直接取代正式進料規格；"
        "正式規格仍須經供應商試配與產線實際驗證。",
        size=9,
        bold=True,
    )

    # -----------------------------------------------------
    # Section 8 — Conclusion
    # -----------------------------------------------------
    add_report_heading(
        doc,
        "八、結論",
        level=1,
    )

    add_report_paragraph(
        doc,
        "本分析透過季節差異、生產量、稀釋劑使用及添加後黏度分布，"
        "將現場黏度調整結果轉換為可供供應商試配之數據化進料條件。"
        "最終目的為降低現場稀釋劑添加需求、縮短調整時間，並提升不同季節"
        "與不同膜厚條件下之作業一致性。",
        size=10,
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def safe_unique(series):
    return sorted(
        [
            str(value).strip()
            for value in series.dropna().unique().tolist()
            if str(value).strip()
            not in {
                "",
                "Unknown",
                "UNKNOWN",
                "nan",
                "NaN",
                "None",
                "<NA>",
            }
        ]
    )


def fmt_thickness(value):
    if pd.isna(value):
        return "—"

    value = float(value)

    if value.is_integer():
        return str(int(value))

    return f"{value:.1f}".rstrip("0").rstrip(".")


def build_structure_label(row):
    primer = row["Primer_Thickness"]
    main = row["Main_Coat_Thickness"]

    if pd.isna(primer) and pd.isna(main):
        return "Unknown"

    return (
        f"{fmt_thickness(primer)} µm + "
        f"{fmt_thickness(main)} µm"
    )


def assign_season(month):
    if pd.isna(month):
        return np.nan

    month = int(month)

    if month in [12, 1, 2]:
        return "冬季 (12–02月)"

    if month in [3, 4, 5]:
        return "春季 (03–05月)"

    if month in [6, 7, 8]:
        return "夏季 (06–08月)"

    return "秋季 (09–11月)"


SEASON_ORDER = {
    "冬季 (12–02月)": 1,
    "春季 (03–05月)": 2,
    "夏季 (06–08月)": 3,
    "秋季 (09–11月)": 4,
}


# =========================================================
# 4. DATA PREPARATION
# =========================================================
text_cols = [
    "Vendor",
    "Resin",
    "Solvent_Type",
    "塗料批號",
    "線別",
    "塗裝位置",
]

for col in text_cols:
    if col not in df.columns:
        df[col] = "Unknown"

    df[col] = (
        df[col]
        .fillna("Unknown")
        .astype(str)
        .str.strip()
    )


df["Paint_Code"] = (
    df.get(
        "塗料編號",
        pd.Series("Unknown", index=df.index),
    )
    .fillna("Unknown")
    .astype(str)
    .str.strip()
    .str.upper()
)

df["Solvent_Type"] = (
    df["Solvent_Type"]
    .astype(str)
    .str.strip()
    .str.upper()
)

df["Batch_ID"] = (
    df["塗料批號"]
    .fillna("Unknown")
    .astype(str)
    .str.strip()
)

df["Bucket_Number"] = (
    df.get(
        "塗料桶號",
        pd.Series("Unknown", index=df.index),
    )
    .fillna("Unknown")
    .astype(str)
    .str.strip()
)


# ---------------------------------------------------------
# 4.1 COATING POSITION
# ---------------------------------------------------------
position_map = {
    "TP": "Primer",
    "正底漆": "Primer",
    "BP": "Primer",
    "背底漆": "Primer",
    "TF": "Top Finish",
    "正面漆": "Top Finish",
    "BF": "Back Finish",
    "背面漆": "Back Finish",
    "PRIMER": "Primer",
    "TOP FINISH": "Top Finish",
    "BACK FINISH": "Back Finish",
}

position_key = (
    df["塗裝位置"]
    .fillna("")
    .astype(str)
    .str.strip()
    .str.upper()
)

invalid_position_values = {
    "",
    "0",
    "0.0",
    "NAN",
    "NONE",
    "NULL",
    "N/A",
    "NA",
    "-",
    "--",
    "<NA>",
    "UNKNOWN",
}

position_key = position_key.mask(
    position_key.isin(invalid_position_values),
    "",
)

df["_Position_Raw"] = position_key

df["Position_UI"] = (
    position_key
    .map(position_map)
    .fillna("Unknown")
)


# ---------------------------------------------------------
# 4.2 SPECIAL EXCEPTION — PS30213Z2
# ---------------------------------------------------------
# PS30213Z2 is operationally a primer-group paint.
# Some multi-pass records may be marked as TF / Top Finish,
# but it must never appear when Top Finish is selected.
special_primer_paint_codes = {
    "PS30213Z2",
}

df.loc[
    df["Paint_Code"].isin(special_primer_paint_codes),
    "Position_UI",
] = "Primer"


# ---------------------------------------------------------
# 4.3 ORDER FILM THICKNESS
# ---------------------------------------------------------
thickness_cols = [
    "TOPFILM_THICK",
    "TTMFILM_THICK",
    "BACKFILM_THICK",
    "BTMFILM_THICK",
]

for col in thickness_cols:
    if col not in df.columns:
        df[col] = np.nan

    df[col] = pd.to_numeric(
        df[col],
        errors="coerce",
    )


# Determine coating side from raw source position.
is_top_side = df["_Position_Raw"].isin(
    [
        "TF",
        "TP",
        "正面漆",
        "正底漆",
        "TOP FINISH",
    ]
)

is_back_side = df["_Position_Raw"].isin(
    [
        "BF",
        "BP",
        "背面漆",
        "背底漆",
        "BACK FINISH",
    ]
)


df["Primer_Thickness"] = np.select(
    [
        is_top_side,
        is_back_side,
    ],
    [
        df["TTMFILM_THICK"],
        df["BTMFILM_THICK"],
    ],
    default=np.nan,
)

df["Main_Coat_Thickness"] = np.select(
    [
        is_top_side,
        is_back_side,
    ],
    [
        df["TOPFILM_THICK"],
        df["BACKFILM_THICK"],
    ],
    default=np.nan,
)

df["Total_Coating_Thickness"] = (
    df["Primer_Thickness"]
    + df["Main_Coat_Thickness"]
)

df["Coating_Structure"] = df.apply(
    build_structure_label,
    axis=1,
)


# ---------------------------------------------------------
# 4.4 NUMERIC FIELDS
# ---------------------------------------------------------
numeric_cols = [
    "塗料重量",
    "添加重量",
    "黏度(秒)",
    "黏度(秒)_1",
    "溫度",
]

for col in numeric_cols:
    if col not in df.columns:
        df[col] = np.nan

    df[col] = pd.to_numeric(
        df[col],
        errors="coerce",
    )


# ---------------------------------------------------------
# 4.5 CORE CALCULATIONS
# ---------------------------------------------------------
df["Base_Paint_kg"] = (
    df["塗料重量"]
    - df["添加重量"]
)

df["Delta_V"] = (
    df["黏度(秒)"]
    - df["黏度(秒)_1"]
)

df["Solvent_Ratio_Percent"] = np.where(
    df["Base_Paint_kg"] > 0,
    df["添加重量"]
    / df["Base_Paint_kg"]
    * 100,
    np.nan,
)

df["Viscosity_Sensitivity"] = np.where(
    df["Solvent_Ratio_Percent"] > 0,
    df["Delta_V"]
    / df["Solvent_Ratio_Percent"],
    np.nan,
)


# Group A strict valid records.
df = df[
    (df["Base_Paint_kg"] > 0)
    & (df["添加重量"] > 0)
    & (df["黏度(秒)"] > 0)
    & (df["黏度(秒)_1"] > 0)
    & (df["Delta_V"] > 0)
    & (df["Position_UI"].isin(
        [
            "Primer",
            "Top Finish",
            "Back Finish",
        ]
    ))
].copy()

if df.empty:
    st.warning(
        "⚠️ 資料清理後無有效黏度調整紀錄。"
    )
    st.stop()


# =========================================================
# 5. DATE / RECORD LOGIC
# =========================================================
date_candidates = [
    "攪拌日期",
    "調整日期",
    "生產日期",
    "Date",
]

date_col = next(
    (
        col
        for col in date_candidates
        if col in df.columns
    ),
    None,
)

if date_col is None:
    st.error(
        "❌ 找不到日期欄位。"
    )
    st.stop()

df["_Analysis_Date"] = pd.to_datetime(
    df[date_col],
    errors="coerce",
)

df = df[
    df["_Analysis_Date"].notna()
].copy()

if df.empty:
    st.warning(
        "⚠️ 日期格式無法解析。"
    )
    st.stop()


time_candidates = [
    "攪拌時間(迄)",
    "攪拌時間",
    "Time",
]

time_col = next(
    (
        col
        for col in time_candidates
        if col in df.columns
    ),
    None,
)

sort_cols = [
    "Batch_ID",
    "Bucket_Number",
    "_Analysis_Date",
]

if time_col is not None:
    sort_cols.append(time_col)

df = df.sort_values(
    sort_cols,
    ascending=True,
    na_position="last",
)


# PS30213X8: every row is an independent use/adjustment event.
special_record_codes = {
    "PS30213X8",
}

is_special_record = df[
    "Paint_Code"
].isin(
    special_record_codes
)

df_standard = (
    df.loc[~is_special_record]
    .drop_duplicates(
        subset=[
            "Batch_ID",
            "Bucket_Number",
        ],
        keep="last",
    )
    .copy()
)

df_special = (
    df.loc[is_special_record]
    .copy()
)

df = pd.concat(
    [
        df_standard,
        df_special,
    ],
    ignore_index=True,
)

df = df.sort_values(
    [
        "_Analysis_Date",
        "Batch_ID",
        "Bucket_Number",
    ],
    ascending=True,
    na_position="last",
).reset_index(drop=True)


# =========================================================
# 6. SEASON CLASSIFICATION
# =========================================================
df["Month"] = (
    df["_Analysis_Date"]
    .dt.month
)

df["Season"] = (
    df["Month"]
    .apply(assign_season)
)

df["Season_Order"] = (
    df["Season"]
    .map(SEASON_ORDER)
)

df["Season_Year"] = np.where(
    df["Month"] == 12,
    df["_Analysis_Date"].dt.year + 1,
    df["_Analysis_Date"].dt.year,
)

df["Season_Year"] = pd.Series(
    df["Season_Year"],
    index=df.index,
).astype("Int64")

df["Season_Period"] = (
    df["Season_Year"].astype(str)
    + " "
    + df["Season"].fillna("Unknown")
)


# =========================================================
# 7. LINKED FILTERS
#    Vendor → Paint Code → Position → Thickness
#    → Resin → Solvent → Line
# =========================================================
st.markdown("---")
st.subheader("🔍 分析篩選條件")

filter_source = df.copy()

row1_col1, row1_col2, row1_col3, row1_col4 = st.columns(4)
row2_col1, row2_col2, row2_col3, row2_col4 = st.columns(4)


# Vendor
vendor_options = safe_unique(
    filter_source["Vendor"]
)

if not vendor_options:
    st.warning("⚠️ 無供應商資料。")
    st.stop()

selected_vendor = row1_col1.selectbox(
    "Vendor (供應商)",
    vendor_options,
)

vendor_df = filter_source[
    filter_source["Vendor"] == selected_vendor
].copy()


# Paint Code
paint_code_options = safe_unique(
    vendor_df["Paint_Code"]
)

if not paint_code_options:
    st.warning(
        "⚠️ 此供應商無有效色號。"
    )
    st.stop()

selected_paint_code = row1_col2.selectbox(
    "Paint Code (色號)",
    paint_code_options,
)

paint_df = vendor_df[
    vendor_df["Paint_Code"] == selected_paint_code
].copy()


# Position
valid_position_order = [
    "Primer",
    "Top Finish",
    "Back Finish",
]

position_available = set(
    paint_df["Position_UI"]
    .dropna()
    .astype(str)
    .tolist()
)

position_options = [
    value
    for value in valid_position_order
    if value in position_available
]

if not position_options:
    st.warning(
        "⚠️ 此色號無有效塗裝位置。"
    )
    st.stop()

selected_position = row1_col3.selectbox(
    "Coating Position (塗裝位置)",
    position_options,
)

position_df = paint_df[
    paint_df["Position_UI"] == selected_position
].copy()


# Coating Structure
structure_options = safe_unique(
    position_df["Coating_Structure"]
)

structure_options = [
    value
    for value in structure_options
    if value != "Unknown"
]

if not structure_options:
    structure_options = ["All"]

selected_structure = row1_col4.selectbox(
    "Coating Structure (膜厚組合)",
    ["All"] + structure_options
    if structure_options != ["All"]
    else ["All"],
    help=(
        "Top side = Primer + Top Finish；"
        "Back side = Primer + Back Finish。"
        "例如：5 µm + 20 µm。"
    ),
)

if selected_structure != "All":
    structure_df = position_df[
        position_df["Coating_Structure"]
        == selected_structure
    ].copy()
else:
    structure_df = position_df.copy()


# Resin
resin_options = safe_unique(
    structure_df["Resin"]
)

if not resin_options:
    resin_options = ["All"]

selected_resin = row2_col1.selectbox(
    "Resin Type (樹脂種類)",
    ["All"] + resin_options
    if resin_options != ["All"]
    else ["All"],
)

if selected_resin != "All":
    resin_df = structure_df[
        structure_df["Resin"] == selected_resin
    ].copy()
else:
    resin_df = structure_df.copy()


# Solvent
solvent_options = safe_unique(
    resin_df["Solvent_Type"]
)

if not solvent_options:
    solvent_options = ["All"]

selected_solvent = row2_col2.selectbox(
    "Solvent Type (稀釋劑種類)",
    ["All"] + solvent_options
    if solvent_options != ["All"]
    else ["All"],
)

if selected_solvent != "All":
    solvent_df = resin_df[
        resin_df["Solvent_Type"]
        == selected_solvent
    ].copy()
else:
    solvent_df = resin_df.copy()


# Production Line
line_options = safe_unique(
    solvent_df["線別"]
)

selected_lines = row2_col3.multiselect(
    "Production Line (產線)",
    line_options,
    default=line_options,
)

if not selected_lines:
    st.warning(
        "⚠️ 請至少選擇一條產線。"
    )
    st.stop()

analysis_df = solvent_df[
    solvent_df["線別"].isin(
        selected_lines
    )
].copy()


# Analysis mode
analysis_mode = row2_col4.selectbox(
    "Analysis Mode (分析方式)",
    [
        "合併各年度比較四季",
        "依季節年度比較",
    ],
)

if analysis_df.empty:
    st.warning(
        "⚠️ 無符合目前篩選條件的資料。"
    )
    st.stop()


# =========================================================
# 8. FILTER / DATA SUMMARY
# =========================================================
min_date = (
    analysis_df["_Analysis_Date"]
    .min()
    .strftime("%Y-%m-%d")
)

max_date = (
    analysis_df["_Analysis_Date"]
    .max()
    .strftime("%Y-%m-%d")
)

unique_structures = safe_unique(
    analysis_df["Coating_Structure"]
)

if selected_structure == "All":
    if len(unique_structures) == 1:
        structure_display = unique_structures[0]
    elif len(unique_structures) <= 4:
        structure_display = " / ".join(
            unique_structures
        )
    else:
        structure_display = (
            f"{len(unique_structures)} structures"
        )
else:
    structure_display = selected_structure

filter_details = (
    f"Vendor: {selected_vendor} | "
    f"Paint Code: {selected_paint_code} | "
    f"Position: {selected_position} | "
    f"Thickness: {structure_display} | "
    f"Resin: {selected_resin} | "
    f"Solvent: {selected_solvent}"
)

st.info(
    f"📅 **資料期間：** {min_date} ➔ {max_date}"
    f" ｜ 📊 **有效紀錄：** {len(analysis_df):,} 筆"
    f" ｜ 🎨 **色號：** {selected_paint_code}"
    f" ｜ 📏 **膜厚：** {structure_display}"
)


# =========================================================
# 9. SEASONAL AGGREGATION
# =========================================================
if analysis_mode == "合併各年度比較四季":
    group_cols = [
        "Season_Order",
        "Season",
    ]
    period_col = "Season"
else:
    group_cols = [
        "Season_Year",
        "Season_Order",
        "Season",
        "Season_Period",
    ]
    period_col = "Season_Period"

season_summary = (
    analysis_df
    .groupby(
        group_cols,
        dropna=False,
    )
    .agg(
        Historical_Records=(
            "Paint_Code",
            "size",
        ),
        Historical_Batches=(
            "Batch_ID",
            "nunique",
        ),
        Median_Before_Viscosity=(
            "黏度(秒)",
            "median",
        ),
        Before_P25=(
            "黏度(秒)",
            lambda x: x.quantile(0.25),
        ),
        Before_P75=(
            "黏度(秒)",
            lambda x: x.quantile(0.75),
        ),
        Median_After_Viscosity=(
            "黏度(秒)_1",
            "median",
        ),
        After_P25=(
            "黏度(秒)_1",
            lambda x: x.quantile(0.25),
        ),
        After_P75=(
            "黏度(秒)_1",
            lambda x: x.quantile(0.75),
        ),
        Median_Viscosity_Drop=(
            "Delta_V",
            "median",
        ),
        Median_Solvent_Ratio=(
            "Solvent_Ratio_Percent",
            "median",
        ),
        Median_Temperature=(
            "溫度",
            "median",
        ),
        Total_Base_Paint_kg=(
            "Base_Paint_kg",
            "sum",
        ),
        Total_Solvent_kg=(
            "添加重量",
            "sum",
        ),
    )
    .reset_index()
)

# Weighted solvent ratio is calculated from seasonal totals.
# This makes the relationship between production volume and solvent use clear.
season_summary["Weighted_Solvent_Ratio"] = np.where(
    season_summary["Total_Base_Paint_kg"] > 0,
    season_summary["Total_Solvent_kg"]
    / season_summary["Total_Base_Paint_kg"]
    * 100,
    np.nan,
)

if analysis_mode == "合併各年度比較四季":
    season_summary = (
        season_summary
        .sort_values(
            ["Season_Order"]
        )
        .reset_index(drop=True)
    )
else:
    season_summary = (
        season_summary
        .sort_values(
            [
                "Season_Year",
                "Season_Order",
            ]
        )
        .reset_index(drop=True)
    )

if season_summary.empty:
    st.warning(
        "⚠️ 此條件無足夠季節資料。"
    )
    st.stop()


# =========================================================
# 10. COMPACT SEASONAL OVERVIEW
# =========================================================
st.markdown("---")
st.subheader("1. Seasonal Overview")
st.caption(filter_details)

period_values = (
    season_summary[period_col]
    .astype(str)
    .tolist()
)

cell_text = []

for _, row in season_summary.iterrows():
    temp_text = (
        f"{row['Median_Temperature']:.1f} °C"
        if pd.notna(
            row["Median_Temperature"]
        )
        else "—"
    )

    cell_text.append(
        (
            f"<b>{row['Median_Before_Viscosity']:.0f}"
            f" → {row['Median_After_Viscosity']:.0f} s</b>"
            f"<br>{row['Median_Solvent_Ratio']:.1f}%"
            f"<br>{temp_text}"
        )
    )

z_values = np.array(
    [
        season_summary[
            "Median_Viscosity_Drop"
        ].tolist()
    ],
    dtype=float,
)

customdata = np.array(
    [
        [
            [
                row["Median_Before_Viscosity"],
                row["Median_After_Viscosity"],
                row["Median_Viscosity_Drop"],
                row["Median_Solvent_Ratio"],
                row["Median_Temperature"],
                row["Historical_Records"],
                row["Historical_Batches"],
                row["Total_Base_Paint_kg"],
                row["Total_Solvent_kg"],
                row["Weighted_Solvent_Ratio"],
            ]
            for _, row in season_summary.iterrows()
        ]
    ],
    dtype=object,
)

fig_overview = go.Figure(
    data=go.Heatmap(
        z=z_values,
        x=period_values,
        y=[selected_paint_code],
        text=[cell_text],
        texttemplate="%{text}",
        textfont=dict(
            size=13,
            color="#111827",
        ),
        customdata=customdata,
        colorscale=[
            [0.00, "#EFF6FF"],
            [0.35, "#BFDBFE"],
            [0.70, "#60A5FA"],
            [1.00, "#1D4ED8"],
        ],
        colorbar=dict(
            title="ΔV<br>(s)",
            thickness=13,
            len=0.70,
        ),
        xgap=3,
        ygap=3,
        hovertemplate=(
            "<b>%{x}</b><br>"
            "Before: %{customdata[0]:.1f} s<br>"
            "After: %{customdata[1]:.1f} s<br>"
            "Drop: %{customdata[2]:.1f} s<br>"
            "Solvent Ratio: %{customdata[3]:.2f}%<br>"
            "Temperature: %{customdata[4]:.1f} °C<br>"
            "Records: %{customdata[5]:,.0f}<br>"
            "Batches: %{customdata[6]:,.0f}<br>"
            "Total Base Paint: %{customdata[7]:,.1f} kg<br>"
            "Total Solvent: %{customdata[8]:,.1f} kg<br>"
            "Weighted Solvent Ratio: %{customdata[9]:.2f}%"
            "<extra></extra>"
        ),
    )
)

fig_overview.update_layout(
    height=315,
    margin=dict(
        l=130,
        r=80,
        t=55,
        b=65,
    ),
    paper_bgcolor="white",
    plot_bgcolor="white",
    font=dict(
        family=(
            "Arial, Microsoft JhengHei, "
            "sans-serif"
        ),
        color="#334155",
    ),
)

fig_overview.update_xaxes(
    title_text="",
    side="top",
    showline=True,
    linecolor="#475569",
    mirror=True,
    ticks="outside",
    ticklen=6,
    tickfont=dict(size=12),
)

fig_overview.update_yaxes(
    title="Paint Code",
    showline=True,
    linecolor="#475569",
    mirror=True,
)

fig_overview.update_layout(
    title_text=""
)

st.plotly_chart(
    fig_overview,
    use_container_width=True,
)

st.caption(
    "每格依序顯示：添加前 → 添加後黏度、稀釋劑添加比例、溫度；底色代表典型降黏幅度。"
)


# =========================================================
# 11. KPI SUMMARY
# =========================================================
st.markdown("---")
st.subheader("2. Seasonal Summary")

highest_before_row = season_summary.loc[
    season_summary[
        "Median_Before_Viscosity"
    ].idxmax()
]

lowest_before_row = season_summary.loc[
    season_summary[
        "Median_Before_Viscosity"
    ].idxmin()
]

highest_ratio_row = season_summary.loc[
    season_summary[
        "Median_Solvent_Ratio"
    ].idxmax()
]

largest_drop_row = season_summary.loc[
    season_summary[
        "Median_Viscosity_Drop"
    ].idxmax()
]

k1, k2, k3, k4 = st.columns(4)

k1.metric(
    "Highest Incoming Viscosity",
    str(
        highest_before_row[
            period_col
        ]
    ),
    (
        f"{highest_before_row['Median_Before_Viscosity']:.1f} s"
    ),
)

k2.metric(
    "Lowest Incoming Viscosity",
    str(
        lowest_before_row[
            period_col
        ]
    ),
    (
        f"{lowest_before_row['Median_Before_Viscosity']:.1f} s"
    ),
)

k3.metric(
    "Highest Solvent Ratio",
    str(
        highest_ratio_row[
            period_col
        ]
    ),
    (
        f"{highest_ratio_row['Median_Solvent_Ratio']:.1f}%"
    ),
)

k4.metric(
    "Largest Viscosity Drop",
    str(
        largest_drop_row[
            period_col
        ]
    ),
    (
        f"{largest_drop_row['Median_Viscosity_Drop']:.1f} s"
    ),
)


# =========================================================
# 12. CHART 1 — BEFORE VS AFTER
# =========================================================
st.markdown("---")
st.subheader("3. Before vs. After Viscosity by Season")

fig_before_after = go.Figure()

for _, row in season_summary.iterrows():
    period_name = str(
        row[period_col]
    )

    fig_before_after.add_trace(
        go.Scatter(
            x=[
                row[
                    "Median_After_Viscosity"
                ],
                row[
                    "Median_Before_Viscosity"
                ],
            ],
            y=[
                period_name,
                period_name,
            ],
            mode="lines",
            line=dict(
                color="#94A3B8",
                width=5,
            ),
            hoverinfo="skip",
            showlegend=False,
        )
    )


fig_before_after.add_trace(
    go.Scatter(
        x=season_summary[
            "Median_Before_Viscosity"
        ],
        y=season_summary[
            period_col
        ].astype(str),
        mode="markers+text",
        name="Before Viscosity",
        marker=dict(
            size=17,
            color="#D97706",
            line=dict(
                color="white",
                width=1.5,
            ),
        ),
        text=season_summary[
            "Median_Before_Viscosity"
        ].map(
            lambda value: f"{value:.1f}"
        ),
        textposition="top center",
        customdata=np.column_stack(
            [
                season_summary[
                    "Before_P25"
                ],
                season_summary[
                    "Before_P75"
                ],
                season_summary[
                    "Historical_Records"
                ],
                season_summary[
                    "Historical_Batches"
                ],
            ]
        ),
        hovertemplate=(
            "<b>%{y}</b><br>"
            "Before Median: %{x:.1f} s<br>"
            "P25–P75: %{customdata[0]:.1f}"
            "–%{customdata[1]:.1f} s<br>"
            "Records: %{customdata[2]:,.0f}<br>"
            "Batches: %{customdata[3]:,.0f}"
            "<extra></extra>"
        ),
    )
)


fig_before_after.add_trace(
    go.Scatter(
        x=season_summary[
            "Median_After_Viscosity"
        ],
        y=season_summary[
            period_col
        ].astype(str),
        mode="markers+text",
        name="After Viscosity",
        marker=dict(
            size=17,
            color="#2563EB",
            line=dict(
                color="white",
                width=1.5,
            ),
        ),
        text=season_summary[
            "Median_After_Viscosity"
        ].map(
            lambda value: f"{value:.1f}"
        ),
        textposition="bottom center",
        customdata=np.column_stack(
            [
                season_summary[
                    "After_P25"
                ],
                season_summary[
                    "After_P75"
                ],
                season_summary[
                    "Median_Viscosity_Drop"
                ],
            ]
        ),
        hovertemplate=(
            "<b>%{y}</b><br>"
            "After Median: %{x:.1f} s<br>"
            "P25–P75: %{customdata[0]:.1f}"
            "–%{customdata[1]:.1f} s<br>"
            "Median Drop: %{customdata[2]:.1f} s"
            "<extra></extra>"
        ),
    )
)


fig_before_after.update_layout(
    title=dict(
        text=(
            f"<b>{selected_paint_code} — "
            "Seasonal Before vs. After Viscosity</b>"
            f"<br><sup>{filter_details}</sup>"
        ),
        x=0.5,
        xanchor="center",
    ),
    height=max(
        520,
        len(season_summary) * 72,
    ),
    margin=dict(
        l=160,
        r=60,
        t=125,
        b=75,
    ),
    template="plotly_white",
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=1.03,
        xanchor="center",
        x=0.5,
    ),
    xaxis=dict(
        title="Viscosity (s)",
        showgrid=True,
        gridcolor="#E5E7EB",
        showline=True,
        linecolor="#475569",
        mirror=True,
    ),
    yaxis=dict(
        title="Season",
        categoryorder="array",
        categoryarray=period_values,
        showgrid=True,
        gridcolor="#E5E7EB",
        showline=True,
        linecolor="#475569",
        mirror=True,
    ),
)

st.plotly_chart(
    fig_before_after,
    use_container_width=True,
)


# =========================================================
# 13. CHART 2 — SEASONAL VISCOSITY + SOLVENT RATIO + TEMPERATURE
#     Preferred line-chart style:
#     Before viscosity + After viscosity + Solvent ratio + Temperature
# =========================================================
st.markdown("---")
st.subheader("4. Seasonal Viscosity, Solvent Ratio and Temperature")

fig_condition = go.Figure()

period_series = season_summary[
    period_col
].astype(str)

# ---------------------------------------------------------
# Before viscosity
# ---------------------------------------------------------
fig_condition.add_trace(
    go.Scatter(
        x=period_series,
        y=season_summary[
            "Median_Before_Viscosity"
        ],
        mode="lines+markers+text",
        name="Before Viscosity (s)",
        line=dict(
            color="#D97706",
            width=3,
        ),
        marker=dict(
            size=9,
            color="#D97706",
        ),
        text=season_summary[
            "Median_Before_Viscosity"
        ].map(
            lambda value: f"{value:.1f}"
        ),
        textposition="top center",
        textfont=dict(
            size=11,
            color="#92400E",
        ),
        yaxis="y1",
        cliponaxis=False,
        hovertemplate=(
            "<b>%{x}</b><br>"
            "Before Viscosity: %{y:.1f} s"
            "<extra></extra>"
        ),
    )
)

# ---------------------------------------------------------
# After viscosity
# ---------------------------------------------------------
fig_condition.add_trace(
    go.Scatter(
        x=period_series,
        y=season_summary[
            "Median_After_Viscosity"
        ],
        mode="lines+markers+text",
        name="After Viscosity (s)",
        line=dict(
            color="#2563EB",
            width=3,
        ),
        marker=dict(
            size=9,
            color="#2563EB",
        ),
        text=season_summary[
            "Median_After_Viscosity"
        ].map(
            lambda value: f"{value:.1f}"
        ),
        textposition="bottom center",
        textfont=dict(
            size=11,
            color="#1D4ED8",
        ),
        yaxis="y1",
        cliponaxis=False,
        hovertemplate=(
            "<b>%{x}</b><br>"
            "After Viscosity: %{y:.1f} s"
            "<extra></extra>"
        ),
    )
)

# ---------------------------------------------------------
# Solvent ratio
# ---------------------------------------------------------
fig_condition.add_trace(
    go.Scatter(
        x=period_series,
        y=season_summary[
            "Median_Solvent_Ratio"
        ],
        mode="lines+markers",
        name="Solvent Ratio (%)",
        line=dict(
            color="#059669",
            width=3,
            dash="dot",
        ),
        marker=dict(
            size=10,
            color="#059669",
            symbol="diamond",
        ),
        yaxis="y2",
        customdata=np.column_stack(
            [
                season_summary[
                    "Total_Base_Paint_kg"
                ],
                season_summary[
                    "Total_Solvent_kg"
                ],
                season_summary[
                    "Weighted_Solvent_Ratio"
                ],
                season_summary[
                    "Historical_Records"
                ],
            ]
        ),
        hovertemplate=(
            "<b>%{x}</b><br>"
            "Median Solvent Ratio: %{y:.2f}%<br>"
            "Total Base Paint: %{customdata[0]:,.1f} kg<br>"
            "Total Solvent: %{customdata[1]:,.1f} kg<br>"
            "Weighted Solvent Ratio: %{customdata[2]:.2f}%<br>"
            "Records: %{customdata[3]:,.0f}"
            "<extra></extra>"
        ),
    )
)

# Solvent ratio labels
for i, row in season_summary.iterrows():
    fig_condition.add_annotation(
        x=str(row[period_col]),
        y=row[
            "Median_Solvent_Ratio"
        ],
        xref="x",
        yref="y2",
        text=(
            f"{row['Median_Solvent_Ratio']:.1f}%"
        ),
        showarrow=False,
        xshift=12 if i % 2 == 0 else -12,
        yshift=13,
        font=dict(
            size=10,
            color="#047857",
        ),
        bgcolor="rgba(255,255,255,0.90)",
        bordercolor="rgba(5,150,105,0.25)",
        borderwidth=1,
        borderpad=2,
    )

# ---------------------------------------------------------
# Temperature
# Use a third Y-axis on the right so the temperature line
# remains readable and does not compress the solvent-ratio line.
# ---------------------------------------------------------
if season_summary[
    "Median_Temperature"
].notna().any():
    fig_condition.add_trace(
        go.Scatter(
            x=period_series,
            y=season_summary[
                "Median_Temperature"
            ],
            mode="lines+markers+text",
            name="Temperature (°C)",
            line=dict(
                color="#7E22CE",
                width=3,
                dash="dash",
            ),
            marker=dict(
                size=9,
                color="#7E22CE",
                symbol="circle",
            ),
            text=season_summary[
                "Median_Temperature"
            ].map(
                lambda value: (
                    f"{value:.1f}°"
                    if pd.notna(value)
                    else ""
                )
            ),
            textposition="top center",
            textfont=dict(
                size=10,
                color="#6B21A8",
            ),
            yaxis="y3",
            cliponaxis=False,
            hovertemplate=(
                "<b>%{x}</b><br>"
                "Temperature: %{y:.1f} °C"
                "<extra></extra>"
            ),
        )
    )


# ---------------------------------------------------------
# Axis ranges
# ---------------------------------------------------------
ratio_values = pd.to_numeric(
    season_summary[
        "Median_Solvent_Ratio"
    ],
    errors="coerce",
).dropna()

if not ratio_values.empty:
    ratio_min = float(
        ratio_values.min()
    )
    ratio_max = float(
        ratio_values.max()
    )
    ratio_pad = max(
        0.8,
        (ratio_max - ratio_min) * 0.20,
    )
    ratio_range = [
        max(0, ratio_min - ratio_pad),
        ratio_max + ratio_pad,
    ]
else:
    ratio_range = None


temp_values = pd.to_numeric(
    season_summary[
        "Median_Temperature"
    ],
    errors="coerce",
).dropna()

if not temp_values.empty:
    temp_min = float(
        temp_values.min()
    )
    temp_max = float(
        temp_values.max()
    )
    temp_pad = max(
        1.0,
        (temp_max - temp_min) * 0.18,
    )
    temp_range = [
        temp_min - temp_pad,
        temp_max + temp_pad,
    ]
else:
    temp_range = None


# ---------------------------------------------------------
# Layout
# ---------------------------------------------------------
fig_condition.update_layout(
    title=dict(
        text=(
            f"<b>{selected_paint_code} — "
            "Seasonal Viscosity, Solvent Ratio & Temperature</b>"
            f"<br><sup>Thickness: "
            f"{structure_display}</sup>"
        ),
        x=0.5,
        xanchor="center",
    ),
    height=590,
    template="plotly_white",
    margin=dict(
        l=80,
        r=230,
        t=145,
        b=80,
    ),
    xaxis=dict(
        title="Season",
        categoryorder="array",
        categoryarray=period_values,
        domain=[0.0, 0.84],
        showgrid=False,
        showline=True,
        linecolor="#4B5563",
        linewidth=1.4,
        mirror=True,
        ticks="outside",
    ),
    yaxis=dict(
        title="Viscosity (s)",
        side="left",
        showgrid=True,
        gridcolor="#D6DCE5",
        gridwidth=1,
        showline=True,
        linecolor="#4B5563",
        linewidth=1.4,
        zeroline=False,
    ),
    yaxis2=dict(
        title=dict(
            text="Solvent Ratio (%)",
            font=dict(color="#047857"),
            standoff=8,
        ),
        overlaying="y",
        side="right",
        anchor="free",
        position=0.88,
        range=ratio_range,
        showgrid=False,
        showline=True,
        linecolor="#059669",
        linewidth=1.4,
        tickfont=dict(
            color="#047857",
            size=11,
        ),
        ticks="outside",
        ticklen=4,
        tickcolor="#059669",
        zeroline=False,
    ),
    yaxis3=dict(
        title="Temperature (°C)",
        overlaying="y",
        side="right",
        anchor="free",
        position=1.0,
        range=temp_range,
        showgrid=False,
        showline=False,
        tickfont=dict(
            color="#7E22CE"
        ),
        title_font=dict(
            color="#7E22CE"
        ),
        ticksuffix="°",
        zeroline=False,
    ),
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=1.07,
        xanchor="center",
        x=0.5,
        bgcolor="rgba(255,255,255,0.94)",
        bordercolor="#D1D5DB",
        borderwidth=1,
    ),
    font=dict(
        family=(
            "Arial, Microsoft JhengHei, "
            "sans-serif"
        ),
        size=12,
        color="#374151",
    ),
)

# Shift the temperature axis title farther right so it does not overlap
# with the solvent-ratio axis.
fig_condition.update_layout(
    yaxis3=dict(
        title=dict(
            text="Temperature (°C)",
            font=dict(color="#7E22CE"),
            standoff=10,
        ),
        overlaying="y",
        side="right",
        anchor="free",
        position=0.985,
        range=temp_range,
        showgrid=False,
        showline=True,
        linecolor="#7E22CE",
        linewidth=1.2,
        tickfont=dict(
            color="#7E22CE",
            size=11,
        ),
        ticks="outside",
        ticklen=4,
        tickcolor="#7E22CE",
        ticksuffix="°",
        zeroline=False,
    )
)

st.plotly_chart(
    fig_condition,
    use_container_width=True,
)

st.caption(
    "橘線＝添加前黏度；藍線＝添加後黏度；"
    "綠色虛線＝稀釋劑添加比例；紫色虛線＝溫度。"
)


# =========================================================
# 14. SINGLE SUMMARY TABLE
# =========================================================
st.markdown("---")
st.subheader("5. Seasonal Summary Table")
st.caption(
    "為避免將添加比例與稀釋劑總用量直接比較，表中新增「季節塗料使用量」與"
    "「加權添加比例」。生產量較高的季節，即使添加比例較低，也可能有較高的"
    "稀釋劑總用量。"
)

season_display = season_summary[
    [
        period_col,
        "Historical_Records",
        "Historical_Batches",
        "Median_Before_Viscosity",
        "Median_After_Viscosity",
        "Median_Viscosity_Drop",
        "Median_Solvent_Ratio",
        "Total_Base_Paint_kg",
        "Weighted_Solvent_Ratio",
        "Total_Solvent_kg",
        "Median_Temperature",
    ]
].copy()

season_display.insert(
    1,
    "膜厚組合",
    structure_display,
)

season_display = season_display.rename(
    columns={
        period_col: "季節期間",
        "Historical_Records": "歷史紀錄數",
        "Historical_Batches": "歷史批數",
        "Median_Before_Viscosity": "添加前黏度中位數",
        "Median_After_Viscosity": "添加後黏度中位數",
        "Median_Viscosity_Drop": "降黏幅度中位數",
        "Median_Solvent_Ratio": "添加比例中位數",
        "Total_Base_Paint_kg": "季節塗料使用量",
        "Weighted_Solvent_Ratio": "加權添加比例",
        "Total_Solvent_kg": "稀釋劑總用量",
        "Median_Temperature": "溫度中位數",
    }
)

round_cols = [
    "添加前黏度中位數",
    "添加後黏度中位數",
    "降黏幅度中位數",
    "添加比例中位數",
    "季節塗料使用量",
    "加權添加比例",
    "稀釋劑總用量",
    "溫度中位數",
]

season_display[
    round_cols
] = season_display[
    round_cols
].round(1)

st.dataframe(
    season_display,
    column_config={
        "季節期間": st.column_config.TextColumn(
            "季節期間",
            width="medium",
        ),
        "膜厚組合": st.column_config.TextColumn(
            "膜厚組合",
            width="medium",
        ),
        "歷史紀錄數": st.column_config.NumberColumn(
            "歷史紀錄數",
            format="%d",
        ),
        "歷史批數": st.column_config.NumberColumn(
            "歷史批數",
            format="%d",
        ),
        "添加前黏度中位數": st.column_config.NumberColumn(
            "添加前黏度中位數 (s)",
            format="%.1f",
        ),
        "添加後黏度中位數": st.column_config.NumberColumn(
            "添加後黏度中位數 (s)",
            format="%.1f",
        ),
        "降黏幅度中位數": st.column_config.NumberColumn(
            "降黏幅度中位數 (s)",
            format="%.1f",
        ),
        "添加比例中位數": st.column_config.NumberColumn(
            "添加比例中位數 (%)",
            format="%.1f",
            help="各筆紀錄添加比例的中位數。",
        ),
        "季節塗料使用量": st.column_config.NumberColumn(
            "季節塗料使用量 (kg)",
            format="%.1f",
            help="該季節所有有效紀錄之原始塗料重量合計，用於表示生產使用量。",
        ),
        "加權添加比例": st.column_config.NumberColumn(
            "加權添加比例 (%)",
            format="%.2f",
            help="季節稀釋劑總用量 ÷ 季節塗料使用量 × 100。",
        ),
        "溫度中位數": st.column_config.NumberColumn(
            "溫度中位數 (°C)",
            format="%.1f",
        ),
        "稀釋劑總用量": st.column_config.NumberColumn(
            "稀釋劑總用量 (kg)",
            format="%.1f",
        ),
    },
    use_container_width=True,
    hide_index=True,
)


# =========================================================
# 15. SHORT AUTOMATIC CONCLUSION
# =========================================================
st.markdown("---")
st.subheader("6. Automatic Conclusion")

before_range = (
    season_summary[
        "Median_Before_Viscosity"
    ].max()
    - season_summary[
        "Median_Before_Viscosity"
    ].min()
)

after_range = (
    season_summary[
        "Median_After_Viscosity"
    ].max()
    - season_summary[
        "Median_After_Viscosity"
    ].min()
)

ratio_range = (
    season_summary[
        "Median_Solvent_Ratio"
    ].max()
    - season_summary[
        "Median_Solvent_Ratio"
    ].min()
)

available_seasons = int(
    season_summary[
        "Season"
    ].nunique()
)

highest_production_row = season_summary.loc[
    season_summary["Total_Base_Paint_kg"].idxmax()
]

highest_production_period = str(
    highest_production_row[period_col]
)

highest_production_kg = float(
    highest_production_row["Total_Base_Paint_kg"]
)

if available_seasons < 2:
    st.info(
        "⚪ 可比較季節少於 2 個，目前資料不足以判定季節差異。"
    )
else:
    if before_range <= 5:
        before_text = (
            "添加前黏度季節差異小"
        )
    elif before_range <= 10:
        before_text = (
            "添加前黏度有輕微季節差異"
        )
    else:
        before_text = (
            "添加前黏度季節差異明顯"
        )

    if after_range <= 5:
        after_text = (
            "現場調整後黏度大致一致"
        )
    else:
        after_text = (
            "調整後黏度仍有季節差異"
        )

    st.markdown(
        f"- **{before_text}**：最大差異約 **{before_range:.1f} s**。  \n"
        f"- **{after_text}**：最大差異約 **{after_range:.1f} s**。  \n"
        f"- 稀釋劑添加比例季節差異約 **{ratio_range:.1f} 個百分點**。  \n"
        f"- **生產使用量最高期間：{highest_production_period}**，"
        f"塗料使用量約 **{highest_production_kg:,.1f} kg**。"
    )



# =========================================================
# 16. OPTIMAL & SAFE INCOMING VISCOSITY RECOMMENDATION
# =========================================================
st.markdown("---")
st.subheader("7. Optimal & Safe Incoming Viscosity Recommendation")

st.caption(
    "依目前篩選之 Vendor × Paint Code × Position × Coating Structure × "
    "Resin × Solvent 條件，使用歷史添加後黏度分布評估免加稀釋劑直接上線之"
    "試驗進料黏度範圍。"
)

recommendation_df = analysis_df[
    analysis_df["黏度(秒)_1"].notna()
    & (analysis_df["黏度(秒)_1"] > 0)
].copy()

if recommendation_df.empty:
    st.info(
        "⚪ 無有效添加後黏度資料，暫時無法建立建議進料黏度範圍。"
    )
else:
    # -----------------------------------------------------
    # 16.1 Core distribution
    # -----------------------------------------------------
    rec_records = int(len(recommendation_df))
    rec_batches = int(
        recommendation_df["Batch_ID"].nunique()
    )

    final_p25 = float(
        recommendation_df["黏度(秒)_1"].quantile(0.25)
    )
    final_median = float(
        recommendation_df["黏度(秒)_1"].median()
    )
    final_p75 = float(
        recommendation_df["黏度(秒)_1"].quantile(0.75)
    )

    final_iqr = final_p75 - final_p25

    current_before_p25 = float(
        recommendation_df["黏度(秒)"].quantile(0.25)
    )
    current_before_median = float(
        recommendation_df["黏度(秒)"].median()
    )
    current_before_p75 = float(
        recommendation_df["黏度(秒)"].quantile(0.75)
    )

    # -----------------------------------------------------
    # 16.2 Seasonal stability
    # -----------------------------------------------------
    seasonal_final = (
        recommendation_df
        .groupby(
            ["Season_Order", "Season"],
            dropna=False,
        )
        .agg(
            Records=("Paint_Code", "size"),
            Batches=("Batch_ID", "nunique"),
            Median_Final_Viscosity=("黏度(秒)_1", "median"),
            Median_Before_Viscosity=("黏度(秒)", "median"),
            Median_Solvent_Ratio=("Solvent_Ratio_Percent", "median"),
            Median_Temperature=("溫度", "median"),
        )
        .reset_index()
        .sort_values("Season_Order")
    )

    available_seasons_for_rec = int(
        seasonal_final["Season"].nunique()
    )

    if available_seasons_for_rec >= 2:
        seasonal_gap = float(
            seasonal_final["Median_Final_Viscosity"].max()
            - seasonal_final["Median_Final_Viscosity"].min()
        )
    else:
        seasonal_gap = np.nan

    # -----------------------------------------------------
    # 16.3 Evidence / safety screening
    # -----------------------------------------------------
    if rec_records < 10 or rec_batches < 3:
        evidence_status = "Insufficient Data"
        evidence_icon = "⚪"
        evidence_score = 0

    elif rec_records < 30 or rec_batches < 5:
        evidence_status = "Pilot Only"
        evidence_icon = "🟡"
        evidence_score = 1

    else:
        evidence_status = "Adequate Historical Evidence"
        evidence_icon = "🟢"
        evidence_score = 2


    if pd.isna(seasonal_gap):
        season_status = "Insufficient Seasonal Coverage"
        season_score = 0

    elif seasonal_gap <= 3:
        season_status = "Seasonally Stable"
        season_score = 2

    elif seasonal_gap <= 5:
        season_status = "Seasonal Monitoring Required"
        season_score = 1

    else:
        season_status = "Seasonal Difference Significant"
        season_score = 0


    if final_iqr <= 3:
        dispersion_status = "Tight Final Viscosity Distribution"
        dispersion_score = 2

    elif final_iqr <= 5:
        dispersion_status = "Moderate Final Viscosity Distribution"
        dispersion_score = 1

    else:
        dispersion_status = "Wide Final Viscosity Distribution"
        dispersion_score = 0


    total_safety_score = (
        evidence_score
        + season_score
        + dispersion_score
    )

    # -----------------------------------------------------
    # 16.4 Final recommendation classification
    # -----------------------------------------------------
    if (
        total_safety_score >= 5
        and rec_records >= 30
        and rec_batches >= 5
        and available_seasons_for_rec >= 2
    ):
        recommendation_status = (
            "Ready for No-Solvent Pilot"
        )
        recommendation_icon = "🟢"
        recommendation_message = (
            "歷史添加後黏度集中且季節差異小，"
            "可優先以此範圍進行供應商小批量免加稀釋劑直接上線試驗。"
        )

    elif total_safety_score >= 3:
        recommendation_status = (
            "Pilot with Monitoring"
        )
        recommendation_icon = "🟡"
        recommendation_message = (
            "可進行小批量試驗，但需同步監控膜厚、光澤、色差及成品品質；"
            "若季節差異較大，應持續累積資料後再決定是否建立正式規格。"
        )

    else:
        recommendation_status = (
            "Not Ready for Direct Specification"
        )
        recommendation_icon = "🟠"
        recommendation_message = (
            "目前資料量、季節穩定性或添加後黏度分布尚不足以建立單一安全進料範圍。"
        )


    # -----------------------------------------------------
    # 16.5 KPI cards
    # -----------------------------------------------------
    rec_col1, rec_col2, rec_col3, rec_col4 = st.columns(4)

    rec_col1.metric(
        "Recommended Lower",
        f"{final_p25:.1f} s",
    )

    rec_col2.metric(
        "Recommended Target",
        f"{final_median:.1f} s",
    )

    rec_col3.metric(
        "Recommended Upper",
        f"{final_p75:.1f} s",
    )

    rec_col4.metric(
        "Seasonal Gap",
        (
            f"{seasonal_gap:.1f} s"
            if pd.notna(seasonal_gap)
            else "N/A"
        ),
    )


    # -----------------------------------------------------
    # 16.6 Decision table
    # -----------------------------------------------------
    recommendation_table = pd.DataFrame(
        [
            {
                "Vendor": selected_vendor,
                "Paint Code": selected_paint_code,
                "Position": selected_position,
                "Coating Structure": structure_display,
                "Resin": selected_resin,
                "Solvent": selected_solvent,
                "Records": rec_records,
                "Batches": rec_batches,
                "Available Seasons": available_seasons_for_rec,
                "Current Incoming P25 (s)": current_before_p25,
                "Current Incoming Median (s)": current_before_median,
                "Current Incoming P75 (s)": current_before_p75,
                "Recommended Lower (s)": final_p25,
                "Recommended Target (s)": final_median,
                "Recommended Upper (s)": final_p75,
                "Final Viscosity IQR (s)": final_iqr,
                "Seasonal Final Viscosity Gap (s)": seasonal_gap,
                "Evidence": evidence_status,
                "Seasonal Stability": season_status,
                "Final Viscosity Stability": dispersion_status,
                "Recommendation": recommendation_status,
            }
        ]
    )

    st.dataframe(
        recommendation_table,
        column_config={
            "Records": st.column_config.NumberColumn(
                "Records",
                format="%d",
            ),
            "Batches": st.column_config.NumberColumn(
                "Batches",
                format="%d",
            ),
            "Available Seasons": st.column_config.NumberColumn(
                "Available Seasons",
                format="%d",
            ),
            "Current Incoming P25 (s)": st.column_config.NumberColumn(
                "Current Incoming P25 (s)",
                format="%.1f",
            ),
            "Current Incoming Median (s)": st.column_config.NumberColumn(
                "Current Incoming Median (s)",
                format="%.1f",
            ),
            "Current Incoming P75 (s)": st.column_config.NumberColumn(
                "Current Incoming P75 (s)",
                format="%.1f",
            ),
            "Recommended Lower (s)": st.column_config.NumberColumn(
                "Recommended Lower (s)",
                format="%.1f",
            ),
            "Recommended Target (s)": st.column_config.NumberColumn(
                "Recommended Target (s)",
                format="%.1f",
            ),
            "Recommended Upper (s)": st.column_config.NumberColumn(
                "Recommended Upper (s)",
                format="%.1f",
            ),
            "Final Viscosity IQR (s)": st.column_config.NumberColumn(
                "Final Viscosity IQR (s)",
                format="%.1f",
            ),
            "Seasonal Final Viscosity Gap (s)": st.column_config.NumberColumn(
                "Seasonal Final Viscosity Gap (s)",
                format="%.1f",
            ),
        },
        use_container_width=True,
        hide_index=True,
    )


    # -----------------------------------------------------
    # 16.7 Current vs recommended viscosity interval
    # -----------------------------------------------------
    fig_recommend = go.Figure()

    fig_recommend.add_trace(
        go.Scatter(
            x=[
                current_before_p25,
                current_before_p75,
            ],
            y=[
                "Current Incoming",
                "Current Incoming",
            ],
            mode="lines",
            line=dict(
                width=18,
                color="#D97706",
            ),
            name="Current Incoming P25–P75",
            hovertemplate=(
                f"Current Incoming: "
                f"{current_before_p25:.1f}–"
                f"{current_before_p75:.1f} s"
                "<extra></extra>"
            ),
        )
    )

    fig_recommend.add_trace(
        go.Scatter(
            x=[current_before_median],
            y=["Current Incoming"],
            mode="markers+text",
            marker=dict(
                size=16,
                symbol="diamond",
                color="#92400E",
            ),
            text=[
                f"{current_before_median:.1f} s"
            ],
            textposition="top center",
            name="Current Median",
            hovertemplate=(
                f"Current Median: "
                f"{current_before_median:.1f} s"
                "<extra></extra>"
            ),
        )
    )

    fig_recommend.add_trace(
        go.Scatter(
            x=[
                final_p25,
                final_p75,
            ],
            y=[
                "Recommended Incoming",
                "Recommended Incoming",
            ],
            mode="lines",
            line=dict(
                width=18,
                color="#2563EB",
            ),
            name="Recommended P25–P75",
            hovertemplate=(
                f"Recommended Range: "
                f"{final_p25:.1f}–"
                f"{final_p75:.1f} s"
                "<extra></extra>"
            ),
        )
    )

    fig_recommend.add_trace(
        go.Scatter(
            x=[final_median],
            y=["Recommended Incoming"],
            mode="markers+text",
            marker=dict(
                size=17,
                symbol="diamond",
                color="#1D4ED8",
            ),
            text=[
                f"Target {final_median:.1f} s"
            ],
            textposition="bottom center",
            name="Recommended Target",
            hovertemplate=(
                f"Recommended Target: "
                f"{final_median:.1f} s"
                "<extra></extra>"
            ),
        )
    )

    fig_recommend.add_annotation(
        x=final_p25,
        y="Recommended Incoming",
        text=f"Lower {final_p25:.1f}",
        showarrow=False,
        yshift=28,
        font=dict(
            size=11,
            color="#1D4ED8",
        ),
        bgcolor="rgba(255,255,255,0.92)",
    )

    fig_recommend.add_annotation(
        x=final_p75,
        y="Recommended Incoming",
        text=f"Upper {final_p75:.1f}",
        showarrow=False,
        yshift=28,
        font=dict(
            size=11,
            color="#1D4ED8",
        ),
        bgcolor="rgba(255,255,255,0.92)",
    )

    rec_x_min = min(
        current_before_p25,
        final_p25,
    )

    rec_x_max = max(
        current_before_p75,
        final_p75,
    )

    rec_x_pad = max(
        (rec_x_max - rec_x_min) * 0.12,
        5,
    )

    fig_recommend.update_layout(
        title=dict(
            text=(
                f"<b>{selected_paint_code} — "
                "Current vs. Recommended Incoming Viscosity</b>"
                f"<br><sup>Thickness: "
                f"{structure_display}</sup>"
            ),
            x=0.5,
            xanchor="center",
        ),
        height=470,
        template="plotly_white",
        margin=dict(
            l=165,
            r=60,
            t=120,
            b=75,
        ),
        xaxis=dict(
            title="Viscosity (s)",
            range=[
                rec_x_min - rec_x_pad,
                rec_x_max + rec_x_pad,
            ],
            showgrid=True,
            gridcolor="#E5E7EB",
            showline=True,
            linecolor="#475569",
            mirror=True,
        ),
        yaxis=dict(
            title="",
            categoryorder="array",
            categoryarray=[
                "Recommended Incoming",
                "Current Incoming",
            ],
            showgrid=True,
            gridcolor="#E5E7EB",
            showline=True,
            linecolor="#475569",
            mirror=True,
        ),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.03,
            xanchor="center",
            x=0.5,
        ),
    )

    st.plotly_chart(
        fig_recommend,
        use_container_width=True,
    )


    # -----------------------------------------------------
    # 16.8 Seasonal final viscosity verification
    # -----------------------------------------------------
    st.markdown(
        "#### Seasonal Final Viscosity Verification"
    )

    seasonal_final_display = (
        seasonal_final[
            [
                "Season",
                "Records",
                "Batches",
                "Median_Before_Viscosity",
                "Median_Final_Viscosity",
                "Median_Solvent_Ratio",
                "Median_Temperature",
            ]
        ]
        .copy()
        .rename(
            columns={
                "Season": "Season",
                "Records": "Records",
                "Batches": "Batches",
                "Median_Before_Viscosity": "Before Median (s)",
                "Median_Final_Viscosity": "Final Median (s)",
                "Median_Solvent_Ratio": "Solvent Ratio Median (%)",
                "Median_Temperature": "Temperature Median (°C)",
            }
        )
    )

    st.dataframe(
        seasonal_final_display,
        column_config={
            "Records": st.column_config.NumberColumn(
                "Records",
                format="%d",
            ),
            "Batches": st.column_config.NumberColumn(
                "Batches",
                format="%d",
            ),
            "Before Median (s)": st.column_config.NumberColumn(
                "Before Median (s)",
                format="%.1f",
            ),
            "Final Median (s)": st.column_config.NumberColumn(
                "Final Median (s)",
                format="%.1f",
            ),
            "Solvent Ratio Median (%)": st.column_config.NumberColumn(
                "Solvent Ratio Median (%)",
                format="%.1f",
            ),
            "Temperature Median (°C)": st.column_config.NumberColumn(
                "Temperature Median (°C)",
                format="%.1f",
            ),
        },
        hide_index=True,
        use_container_width=True,
    )


    # -----------------------------------------------------
    # 16.9 Final decision
    # -----------------------------------------------------
    if recommendation_status == "Ready for No-Solvent Pilot":
        st.success(
            f"{recommendation_icon} **{recommendation_status}**  \n"
            f"{recommendation_message}"
        )

    elif recommendation_status == "Pilot with Monitoring":
        st.warning(
            f"{recommendation_icon} **{recommendation_status}**  \n"
            f"{recommendation_message}"
        )

    else:
        st.warning(
            f"{recommendation_icon} **{recommendation_status}**  \n"
            f"{recommendation_message}"
        )


    st.caption(
        "建議下限／目標／上限分別採歷史添加後黏度 P25／Median／P75。"
        "此範圍僅作為供應商小批量試配及免加稀釋劑直接上線之試驗參考，"
        "正式進料規格仍須確認膜厚、光澤、色差、表面品質及成品品質皆符合要求。"
    )

# =========================================================
# 17. MANAGEMENT REPORT EXPORT — WORD
# =========================================================
st.markdown("---")
st.subheader("📥 Management Report Export")

csv_data = dataframe_to_csv_bytes(
    season_display
)

st.download_button(
    label="下載季節比較表 CSV",
    data=csv_data,
    file_name=(
        f"Seasonal_Viscosity_"
        f"{selected_paint_code}.csv"
    ),
    mime="text/csv",
    on_click="ignore",
)

if not HAS_DOCX:
    st.warning(
        "⚠️ Word export requires python-docx. "
        "Please add `python-docx` to requirements.txt."
    )
else:
    if st.button(
        "產生管理報告 Word",
        type="primary",
    ):
        try:
            # Prepare narrative variables used by the Word report.
            if available_seasons >= 2:
                if before_range <= 5:
                    seasonal_before_comment = (
                        "不同季節之添加前黏度差異小，整體季節影響有限。"
                    )
                elif before_range <= 10:
                    seasonal_before_comment = (
                        "不同季節之添加前黏度存在輕微差異，"
                        "建議持續觀察溫度及儲存條件。"
                    )
                else:
                    seasonal_before_comment = (
                        "不同季節之添加前黏度差異明顯，"
                        "建議確認溫度、儲存及批次因素。"
                    )

                if after_range <= 5:
                    seasonal_after_comment = (
                        "現場調整後黏度差異小，"
                        "顯示實際生產黏度具有一定穩定性。"
                    )
                else:
                    seasonal_after_comment = (
                        "現場調整後黏度仍存在季節差異，"
                        "需評估是否應分季節管理。"
                    )
            else:
                seasonal_before_comment = (
                    "目前季節資料不足，暫無法判定季節穩定性。"
                )
                seasonal_after_comment = (
                    "建議持續累積至少兩個以上季節之有效資料。"
                )

            if "recommendation_status" in locals():
                recommendation_summary = (
                    f"{recommendation_icon} {recommendation_status}"
                )
                recommendation_detail = recommendation_message
            else:
                recommendation_summary = "⚪ Insufficient Data"
                recommendation_detail = (
                    "目前無足夠資料建立建議進料黏度範圍。"
                )

            if "final_p25" in locals():
                recommended_range_text = (
                    f"{final_p25:.1f}–{final_p75:.1f} s，"
                    f"目標值 {final_median:.1f} s"
                )
            else:
                recommended_range_text = "N/A"

            word_buffer = create_management_word_report()

            st.success(
                "✅ Word 管理報告已產生。"
            )

            st.download_button(
                label="下載管理報告 Word",
                data=word_buffer.getvalue(),
                file_name=(
                    f"Seasonal_Viscosity_Management_Report_"
                    f"{selected_paint_code}_"
                    f"{pd.Timestamp.now().strftime('%Y%m%d_%H%M')}"
                    f".docx"
                ),
                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.wordprocessingml.document"
                ),
                on_click="ignore",
            )

            st.caption(
                "Word 圖表使用 Matplotlib 產生，不需要 kaleido。"
            )

        except Exception as error:
            st.error(
                f"❌ 產生 Word 管理報告時發生錯誤：{error}"
            )
