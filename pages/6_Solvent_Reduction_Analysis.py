import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import matplotlib.pyplot as plt
from matplotlib.transforms import blended_transform_factory
import io

# ==========================================
# WORD EXPORT LIBRARY (Check if installed)
# ==========================================
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ==========================================
# 1. PAGE CONFIGURATION & DATA LOAD
# ==========================================
st.set_page_config(page_title="Solvent Reduction Analysis", page_icon="🎨", layout="wide")
st.title("🎨 Solvent Reduction Opportunity Analysis")

if not st.session_state.get("raw_data_loaded", False) or st.session_state.get("group_a_data") is None:
    st.warning("⚠️ Please return to the Main App page and upload the raw data first.")
    st.stop()

df = st.session_state["group_a_data"].copy()

# ==========================================
# 2. DATA CLEANING & PREPARATION
# ==========================================
text_cols = ["Vendor", "Resin", "Solvent_Type", "塗料批號", "線別", "塗裝位置"]
for col in text_cols:
    df[col] = df.get(col, "Unknown").fillna("Unknown").astype(str).str.strip()

df["Paint_Code"] = df.get("塗料編號", pd.Series("Unknown", index=df.index)).fillna("Unknown").astype(str).str.strip().str.upper()
df["Solvent_Type"] = df["Solvent_Type"].str.upper()

invalid_vals = {"", "NAN", "NONE", "NULL", "N/A", "NA", "-", "--"}
df.replace(list(invalid_vals), "Unknown", inplace=True)

# Standardize Identifiers
df["Batch_ID"] = df["塗料批號"]
df["Bucket_Number"] = df.get("塗料桶號", pd.Series("Unknown", index=df.index)).fillna("Unknown").astype(str).str.strip()

pos_map = {"TP": "Primer", "正底漆": "Primer", "BP": "Primer", "背底漆": "Primer", "TF": "Top Finish", "正面漆": "Top Finish", "BF": "Back Finish", "背面漆": "Back Finish"}
df["Position_UI"] = df["塗裝位置"].map(pos_map).fillna(df["塗裝位置"])

num_cols = ["塗料重量", "添加重量", "黏度(秒)", "黏度(秒)_1", "溫度"]
for col in num_cols:
    df[col] = pd.to_numeric(df.get(col, np.nan), errors="coerce")

# ---------------------------------------------------------
# WEIGHT LOGIC
# Source column 塗料重量 is the total weight after solvent addition.
# Preserve the source value and calculate the original paint weight separately.
# ---------------------------------------------------------
df["Mixture_Weight_kg"] = df["塗料重量"]
df["Base_Paint_Weight_kg"] = df["Mixture_Weight_kg"] - df["添加重量"]

df["Delta_V"] = df["黏度(秒)"] - df["黏度(秒)_1"]
df["Solvent_Ratio_Percent"] = np.where(
    df["Base_Paint_Weight_kg"] > 0,
    df["添加重量"] / df["Base_Paint_Weight_kg"] * 100,
    np.nan,
)
df["Viscosity_Sensitivity"] = np.where(
    df["Solvent_Ratio_Percent"] > 0,
    df["Delta_V"] / df["Solvent_Ratio_Percent"],
    np.nan,
)

# Strict validation after reconstructing the original paint weight
df = df[
    (df["Base_Paint_Weight_kg"] > 0)
    & (df["添加重量"] > 0)
    & (df["黏度(秒)"] > 0)
    & (df["黏度(秒)_1"] > 0)
    & (df["Delta_V"] > 0)
].copy()

if df.empty:
    st.warning("⚠️ No valid dilution records remain after data cleaning.")
    st.stop()

# ==========================================
# 2.1 TIME PARSING & RECORD SELECTION LOGIC
# ==========================================
date_col = next((c for c in ["攪拌日期", "調整日期", "生產日期", "Date"] if c in df.columns), None)
time_col = next((c for c in ["攪拌時間(迄)", "攪拌時間", "Time"] if c in df.columns), None)

if date_col:
    date_part = pd.to_datetime(df[date_col], errors="coerce")
    df["_Analysis_Date"] = date_part
else:
    date_part = pd.Series(pd.NaT, index=df.index)
    df["_Analysis_Date"] = pd.NaT

if time_col:
    time_text = (
        df[time_col]
        .fillna("")
        .astype(str)
        .str.replace(r"\.0$", "", regex=True)
        .str.replace(":", "", regex=False)
        .str.zfill(4)
    )
    valid_time = time_text.str.fullmatch(r"\d{4}")
    hours = pd.to_numeric(time_text.str[:2], errors="coerce")
    minutes = pd.to_numeric(time_text.str[2:4], errors="coerce")
    valid_time &= hours.between(0, 23) & minutes.between(0, 59)
    time_delta = pd.to_timedelta(
        np.where(valid_time, hours * 60 + minutes, np.nan), unit="m"
    )
    df["_Analysis_DateTime"] = date_part.dt.normalize() + time_delta
else:
    df["_Analysis_DateTime"] = date_part

# Preserve source row order as a final tie-breaker when timestamps are missing/equal.
df["_Source_Order"] = np.arange(len(df))
df = df.sort_values(
    ["Batch_ID", "Bucket_Number", "_Analysis_DateTime", "_Source_Order"],
    ascending=True,
    na_position="first",
)

# PS30213X8: every row is an independent adjustment and must be retained.
# Other paint codes: repeated rows for the same batch + bucket are cumulative
# adjustments, therefore only the final chronological record is representative.
special_paint_codes = {"PS30213X8"}
is_special_paint = df["Paint_Code"].isin(special_paint_codes)

df_standard = (
    df.loc[~is_special_paint]
    .drop_duplicates(subset=["Batch_ID", "Bucket_Number"], keep="last")
    .copy()
)
df_special = df.loc[is_special_paint].copy()

df = pd.concat([df_standard, df_special], ignore_index=True)
df = df.sort_values(
    ["_Analysis_DateTime", "Batch_ID", "Bucket_Number", "_Source_Order"],
    ascending=True,
    na_position="last",
).reset_index(drop=True)

# ==========================================
# 3. CORE LOGIC HELPER
# ==========================================
def build_summary(source_df, group_cols):
    if source_df.empty: return pd.DataFrame()
    
    agg_dict = {
        "Adjustment_Records": ("Paint_Code", "size"),
        "Historical_Batches": ("Batch_ID", "nunique"),
        "Total_Paint_kg": ("Base_Paint_Weight_kg", "sum"),
        "Total_Solvent_kg": ("添加重量", "sum"),
        "Median_Paint_kg": ("Base_Paint_Weight_kg", "median"),
        "Median_Solvent_kg": ("添加重量", "median"),
        "Median_Ratio_Percent": ("Solvent_Ratio_Percent", "median"),
        "Median_Before_Viscosity": ("黏度(秒)", "median"),
        "Median_After_Viscosity": ("黏度(秒)_1", "median"),
        "Median_Viscosity_Drop": ("Delta_V", "median"),
        "Median_Dilution_Efficiency": ("Viscosity_Sensitivity", "median"),
    }
    if "線別" not in group_cols:
        agg_dict["Production_Lines"] = ("線別", lambda x: x[x != "Unknown"].nunique())

    summary = source_df.groupby(group_cols, dropna=False).agg(**agg_dict).reset_index()
    summary["Weighted_Ratio_Percent"] = np.where(summary["Total_Paint_kg"] > 0, summary["Total_Solvent_kg"] / summary["Total_Paint_kg"] * 100, np.nan)
    return summary


def apply_professional_layout(fig, title_text=None, subtitle_text=None, height=None):
    """Apply a clean business-style Plotly layout with black text."""
    annotations = list(fig.layout.annotations) if fig.layout.annotations else []

    if title_text is not None:
        fig.update_layout(
            title=dict(
                text=f"<b>{title_text}</b>",
                x=0.0,
                xanchor="left",
                y=0.985,
                yanchor="top",
                font=dict(size=22, color="#000000"),
                pad=dict(t=0, b=0),
            )
        )

    if subtitle_text:
        annotations.append(
            dict(
                x=0.0,
                y=1.205,
                xref="paper",
                yref="paper",
                text=subtitle_text,
                showarrow=False,
                xanchor="left",
                yanchor="bottom",
                align="left",
                font=dict(size=13, color="#000000"),
            )
        )

    fig.update_layout(
        template="plotly_white",
        paper_bgcolor="white",
        plot_bgcolor="white",
        font=dict(color="#000000", size=13),
        annotations=annotations,
        hoverlabel=dict(
            bgcolor="white",
            font=dict(color="#000000", size=12),
            bordercolor="#D1D5DB",
        ),
        legend=dict(
            bgcolor="rgba(255,255,255,0.95)",
            bordercolor="rgba(0,0,0,0)",
            font=dict(color="#000000", size=12),
            title=dict(font=dict(color="#000000", size=12)),
        ),
        margin=dict(l=90, r=50, t=235, b=90),
    )
    if height is not None:
        fig.update_layout(height=height)

    fig.update_xaxes(
        showline=True,
        linewidth=1.2,
        linecolor="#111827",
        mirror=True,
        ticks="outside",
        tickfont=dict(color="#000000", size=12),
        title=dict(font=dict(color="#000000", size=14)),
        gridcolor="#E5E7EB",
        zeroline=False,
    )
    fig.update_yaxes(
        showline=True,
        linewidth=1.2,
        linecolor="#111827",
        mirror=True,
        ticks="outside",
        tickfont=dict(color="#000000", size=12),
        title=dict(font=dict(color="#000000", size=14)),
        gridcolor="#E5E7EB",
        zerolinecolor="#D1D5DB",
    )
    return fig


def create_top10_usage_ratio_png(summary_df, filter_details):
    """Create the Top 10 usage/ratio chart for Word export with clean labels.

    Layout rules:
    1. Paint values are placed inside tall blue bars.
    2. Solvent values stay above the short yellow bars.
    3. Ratio labels are placed above the line points with alternating
       horizontal offsets and white backgrounds.
    4. Extra top space is reserved so labels never touch the chart border.
    """
    fig, ax1 = plt.subplots(figsize=(11.2, 6.3), dpi=180)

    if summary_df is None or summary_df.empty:
        ax1.text(
            0.5, 0.5, "No data available",
            ha="center", va="center", fontsize=12, color="black"
        )
        ax1.set_axis_off()
    else:
        chart_df = summary_df.copy().reset_index(drop=True)
        x = np.arange(len(chart_df))
        bar_width = 0.36

        paint_values = pd.to_numeric(
            chart_df["Total_Paint_kg"], errors="coerce"
        ).fillna(0)
        solvent_values = pd.to_numeric(
            chart_df["Total_Solvent_kg"], errors="coerce"
        ).fillna(0)
        ratio_values = pd.to_numeric(
            chart_df["Weighted_Ratio_Percent"], errors="coerce"
        )

        paint_bars = ax1.bar(
            x - bar_width / 2,
            paint_values,
            width=bar_width,
            label="Paint (kg)",
            color="#5B8FF9",
            edgecolor="white",
            linewidth=0.8,
            zorder=3,
        )
        solvent_bars = ax1.bar(
            x + bar_width / 2,
            solvent_values,
            width=bar_width,
            label="Solvent (kg)",
            color="#F6BD16",
            edgecolor="white",
            linewidth=0.8,
            zorder=3,
        )

        ax2 = ax1.twinx()
        ax2.plot(
            x,
            ratio_values,
            marker="o",
            markersize=6,
            linewidth=2.4,
            color="#10A9E2",
            label="Solvent Ratio (%)",
            zorder=5,
        )

        max_weight = max(float(max(paint_values.max(), solvent_values.max())), 1.0)
        max_ratio = max(float(ratio_values.max()), 1.0)

        weight_upper = max_weight * 1.24
        ratio_upper = max(5.0, max_ratio * 1.42)

        ax1.set_ylim(0, weight_upper)
        ax2.set_ylim(0, ratio_upper)

        # ---------------------------------------------------------
        # PAINT LABELS
        # Place all paint values inside the blue bars using black text
        # on a white background box. This remains readable in Word/PDF
        # and avoids blending into the blue bar.
        # ---------------------------------------------------------
        for bar in paint_bars:
            value = float(bar.get_height())
            x_center = bar.get_x() + bar.get_width() / 2

            if value >= max_weight * 0.12:
                label_y = value - max(value * 0.08, max_weight * 0.018)
                ax1.annotate(
                    f"{value:,.0f}",
                    xy=(x_center, label_y),
                    xytext=(0, 0),
                    textcoords="offset points",
                    ha="center",
                    va="top",
                    fontsize=8.0,
                    color="black",
                    fontweight="bold",
                    rotation=0,
                    bbox=dict(
                        boxstyle="round,pad=0.16",
                        facecolor="white",
                        edgecolor="none",
                        alpha=0.92,
                    ),
                    clip_on=True,
                    zorder=7,
                )
            else:
                ax1.annotate(
                    f"{value:,.0f}",
                    xy=(x_center, value),
                    xytext=(0, 5),
                    textcoords="offset points",
                    ha="center",
                    va="bottom",
                    fontsize=8.0,
                    color="black",
                    bbox=dict(
                        boxstyle="round,pad=0.14",
                        facecolor="white",
                        edgecolor="none",
                        alpha=0.92,
                    ),
                    clip_on=False,
                    zorder=7,
                )

        # ---------------------------------------------------------
        # SOLVENT LABELS
        # ---------------------------------------------------------
        for bar in solvent_bars:
            value = float(bar.get_height())
            x_center = bar.get_x() + bar.get_width() / 2

            ax1.annotate(
                f"{value:,.0f}",
                xy=(x_center, value),
                xytext=(0, 5),
                textcoords="offset points",
                ha="center",
                va="bottom",
                fontsize=8.0,
                color="black",
                clip_on=False,
                zorder=6,
            )

        # ---------------------------------------------------------
        # RATIO LABELS
        # Keep all percentage labels outside the bars.
        # Use alternating horizontal shifts to avoid crowding.
        # ---------------------------------------------------------
        for i, ratio in enumerate(ratio_values):
            if pd.isna(ratio):
                continue

            # Wider alternating offset for neighbouring labels.
            if i % 3 == 0:
                x_offset = -10
            elif i % 3 == 1:
                x_offset = 0
            else:
                x_offset = 10

            # Standard position above the point.
            y_offset = 15

            # For very low ratio points, move the label slightly higher
            # so it does not touch the paint-bar top or x-axis area.
            if float(ratio) <= ratio_upper * 0.36:
                y_offset = 18

            ax2.annotate(
                f"{ratio:.2f}%",
                xy=(x[i], ratio),
                xytext=(x_offset, y_offset),
                textcoords="offset points",
                ha="center",
                va="bottom",
                fontsize=8.8,
                fontweight="bold",
                color="black",
                bbox=dict(
                    boxstyle="round,pad=0.20",
                    facecolor="white",
                    edgecolor="none",
                    alpha=0.97,
                ),
                arrowprops=dict(
                    arrowstyle="-",
                    color="#10A9E2",
                    linewidth=0.65,
                    shrinkA=2,
                    shrinkB=4,
                ),
                annotation_clip=False,
                zorder=8,
            )

        ax1.set_xticks(x)
        ax1.set_xticklabels(
            chart_df["Paint_Code"].astype(str),
            fontsize=8.8,
            color="black",
        )
        ax1.set_xlabel("Paint Code", fontsize=11, color="black", labelpad=10)
        ax1.set_ylabel("Weight (kg)", fontsize=11, color="black")
        ax2.set_ylabel("Solvent Ratio (%)", fontsize=11, color="black")

        ax1.grid(axis="y", color="#E5E7EB", linewidth=0.8, zorder=1)
        ax1.tick_params(axis="both", colors="black", labelsize=9.5)
        ax2.tick_params(axis="y", colors="black", labelsize=9.5)
        ax1.set_facecolor("white")

        for spine in ax1.spines.values():
            spine.set_visible(True)
            spine.set_color("#111827")
            spine.set_linewidth(1.0)

        ax2.spines["right"].set_color("#111827")
        ax2.spines["right"].set_linewidth(1.0)
        ax2.spines["top"].set_visible(False)
        ax2.spines["left"].set_visible(False)
        ax2.spines["bottom"].set_visible(False)

        handles1, labels1 = ax1.get_legend_handles_labels()
        handles2, labels2 = ax2.get_legend_handles_labels()
        legend = ax1.legend(
            handles1 + handles2,
            labels1 + labels2,
            loc="upper center",
            bbox_to_anchor=(0.5, 1.13),
            ncol=3,
            frameon=False,
            fontsize=9.5,
        )
        for txt in legend.get_texts():
            txt.set_color("black")

    fig.patch.set_facecolor("white")
    fig.subplots_adjust(left=0.08, right=0.92, bottom=0.16, top=0.86)

    buf = io.BytesIO()
    fig.savefig(
        buf,
        format="png",
        bbox_inches="tight",
        facecolor="white",
        dpi=220,
        pad_inches=0.12,
    )
    plt.close(fig)
    buf.seek(0)
    return buf

def create_viscosity_history_png(chart_df):
    """Create the Tab 2 dumbbell viscosity chart for Word export."""
    fig, ax = plt.subplots(figsize=(11.2, 5.6), dpi=180)

    if chart_df is None or chart_df.empty:
        ax.text(0.5, 0.5, "No data available", ha="center", va="center", fontsize=12, color="black")
        ax.set_axis_off()
    else:
        export_df = chart_df.copy().reset_index(drop=True)
        x = export_df["Record_Index"].to_numpy(dtype=float)
        before = pd.to_numeric(export_df["黏度(秒)"], errors="coerce").to_numpy(dtype=float)
        after = pd.to_numeric(export_df["黏度(秒)_1"], errors="coerce").to_numpy(dtype=float)
        delta = pd.to_numeric(export_df["Delta_V"], errors="coerce").to_numpy(dtype=float)

        for xi, y_after, y_before in zip(x, after, before):
            if np.isfinite(xi) and np.isfinite(y_after) and np.isfinite(y_before):
                ax.plot([xi, xi], [y_after, y_before], color="#B8C2CC", linewidth=1.2, zorder=1)

        ax.scatter(x, before, s=34, color="#0B67C2", edgecolors="white", linewidths=0.6,
                   label="Before Viscosity (s)", zorder=3)
        ax.scatter(x, after, s=34, color="#74BDF2", edgecolors="white", linewidths=0.6,
                   label="After Viscosity (s)", zorder=3)

        if len(export_df) <= 30:
            for xi, y_before, y_after, d in zip(x, before, after, delta):
                if np.isfinite(xi) and np.isfinite(y_before) and np.isfinite(y_after) and np.isfinite(d):
                    ax.text(xi, (y_before + y_after) / 2, f"{d:.0f}", ha="center", va="center",
                            fontsize=7.5, color="#374151",
                            bbox=dict(facecolor="white", edgecolor="none", alpha=0.75, pad=0.5), zorder=4)

        ax.set_xlabel("Historical Record", fontsize=11, color="black")
        ax.set_ylabel("Viscosity (s)", fontsize=11, color="black")
        ax.grid(axis="y", color="#E5E7EB", linewidth=0.8, zorder=0)
        ax.tick_params(axis="both", colors="black", labelsize=9)
        if len(export_df) <= 30:
            ax.set_xticks(x)
        else:
            ax.locator_params(axis="x", nbins=12)
        ax.legend(loc="upper center", bbox_to_anchor=(0.5, 1.10), ncol=2, frameon=False, fontsize=9.5)
        ax.set_facecolor("white")
        for spine in ax.spines.values():
            spine.set_visible(True)
            spine.set_color("#111827")
            spine.set_linewidth(1.0)

    fig.patch.set_facecolor("white")
    fig.subplots_adjust(left=0.08, right=0.98, bottom=0.14, top=0.86)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white", dpi=220)
    plt.close(fig)
    buf.seek(0)
    return buf


def create_ratio_temperature_png(chart_df):
    """Create the Tab 2 solvent-ratio and temperature chart for Word export."""
    fig, ax1 = plt.subplots(figsize=(11.2, 5.2), dpi=180)

    if chart_df is None or chart_df.empty:
        ax1.text(0.5, 0.5, "No data available", ha="center", va="center", fontsize=12, color="black")
        ax1.set_axis_off()
    else:
        export_df = chart_df.copy().reset_index(drop=True)
        x = export_df["Record_Index"].to_numpy(dtype=float)
        ratio = pd.to_numeric(export_df["Solvent_Ratio_Percent"], errors="coerce").to_numpy(dtype=float)
        temp = pd.to_numeric(export_df["溫度"], errors="coerce").to_numpy(dtype=float)

        if len(export_df) > 1:
            diffs = np.diff(np.sort(np.unique(x)))
            bar_width = max(0.45, min(0.8, float(np.nanmedian(diffs)) * 0.72)) if len(diffs) else 0.65
        else:
            bar_width = 0.65

        bars = ax1.bar(x, ratio, width=bar_width, color="#F87171", alpha=0.78,
                       label="Solvent Ratio (%)", edgecolor="white", linewidth=0.5, zorder=2)
        ax1.set_xlabel("Historical Record", fontsize=11, color="black")
        ax1.set_ylabel("Solvent Ratio (%)", fontsize=11, color="#DC2626")
        ratio_max = np.nanmax(ratio) if np.isfinite(ratio).any() else 5.0
        ax1.set_ylim(0, max(5.0, ratio_max * 1.18))
        ax1.tick_params(axis="y", colors="#DC2626", labelsize=9)
        ax1.tick_params(axis="x", colors="black", labelsize=9)
        ax1.grid(axis="y", color="#FEE2E2", linewidth=0.8, zorder=0)

        ax2 = ax1.twinx()
        valid_temp = np.isfinite(temp)
        if valid_temp.any():
            ax2.plot(x, temp, color="#7E22CE", marker="o", markersize=4.5, linewidth=1.8,
                     label="Temperature (°C)", zorder=3)
            temp_min = float(np.nanmin(temp))
            temp_max = float(np.nanmax(temp))
            temp_pad = max(1.0, (temp_max - temp_min) * 0.15)
            ax2.set_ylim(temp_min - temp_pad, temp_max + temp_pad)
        ax2.set_ylabel("Temperature (°C)", fontsize=11, color="#7E22CE")
        ax2.tick_params(axis="y", colors="#7E22CE", labelsize=9)

        if len(export_df) <= 30:
            ax1.set_xticks(x)
        else:
            ax1.locator_params(axis="x", nbins=12)

        handles1, labels1 = ax1.get_legend_handles_labels()
        handles2, labels2 = ax2.get_legend_handles_labels()
        ax1.legend(handles1 + handles2, labels1 + labels2, loc="upper center",
                   bbox_to_anchor=(0.5, 1.10), ncol=2, frameon=False, fontsize=9.5)

        ax1.set_facecolor("white")
        for spine in ax1.spines.values():
            spine.set_visible(True)
            spine.set_color("#111827")
            spine.set_linewidth(1.0)
        ax2.spines["right"].set_color("#111827")
        ax2.spines["right"].set_linewidth(1.0)
        ax2.spines["top"].set_visible(False)
        ax2.spines["left"].set_visible(False)
        ax2.spines["bottom"].set_visible(False)

    fig.patch.set_facecolor("white")
    fig.subplots_adjust(left=0.08, right=0.92, bottom=0.14, top=0.86)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white", dpi=220)
    plt.close(fig)
    buf.seek(0)
    return buf

def create_supplier_priority_png(plot_df, target_solvent_limit):
    """Create a clean Word-exportable priority matrix.

    The chart title and subtitle are intentionally excluded from the PNG and
    are written as separate Word paragraphs during export.
    """
    fig, ax = plt.subplots(figsize=(11.2, 6.5), dpi=180)

    export_colors = {
        "High Supplier Priority": "#1D4ED8",
        "Validate with Supplier": "#F97316",
        "Monitor": "#60A5FA",
        "Low Priority": "#FDBA74",
    }

    if plot_df is None or plot_df.empty:
        ax.text(
            0.5, 0.5, "No data available",
            ha="center", va="center", fontsize=12, color="black",
            transform=ax.transAxes,
        )
        ax.set_axis_off()
    else:
        chart_df = plot_df.copy().reset_index(drop=True)
        max_batches = max(float(chart_df["Historical_Batches"].max()), 1.0)
        bubble_sizes = 45 + chart_df["Historical_Batches"] / max_batches * 620

        for action, subdf in chart_df.groupby("Supplier_Action", sort=False):
            ax.scatter(
                subdf["Matrix_X"],
                subdf["Total_Solvent_kg"],
                s=bubble_sizes.loc[subdf.index],
                color=export_colors.get(action, "#9CA3AF"),
                edgecolors="white",
                linewidths=1.0,
                alpha=0.88,
                label=action,
                zorder=3,
            )

        max_y = max(float(chart_df["Total_Solvent_kg"].max()), 1.0)
        y_lower = -max_y * 0.045
        y_upper = max_y * 1.16
        ax.set_xlim(-0.45, 3.45)
        ax.set_ylim(y_lower, y_upper)

        ax.axvline(2.5, color="#DC2626", linestyle=(0, (4, 4)), linewidth=1.2, zorder=2)
        ax.axhline(target_solvent_limit, color="#DC2626", linestyle=(0, (4, 4)), linewidth=1.2, zorder=2)

        # Label only the most important paint codes and allocate non-overlapping
        # vertical label slots within each stability-score group.
        priority_map = {
            "High Supplier Priority": 4,
            "Validate with Supplier": 3,
            "Monitor": 2,
            "Low Priority": 1,
        }
        important = chart_df.copy()
        important["Label_Priority"] = important["Supplier_Action"].map(priority_map).fillna(0)
        important = important.sort_values(
            ["Label_Priority", "Total_Solvent_kg", "Historical_Batches"],
            ascending=[False, False, False],
        )
        important["Label_Rank_In_X"] = important.groupby("High_Stability_Count").cumcount()
        important = important[important["Label_Rank_In_X"] < 3].head(10).copy()

        min_gap = max(max_y * 0.055, 850.0)
        label_positions = {}
        for score, group in important.groupby("High_Stability_Count", sort=True):
            group = group.sort_values("Total_Solvent_kg")
            previous_y = None
            for idx, row in group.iterrows():
                desired_y = float(row["Total_Solvent_kg"]) + max_y * 0.035
                if previous_y is not None and desired_y - previous_y < min_gap:
                    desired_y = previous_y + min_gap
                desired_y = min(desired_y, y_upper - max_y * 0.045)
                label_positions[idx] = desired_y
                previous_y = desired_y

        for idx, row in important.iterrows():
            point_x = float(row["Matrix_X"])
            point_y = float(row["Total_Solvent_kg"])
            rank = int(row["Label_Rank_In_X"])
            horizontal_shift = -0.12 if rank % 2 == 0 else 0.12
            if point_x >= 2.85:
                horizontal_shift = -0.16
            elif point_x <= 0.10:
                horizontal_shift = 0.16

            label_x = point_x + horizontal_shift
            label_y = label_positions.get(idx, point_y + max_y * 0.04)

            ax.annotate(
                str(row["Paint_Code"]),
                xy=(point_x, point_y),
                xycoords="data",
                xytext=(label_x, label_y),
                textcoords="data",
                fontsize=8.5,
                color="black",
                ha="center",
                va="bottom",
                bbox=dict(
                    boxstyle="round,pad=0.20",
                    facecolor="white",
                    edgecolor="#9CA3AF",
                    linewidth=0.75,
                    alpha=0.97,
                ),
                arrowprops=dict(
                    arrowstyle="-",
                    color="#9CA3AF",
                    linewidth=0.75,
                    shrinkA=2,
                    shrinkB=4,
                ),
                annotation_clip=False,
                zorder=5,
            )

        ax.text(
            2.54, y_upper * 0.965, "Strong Evidence Zone",
            fontsize=9.2, color="black", ha="left", va="top",
            bbox=dict(facecolor="white", edgecolor="#DC2626", linewidth=0.8, pad=2),
            zorder=5,
        )

        threshold_transform = blended_transform_factory(ax.transAxes, ax.transData)
        ax.text(
            1.015, target_solvent_limit,
            f"Threshold: {target_solvent_limit:,.0f} kg",
            transform=threshold_transform,
            fontsize=9.0, color="black", ha="left", va="center",
            bbox=dict(facecolor="white", edgecolor="#DC2626", linewidth=0.8, pad=2),
            clip_on=False, zorder=5,
        )

    ax.set_xlabel("High-Stability Checks Passed (0–3)", fontsize=11.5, color="black", labelpad=10)
    ax.set_ylabel("Total On-site Solvent Adjustment (kg)", fontsize=11.5, color="black", labelpad=10)
    ax.set_xticks([0, 1, 2, 3])
    ax.tick_params(axis="both", colors="black", labelsize=10)
    ax.grid(axis="y", color="#E5E7EB", linewidth=0.8, zorder=0)
    ax.set_facecolor("white")

    for spine in ax.spines.values():
        spine.set_visible(True)
        spine.set_color("#111827")
        spine.set_linewidth(1.1)

    handles, labels = ax.get_legend_handles_labels()
    if handles:
        preferred_order = [
            "High Supplier Priority",
            "Low Priority",
            "Monitor",
            "Validate with Supplier",
        ]
        legend_items = dict(zip(labels, handles))
        ordered_labels = [label for label in preferred_order if label in legend_items]
        leg = ax.legend(
            [legend_items[label] for label in ordered_labels],
            ordered_labels,
            ncol=1,
            loc="upper left",
            bbox_to_anchor=(1.01, 1.00),
            frameon=True,
            fontsize=9.3,
            handletextpad=0.7,
            labelspacing=0.9,
            borderaxespad=0.0,
        )
        leg.get_frame().set_facecolor("white")
        leg.get_frame().set_edgecolor("#D1D5DB")
        leg.get_frame().set_linewidth(0.8)
        for text_item in leg.get_texts():
            text_item.set_color("black")

    # No embedded title/subtitle here. They are added separately in Word.
    fig.patch.set_facecolor("white")
    fig.subplots_adjust(left=0.10, right=0.76, bottom=0.15, top=0.96)

    buf = io.BytesIO()
    fig.savefig(
        buf,
        format="png",
        facecolor="white",
        dpi=220,
        bbox_inches="tight",
        pad_inches=0.16,
    )
    plt.close(fig)
    buf.seek(0)
    return buf


# ==========================================
# 4. GLOBAL FILTERS
# ==========================================
st.markdown("---")
st.subheader("🔍 Global Filters")

filter_df = df.copy()
col1, col2, col3, col4, col5 = st.columns(5)

vendor_opts = ["All"] + sorted([str(x) for x in filter_df["Vendor"].unique() if x != "Unknown"])
selected_vendor = col1.selectbox("Vendor", vendor_opts)
if selected_vendor != "All": filter_df = filter_df[filter_df["Vendor"] == selected_vendor]

resin_opts = ["All"] + sorted([str(x) for x in filter_df["Resin"].unique() if x != "Unknown"])
selected_resin = col2.selectbox("Resin Type", resin_opts)
if selected_resin != "All": filter_df = filter_df[filter_df["Resin"] == selected_resin]

pos_opts = ["All"] + sorted([str(x) for x in filter_df["Position_UI"].unique() if x != "Unknown"])
selected_position = col3.selectbox("Coating Position", pos_opts)
if selected_position != "All": filter_df = filter_df[filter_df["Position_UI"] == selected_position]

solvent_opts = ["All"] + sorted([str(x) for x in filter_df["Solvent_Type"].unique() if x != "Unknown"])
selected_solvent = col4.selectbox("Solvent Type", solvent_opts)
if selected_solvent != "All": filter_df = filter_df[filter_df["Solvent_Type"] == selected_solvent]

line_opts = sorted([str(x) for x in filter_df["線別"].unique() if x != "Unknown"])
selected_lines = col5.multiselect("Production Line", line_opts, default=line_opts)

if selected_lines:
    filter_df = filter_df[filter_df["線別"].isin(selected_lines)]
else:
    st.warning("⚠️ Please select at least one production line.")
    st.stop()

if filter_df.empty:
    st.warning("⚠️ No records match the global analysis filters.")
    st.stop()

st.markdown("---")
if "_Analysis_Date" in filter_df.columns and not filter_df["_Analysis_Date"].isna().all():
    min_date = filter_df["_Analysis_Date"].min().strftime("%Y-%m-%d")
    max_date = filter_df["_Analysis_Date"].max().strftime("%Y-%m-%d")
    period_label = f"{min_date} ➔ {max_date}"
else:
    period_label = "All available data"

st.info(f"📅 **Analysis Period:** {period_label} | 📊 **Valid Records:** {len(filter_df):,}")
filter_details = f"Vendor: {selected_vendor} | Resin: {selected_resin} | Position: {selected_position} | Solvent: {selected_solvent}"
st.markdown("<br>", unsafe_allow_html=True)

# Initialize Dictionary to store charts for HTML export
exported_figs = {}

# ==========================================
# 5. HIERARCHICAL OVERVIEW (TREEMAP)
# ==========================================
st.subheader("🗂️ Hierarchical Overview")
st.markdown("Hierarchy: **Vendor ➔ Resin ➔ Position ➔ Solvent Type ➔ Paint Code**. Box size represents total solvent usage (kg).")

# Prepare Treemap data
tree_df = filter_df.groupby(["Vendor", "Resin", "Position_UI", "Solvent_Type", "Paint_Code"]).agg(
    添加重量=("添加重量", "sum"),
    Delta_V=("Delta_V", "median"), 
    Solvent_Ratio_Percent=("Solvent_Ratio_Percent", "median")
).reset_index()
tree_df = tree_df[tree_df["添加重量"] > 0]

fig_tree = px.treemap(
    tree_df,
    path=[px.Constant("Total"), "Vendor", "Resin", "Position_UI", "Solvent_Type", "Paint_Code"],
    values="添加重量", 
    color="Resin",  
    color_discrete_sequence=px.colors.qualitative.Pastel,
    custom_data=["Delta_V", "Solvent_Ratio_Percent"],
    title=f"Hierarchical Breakdown of Solvent Usage (kg)<br><sup>Filters: {filter_details}</sup>",
    height=700
)

# Update Labels and Tooltips
fig_tree.update_traces(
    texttemplate="<b>%{label}</b><br>%{value:,.0f} kg",
    hovertemplate="<b>%{label}</b><br>Solvent Usage: %{value:,.1f} kg<br>Visc Drop: ~%{customdata[0]:.1f} s<br>Solvent Added: ~%{customdata[1]:.1f}%",
    root_color="#f8f9fa"
)
fig_tree.update_layout(margin=dict(t=90, l=10, r=10, b=10)) 

st.plotly_chart(fig_tree, use_container_width=True)
exported_figs["1. Hierarchical Treemap"] = fig_tree
st.markdown("<br>", unsafe_allow_html=True)

# ==========================================
# 6. TABS & VISUALIZATION
# ==========================================
tab_ranking, tab_detail, tab_line, tab_pilot = st.tabs([
    "1️⃣ Top 10 Paint Code Ranking",
    "2️⃣ Paint Code History",
    "3️⃣ Line Comparison",
    "4️⃣ Supplier Improvement Priority",
])

# ----- TAB 1: RANKING -----
with tab_ranking:
    st.subheader("1. Paint Code Solvent Consumption (Top 10)")

    full_summary_df = build_summary(
        filter_df,
        ["Vendor", "Resin", "Position_UI", "Paint_Code", "Solvent_Type"],
    )
    summary_df = (
        full_summary_df
        .sort_values("Total_Solvent_kg", ascending=False)
        .head(10)
        .reset_index(drop=True)
    )
    summary_df.insert(0, "Rank", np.arange(1, len(summary_df) + 1))

    if summary_df.empty:
        st.info("No paint-code summary is available for the selected filters.")
    else:
        # One combined chart replaces the previous duplicate usage and ratio charts.
        fig_dual = go.Figure()

        fig_dual.add_trace(go.Bar(
            x=summary_df["Paint_Code"],
            y=summary_df["Total_Paint_kg"],
            name="Paint (kg)",
            marker_color="#5B8FF9",
            yaxis="y1",
            text=summary_df["Total_Paint_kg"].apply(lambda x: f"{x:,.0f}"),
            textposition="auto",
        ))

        fig_dual.add_trace(go.Bar(
            x=summary_df["Paint_Code"],
            y=summary_df["Total_Solvent_kg"],
            name="Solvent (kg)",
            marker_color="#F6BD16",
            yaxis="y1",
            text=summary_df["Total_Solvent_kg"].apply(lambda x: f"{x:,.0f}"),
            textposition="auto",
        ))

        fig_dual.add_trace(go.Scatter(
            x=summary_df["Paint_Code"],
            y=summary_df["Weighted_Ratio_Percent"],
            name="Solvent Ratio (%)",
            mode="lines+markers",
            line=dict(color="DeepSkyBlue", width=3),
            marker=dict(size=8),
            yaxis="y2",
        ))

        for _, row in summary_df.iterrows():
            fig_dual.add_annotation(
                x=row["Paint_Code"],
                y=row["Weighted_Ratio_Percent"],
                text=f"<b>{row['Weighted_Ratio_Percent']:.2f}%</b>",
                xref="x",
                yref="y2",
                showarrow=False,
                yshift=18,
                font=dict(color="black", size=12),
                bgcolor="rgba(255,255,255,0.85)",
                borderpad=2,
            )

        ratio_max = summary_df["Weighted_Ratio_Percent"].max()
        ratio_upper = max(5.0, float(ratio_max) * 1.25) if pd.notna(ratio_max) else 5.0

        fig_dual.update_layout(
            title=(
                "Paint & Solvent Usage vs. Solvent Ratio"
                f"<br><sup>Filters Applied: {filter_details}</sup>"
            ),
            xaxis=dict(title="Paint Code"),
            yaxis=dict(title="Weight (kg)", side="left", showgrid=False),
            yaxis2=dict(
                title="Solvent Ratio (%)",
                overlaying="y",
                side="right",
                showgrid=False,
                range=[0, ratio_upper],
            ),
            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1,
            ),
            height=600,
            uniformtext_minsize=8,
            uniformtext_mode="hide",
            barmode="group",
        )
        st.plotly_chart(fig_dual, use_container_width=True)
        exported_figs["4. Top 10 Usage and Ratio"] = fig_dual

        ranking_display = summary_df[[
            "Rank",
            "Paint_Code",
            "Total_Paint_kg",
            "Total_Solvent_kg",
            "Weighted_Ratio_Percent",
            "Adjustment_Records",
            "Historical_Batches",
        ]].rename(columns={
            "Paint_Code": "Paint Code",
            "Total_Paint_kg": "Total Paint (kg)",
            "Total_Solvent_kg": "Total Solvent (kg)",
            "Weighted_Ratio_Percent": "Solvent Ratio (%)",
            "Adjustment_Records": "Records",
            "Historical_Batches": "Batches",
        })

        st.dataframe(
            ranking_display.style.format({
                "Total Paint (kg)": "{:,.2f}",
                "Total Solvent (kg)": "{:,.2f}",
                "Solvent Ratio (%)": "{:.2f}",
            }),
            use_container_width=True,
            hide_index=True,
        )

# ----- TAB 2: PAINT CODE HISTORY -----
with tab_detail:
    st.subheader("2. Paint Code History")
    st.caption(
        "Each record is one independent viscosity-adjustment event. "
        "The upper chart shows the before-to-after viscosity change; "
        "the lower chart shows the corresponding solvent ratio and temperature."
    )

    # Use every paint code available under the global filters, not only the Top 10 list.
    detail_code_options = sorted(
        filter_df["Paint_Code"].dropna().astype(str).unique().tolist()
    )

    if not detail_code_options:
        st.info("No paint code is available for the selected global filters.")
    else:
        select_col, range_col = st.columns([2.2, 1.0])
        selected_code = select_col.selectbox(
            "Select Paint Code",
            detail_code_options,
            key="tab2_selected_paint_code",
        )
        record_window = range_col.selectbox(
            "Records Displayed",
            ["Latest 30", "Latest 50", "Latest 100", "All Records"],
            index=1,
            key="tab2_record_window",
            help="Limit the number of visible records to keep the charts readable. "
                 "All records remain available in the detail table.",
        )

        detail_df = filter_df[
            filter_df["Paint_Code"].astype(str) == str(selected_code)
        ].copy()

        if detail_df.empty:
            st.warning("No historical records were found for the selected paint code.")
        else:
            sort_cols = [
                col for col in ["_Analysis_DateTime", "_Source_Order"]
                if col in detail_df.columns
            ]
            if sort_cols:
                detail_df = detail_df.sort_values(sort_cols, na_position="last")

            detail_df = detail_df.reset_index(drop=True)
            detail_df["Record_Index"] = np.arange(1, len(detail_df) + 1)
            detail_title_filter = f"{filter_details} | Paint Code: {selected_code}"

            # KPI values always use the complete historical data for the selected paint code.
            typical_before = detail_df["黏度(秒)"].median()
            typical_after = detail_df["黏度(秒)_1"].median()
            typical_ratio = detail_df["Solvent_Ratio_Percent"].median()
            typical_temperature = detail_df["溫度"].median()

            m1, m2, m3, m4, m5 = st.columns(5)
            m1.metric("Historical Records", f"{len(detail_df):,}")
            m2.metric(
                "Typical Incoming Viscosity",
                f"{typical_before:.2f} s" if pd.notna(typical_before) else "N/A",
            )
            m3.metric(
                "Typical Adjusted Viscosity",
                f"{typical_after:.2f} s" if pd.notna(typical_after) else "N/A",
            )
            m4.metric(
                "Typical Solvent Ratio",
                f"{typical_ratio:.2f}%" if pd.notna(typical_ratio) else "N/A",
            )
            m5.metric(
                "Typical Temperature",
                f"{typical_temperature:.2f} °C" if pd.notna(typical_temperature) else "N/A",
            )

            # Limit only the chart view. The table below still contains all records.
            window_map = {
                "Latest 30": 30,
                "Latest 50": 50,
                "Latest 100": 100,
                "All Records": None,
            }
            visible_n = window_map[record_window]
            chart_df = detail_df.tail(visible_n).copy() if visible_n else detail_df.copy()
            chart_df = chart_df.reset_index(drop=True)

            # Keep the original historical record number even when only the latest records are shown.
            first_record = int(chart_df["Record_Index"].min())
            last_record = int(chart_df["Record_Index"].max())
            chart_range_text = (
                f"Records {first_record}–{last_record} of {len(detail_df)}"
                if len(chart_df) < len(detail_df)
                else f"All {len(detail_df)} records"
            )

            # =========================================================
            # CHART 1 — DUMBBELL: BEFORE → AFTER VISCOSITY
            # =========================================================
            segment_x, segment_y = [], []
            for _, row in chart_df.iterrows():
                segment_x.extend([row["Record_Index"], row["Record_Index"], None])
                segment_y.extend([row["黏度(秒)_1"], row["黏度(秒)"], None])

            fig_viscosity = go.Figure()
            fig_viscosity.add_trace(go.Scatter(
                x=segment_x,
                y=segment_y,
                mode="lines",
                name="Viscosity Drop",
                line=dict(color="#B8C2CC", width=1.5),
                hoverinfo="skip",
            ))
            fig_viscosity.add_trace(go.Scatter(
                x=chart_df["Record_Index"],
                y=chart_df["黏度(秒)"],
                mode="markers",
                name="Before Viscosity (s)",
                marker=dict(size=8, color="#0B67C2", line=dict(width=0.8, color="white")),
                customdata=np.column_stack([
                    chart_df["黏度(秒)_1"],
                    chart_df["Delta_V"],
                    chart_df["Solvent_Ratio_Percent"],
                    chart_df["溫度"],
                ]),
                hovertemplate=(
                    "<b>Record %{x}</b><br>"
                    "Before Viscosity: %{y:.1f} s<br>"
                    "After Viscosity: %{customdata[0]:.1f} s<br>"
                    "Viscosity Drop: %{customdata[1]:.1f} s<br>"
                    "Solvent Ratio: %{customdata[2]:.2f}%<br>"
                    "Temperature: %{customdata[3]:.1f} °C"
                    "<extra></extra>"
                ),
            ))
            fig_viscosity.add_trace(go.Scatter(
                x=chart_df["Record_Index"],
                y=chart_df["黏度(秒)_1"],
                mode="markers",
                name="After Viscosity (s)",
                marker=dict(size=8, color="#74BDF2", line=dict(width=0.8, color="white")),
                hovertemplate=(
                    "<b>Record %{x}</b><br>"
                    "After Viscosity: %{y:.1f} s"
                    "<extra></extra>"
                ),
            ))

            # Show viscosity-drop labels only when the chart is not crowded.
            if len(chart_df) <= 30:
                label_y = (chart_df["黏度(秒)"] + chart_df["黏度(秒)_1"]) / 2
                fig_viscosity.add_trace(go.Scatter(
                    x=chart_df["Record_Index"],
                    y=label_y,
                    mode="text",
                    text=chart_df["Delta_V"].map(lambda v: f"{v:.0f}" if pd.notna(v) else ""),
                    textfont=dict(size=9, color="#374151"),
                    name="Viscosity Drop (s)",
                    showlegend=False,
                    hoverinfo="skip",
                ))

            fig_viscosity.update_layout(
                title=(
                    "<b>Viscosity Adjustment by Historical Record</b>"
                    f"<br><sup>{detail_title_filter} | {chart_range_text}</sup>"
                ),
                template="plotly_white",
                height=460,
                margin=dict(l=75, r=35, t=105, b=55),
                xaxis=dict(
                    title="Historical Record Order",
                    tickmode="linear" if len(chart_df) <= 30 else "auto",
                    dtick=1 if len(chart_df) <= 30 else None,
                    showgrid=False,
                    showline=True,
                    linecolor="#111827",
                    linewidth=1,
                    mirror=True,
                ),
                yaxis=dict(
                    title="Viscosity (s)",
                    gridcolor="#E5E7EB",
                    showline=True,
                    linecolor="#111827",
                    linewidth=1,
                    mirror=True,
                ),
                legend=dict(
                    orientation="h",
                    yanchor="bottom",
                    y=1.02,
                    xanchor="right",
                    x=1,
                ),
                hovermode="closest",
            )
            st.plotly_chart(fig_viscosity, use_container_width=True)
            exported_figs["5A. Viscosity Adjustment History"] = fig_viscosity

            # =========================================================
            # CHART 2 — SOLVENT RATIO BARS + TEMPERATURE LINE
            # =========================================================
            fig_condition = go.Figure()
            fig_condition.add_trace(go.Bar(
                x=chart_df["Record_Index"],
                y=chart_df["Solvent_Ratio_Percent"],
                name="Solvent Ratio (%)",
                marker_color="#F87171",
                opacity=0.78,
                yaxis="y1",
                customdata=np.column_stack([
                    chart_df["添加重量"],
                    chart_df["Base_Paint_Weight_kg"],
                ]),
                hovertemplate=(
                    "<b>Record %{x}</b><br>"
                    "Solvent Ratio: %{y:.2f}%<br>"
                    "Solvent Added: %{customdata[0]:.1f} kg<br>"
                    "Base Paint: %{customdata[1]:.1f} kg"
                    "<extra></extra>"
                ),
            ))

            if chart_df["溫度"].notna().any():
                fig_condition.add_trace(go.Scatter(
                    x=chart_df["Record_Index"],
                    y=chart_df["溫度"],
                    mode="lines+markers",
                    name="Temperature (°C)",
                    line=dict(color="#7E22CE", width=2),
                    marker=dict(size=6, color="#7E22CE", line=dict(width=0.7, color="white")),
                    yaxis="y2",
                    connectgaps=False,
                    hovertemplate=(
                        "<b>Record %{x}</b><br>"
                        "Temperature: %{y:.1f} °C"
                        "<extra></extra>"
                    ),
                ))

            ratio_max = pd.to_numeric(chart_df["Solvent_Ratio_Percent"], errors="coerce").max()
            ratio_upper = max(5.0, float(ratio_max) * 1.18) if pd.notna(ratio_max) else 5.0

            temp_values = pd.to_numeric(chart_df["溫度"], errors="coerce").dropna()
            if not temp_values.empty:
                temp_min = float(temp_values.min())
                temp_max = float(temp_values.max())
                temp_pad = max(1.0, (temp_max - temp_min) * 0.15)
                temp_range = [temp_min - temp_pad, temp_max + temp_pad]
            else:
                temp_range = None

            fig_condition.update_layout(
                title=(
                    "<b>Solvent Ratio and Temperature by Historical Record</b>"
                    f"<br><sup>{chart_range_text}</sup>"
                ),
                template="plotly_white",
                height=380,
                margin=dict(l=75, r=85, t=95, b=55),
                barmode="overlay",
                xaxis=dict(
                    title="Historical Record Order",
                    tickmode="linear" if len(chart_df) <= 30 else "auto",
                    dtick=1 if len(chart_df) <= 30 else None,
                    showgrid=False,
                    showline=True,
                    linecolor="#111827",
                    linewidth=1,
                    mirror=True,
                ),
                yaxis=dict(
                    title="Solvent Ratio (%)",
                    range=[0, ratio_upper],
                    gridcolor="#FEE2E2",
                    showline=True,
                    linecolor="#F87171",
                    linewidth=1,
                    tickfont=dict(color="#DC2626"),
                    title_font=dict(color="#DC2626"),
                ),
                yaxis2=dict(
                    title="Temperature (°C)",
                    overlaying="y",
                    side="right",
                    range=temp_range,
                    showgrid=False,
                    showline=True,
                    linecolor="#7E22CE",
                    linewidth=1,
                    tickfont=dict(color="#7E22CE"),
                    title_font=dict(color="#7E22CE"),
                ),
                legend=dict(
                    orientation="h",
                    yanchor="bottom",
                    y=1.02,
                    xanchor="right",
                    x=1,
                ),
                hovermode="x unified",
            )
            st.plotly_chart(fig_condition, use_container_width=True)
            exported_figs["5B. Solvent Ratio and Temperature"] = fig_condition

            st.caption(
                "Note: Each record represents one independent viscosity-adjustment event. "
                "The vertical segment in the upper chart represents the viscosity drop (Before − After)."
            )

            detail_columns = [
                col for col in [
                    "Record_Index",
                    "_Analysis_DateTime",
                    "Batch_ID",
                    "Bucket_Number",
                    "線別",
                    "黏度(秒)",
                    "黏度(秒)_1",
                    "Delta_V",
                    "溫度",
                    "Base_Paint_Weight_kg",
                    "添加重量",
                    "Solvent_Ratio_Percent",
                ]
                if col in detail_df.columns
            ]
            detail_table = detail_df[detail_columns].copy().rename(columns={
                "Record_Index": "Record",
                "_Analysis_DateTime": "Adjustment Time",
                "Batch_ID": "Paint Batch",
                "Bucket_Number": "Bucket Number",
                "線別": "Production Line",
                "黏度(秒)": "Before Viscosity (s)",
                "黏度(秒)_1": "After Viscosity (s)",
                "Delta_V": "Viscosity Drop (s)",
                "溫度": "Temperature (°C)",
                "Base_Paint_Weight_kg": "Base Paint Weight (kg)",
                "添加重量": "Solvent Added (kg)",
                "Solvent_Ratio_Percent": "Solvent Ratio (%)",
            })

            with st.expander("View Historical Records"):
                st.dataframe(
                    detail_table.style.format({
                        "Before Viscosity (s)": "{:.2f}",
                        "After Viscosity (s)": "{:.2f}",
                        "Viscosity Drop (s)": "{:.2f}",
                        "Temperature (°C)": "{:.2f}",
                        "Base Paint Weight (kg)": "{:,.2f}",
                        "Solvent Added (kg)": "{:,.2f}",
                        "Solvent Ratio (%)": "{:.2f}",
                    }),
                    use_container_width=True,
                    hide_index=True,
                )


# ----- TAB 3: LINE COMPARISON -----
with tab_line:
    st.subheader("3. Production Line Comparison")

    # Only show paint codes that have valid records on at least two production lines
    # within the current global filter selection.
    comparable_code_summary = (
        filter_df.loc[filter_df["線別"] != "Unknown"]
        .groupby("Paint_Code", dropna=False)
        .agg(
            Production_Line_Count=("線別", "nunique"),
            Comparison_Records=("Paint_Code", "size"),
            Total_Solvent_kg=("添加重量", "sum"),
        )
        .reset_index()
    )

    comparable_code_summary = comparable_code_summary[
        comparable_code_summary["Production_Line_Count"] >= 2
    ].sort_values(
        ["Production_Line_Count", "Comparison_Records", "Total_Solvent_kg", "Paint_Code"],
        ascending=[False, False, False, True],
    )

    comparable_codes = comparable_code_summary["Paint_Code"].tolist()

    if not comparable_codes:
        st.info(
            "No paint code is currently used on at least two production lines "
            "within the selected global filters."
        )
    else:
        st.caption(
            f"Only paint codes used on at least two production lines are listed "
            f"({len(comparable_codes):,} comparable codes)."
        )

        comp_code = st.selectbox(
            "Select Paint Code for Line Comparison",
            comparable_codes,
            key="line_comp",
        )

        comp_df = filter_df[
            (filter_df["Paint_Code"] == comp_code)
            & (filter_df["線別"] != "Unknown")
        ].copy()
        line_summary = build_summary(comp_df, ["線別"]).sort_values("線別")

        ch5, ch6 = st.columns(2)
        with ch5:
            fig5 = go.Figure()
            for _, row in line_summary.iterrows():
                fig5.add_trace(
                    go.Scatter(
                        x=[row["Median_After_Viscosity"], row["Median_Before_Viscosity"]],
                        y=[row["線別"], row["線別"]],
                        mode="lines+markers",
                        marker=dict(size=12),
                        name=row["線別"],
                    )
                )
            fig5.update_layout(
                title="Viscosity Drop (Before vs After)",
                xaxis_title="Viscosity (s)",
                yaxis_title="",
            )
            st.plotly_chart(fig5, use_container_width=True)
            exported_figs["7. Line Comparison - Viscosity Drop"] = fig5

        with ch6:
            fig6 = px.bar(
                line_summary.sort_values("Weighted_Ratio_Percent"),
                x="Weighted_Ratio_Percent",
                y="線別",
                orientation="h",
                text_auto=".2f",
                title="Weighted Solvent Ratio",
                color_discrete_sequence=["#2563EB"],
            )
            fig6.update_traces(texttemplate="%{x:.2f}%", textposition="outside")
            fig6.update_xaxes(title="Weighted Solvent Ratio (%)")
            fig6.update_yaxes(title="")
            st.plotly_chart(fig6, use_container_width=True)
            exported_figs["8. Line Comparison - Solvent Ratio"] = fig6

# ----- TAB 4: PILOT PAINT CODE EVALUATION -----
with tab_pilot:
    st.subheader("4. Supplier Incoming Viscosity Improvement Priority")
    st.markdown(
        "Identify paint codes that repeatedly require substantial on-site solvent adjustment. "
        "These paint codes are candidates for discussion with the supplier so the incoming viscosity can be improved "
        "and operators may reduce or eliminate routine dilution and readjustment work."
    )

    set1, set2, set3 = st.columns(3)
    with set1:
        min_pilot_records = st.number_input(
            "Minimum Historical Records", min_value=5, value=20, step=1
        )
    with set2:
        target_solvent_limit = st.number_input(
            "High Solvent Adjustment Threshold (kg)", min_value=0.0, value=500.0, step=100.0
        )
    with set3:
        min_peer_codes = st.number_input(
            "Minimum Comparable Paint Codes", min_value=3, value=5, step=1,
            help="At least this many paint codes are required within the same Vendor, Resin, Position, and Solvent Type group to use peer-group P33/P67. Otherwise, the overall eligible distribution is used."
        )

    st.info(
        "Decision focus: total on-site solvent adjustment, historical batch coverage, typical incoming viscosity, "
        "typical required viscosity after adjustment, and the repeatability of the adjustment pattern."
    )

    def robust_mad(series):
        values = pd.to_numeric(series, errors="coerce").dropna()
        if values.empty:
            return np.nan
        median_value = values.median()
        return float(np.median(np.abs(values - median_value)))

    def ratio_consistency_stats(series):
        """Return the typical ratio, data-driven stable range, and consistency share."""
        values = pd.to_numeric(series, errors="coerce").dropna()
        if len(values) < 5:
            return {
                "median": np.nan,
                "lower": np.nan,
                "upper": np.nan,
                "consistency": np.nan,
            }

        median_value = float(values.median())
        mad = robust_mad(values)
        tolerance = max(1.0, 1.5 * mad)
        lower = max(0.0, median_value - tolerance)
        upper = median_value + tolerance
        consistency = values.between(lower, upper, inclusive="both").mean()

        return {
            "median": median_value,
            "lower": float(lower),
            "upper": float(upper),
            "consistency": float(consistency),
        }

    def ratio_consistency(series):
        return ratio_consistency_stats(series)["consistency"]

    def ratio_stable_lower(series):
        return ratio_consistency_stats(series)["lower"]

    def ratio_stable_upper(series):
        return ratio_consistency_stats(series)["upper"]

    def robust_relative_variation(series):
        values = pd.to_numeric(series, errors="coerce").dropna()
        if len(values) < 5:
            return np.nan
        median_value = values.median()
        if median_value <= 0:
            return np.nan
        return robust_mad(values) / median_value

    group_keys = ["Vendor", "Resin", "Position_UI", "Solvent_Type", "Paint_Code"]
    peer_keys = ["Vendor", "Resin", "Position_UI", "Solvent_Type"]

    supplier_df = filter_df.groupby(group_keys, dropna=False).agg(
        Historical_Records=("Paint_Code", "size"),
        Historical_Batches=("Batch_ID", "nunique"),
        Total_Paint_kg=("Base_Paint_Weight_kg", "sum"),
        Total_Solvent_kg=("添加重量", "sum"),
        Median_Ratio_Percent=("Solvent_Ratio_Percent", "median"),
        Stable_Ratio_Lower=("Solvent_Ratio_Percent", ratio_stable_lower),
        Stable_Ratio_Upper=("Solvent_Ratio_Percent", ratio_stable_upper),
        Median_Incoming_Viscosity=("黏度(秒)", "median"),
        Median_Required_Viscosity=("黏度(秒)_1", "median"),
        Median_Viscosity_Drop=("Delta_V", "median"),
        Ratio_Consistency=("Solvent_Ratio_Percent", ratio_consistency),
        Efficiency_Relative_Variation=("Viscosity_Sensitivity", robust_relative_variation),
    ).reset_index()

    def trend_per_10_records(group):
        temp = group.dropna(subset=["Solvent_Ratio_Percent"]).sort_values(
            ["_Analysis_DateTime", "_Source_Order"], na_position="last"
        )
        if len(temp) < 5:
            return np.nan
        x = np.arange(len(temp), dtype=float)
        y = temp["Solvent_Ratio_Percent"].to_numpy(dtype=float)
        return float(np.polyfit(x, y, 1)[0] * 10)

    trend_df = (
        filter_df.groupby(group_keys, dropna=False)
        .apply(trend_per_10_records)
        .reset_index(name="Ratio_Trend_Per_10_Records")
    )
    supplier_df = supplier_df.merge(trend_df, on=group_keys, how="left")
    supplier_df["Stable_Ratio_Range"] = supplier_df.apply(
        lambda row: (
            f"{row['Stable_Ratio_Lower']:.2f}–{row['Stable_Ratio_Upper']:.2f}"
            if pd.notna(row["Stable_Ratio_Lower"]) and pd.notna(row["Stable_Ratio_Upper"])
            else "Insufficient Data"
        ),
        axis=1,
    )
    supplier_df["Abs_Ratio_Trend_Per_10_Records"] = supplier_df["Ratio_Trend_Per_10_Records"].abs()
    supplier_df["Weighted_Ratio_Percent"] = np.where(
        supplier_df["Total_Paint_kg"] > 0,
        supplier_df["Total_Solvent_kg"] / supplier_df["Total_Paint_kg"] * 100,
        np.nan,
    )

    eligible_df = supplier_df[
        supplier_df["Historical_Records"] >= min_pilot_records
    ].copy()

    metric_rules = {
        "Ratio_Consistency": "higher",
        "Efficiency_Relative_Variation": "lower",
        "Abs_Ratio_Trend_Per_10_Records": "lower",
    }

    global_thresholds = {}
    for metric in metric_rules:
        valid = pd.to_numeric(eligible_df[metric], errors="coerce").dropna()
        global_thresholds[metric] = {
            "P33": valid.quantile(0.33) if not valid.empty else np.nan,
            "P67": valid.quantile(0.67) if not valid.empty else np.nan,
        }

    peer_stats = (
        eligible_df.groupby(peer_keys, dropna=False)
        .agg(
            Peer_Code_Count=("Paint_Code", "nunique"),
            Ratio_Consistency_P33=("Ratio_Consistency", lambda x: x.dropna().quantile(0.33)),
            Ratio_Consistency_P67=("Ratio_Consistency", lambda x: x.dropna().quantile(0.67)),
            Efficiency_Variation_P33=("Efficiency_Relative_Variation", lambda x: x.dropna().quantile(0.33)),
            Efficiency_Variation_P67=("Efficiency_Relative_Variation", lambda x: x.dropna().quantile(0.67)),
            Trend_Abs_P33=("Abs_Ratio_Trend_Per_10_Records", lambda x: x.dropna().quantile(0.33)),
            Trend_Abs_P67=("Abs_Ratio_Trend_Per_10_Records", lambda x: x.dropna().quantile(0.67)),
        )
        .reset_index()
    )
    supplier_df = supplier_df.merge(peer_stats, on=peer_keys, how="left")

    def select_thresholds(row, metric):
        use_peer = pd.notna(row.get("Peer_Code_Count")) and row["Peer_Code_Count"] >= min_peer_codes
        mapping = {
            "Ratio_Consistency": ("Ratio_Consistency_P33", "Ratio_Consistency_P67"),
            "Efficiency_Relative_Variation": ("Efficiency_Variation_P33", "Efficiency_Variation_P67"),
            "Abs_Ratio_Trend_Per_10_Records": ("Trend_Abs_P33", "Trend_Abs_P67"),
        }
        p33_col, p67_col = mapping[metric]
        if use_peer and pd.notna(row.get(p33_col)) and pd.notna(row.get(p67_col)):
            return row[p33_col], row[p67_col], "Peer Group Distribution"
        return global_thresholds[metric]["P33"], global_thresholds[metric]["P67"], "Overall Distribution Fallback"

    def classify_by_distribution(value, p33, p67, direction):
        if pd.isna(value) or pd.isna(p33) or pd.isna(p67):
            return "Insufficient Data"
        if direction == "higher":
            if value >= p67:
                return "High Stability"
            if value >= p33:
                return "Medium Stability"
            return "Low Stability"
        if value <= p33:
            return "High Stability"
        if value <= p67:
            return "Medium Stability"
        return "Low Stability"

    def evaluate_stability_row(row):
        ratio_p33, ratio_p67, ratio_source = select_thresholds(row, "Ratio_Consistency")
        eff_p33, eff_p67, eff_source = select_thresholds(row, "Efficiency_Relative_Variation")
        trend_p33, trend_p67, trend_source = select_thresholds(row, "Abs_Ratio_Trend_Per_10_Records")

        ratio_level = classify_by_distribution(row["Ratio_Consistency"], ratio_p33, ratio_p67, "higher")
        efficiency_level = classify_by_distribution(row["Efficiency_Relative_Variation"], eff_p33, eff_p67, "lower")
        trend_level = classify_by_distribution(row["Abs_Ratio_Trend_Per_10_Records"], trend_p33, trend_p67, "lower")

        levels = [ratio_level, efficiency_level, trend_level]
        if "Insufficient Data" in levels:
            overall = "Insufficient Data"
            high_count = np.nan
        else:
            high_count = sum(level == "High Stability" for level in levels)
            if high_count == 3:
                overall = "High Stability"
            elif "Low Stability" not in levels:
                overall = "Medium Stability"
            else:
                overall = "Low Stability"

        sources = {ratio_source, eff_source, trend_source}
        benchmark_source = "Peer Group Distribution" if sources == {"Peer Group Distribution"} else "Overall Distribution Fallback"

        return pd.Series({
            "Ratio_Stability_Level": ratio_level,
            "Efficiency_Stability_Level": efficiency_level,
            "Trend_Stability_Level": trend_level,
            "High_Stability_Count": high_count,
            "Stability_Level": overall,
            "Stability_Benchmark_Source": benchmark_source,
        })

    supplier_df = pd.concat([supplier_df, supplier_df.apply(evaluate_stability_row, axis=1)], axis=1)
    supplier_df = supplier_df[supplier_df["Historical_Records"] >= min_pilot_records].copy()

    def classify_supplier_action(row):
        high_adjustment = row["Total_Solvent_kg"] >= target_solvent_limit
        high_stability = row["Stability_Level"] == "High Stability"
        if high_adjustment and high_stability:
            return "High Supplier Priority"
        if high_adjustment and not high_stability:
            return "Validate with Supplier"
        if not high_adjustment and high_stability:
            return "Monitor"
        return "Low Priority"

    if not supplier_df.empty:
        supplier_df["Supplier_Action"] = supplier_df.apply(classify_supplier_action, axis=1)

        color_map = {
            "High Supplier Priority": "#1D4ED8",
            "Validate with Supplier": "#F97316",
            "Monitor": "#60A5FA",
            "Low Priority": "#FDBA74",
        }

        action_order = {
            "High Supplier Priority": 1,
            "Validate with Supplier": 2,
            "Monitor": 3,
            "Low Priority": 4,
        }
        supplier_df["Action_Order"] = supplier_df["Supplier_Action"].map(action_order).fillna(99)

        # =========================================================
        # TABLE 2 — TOP 10 PAINT CODE STABILITY RESULTS
        # =========================================================
        st.markdown("### 表2　主要色號穩定性評估結果")
        st.caption(
            "依累積稀釋劑添加量由高至低，僅列示前10個主要色號；"
            "各指標同時呈現實際數值及穩定性分級。"
        )

        stability_class_map = {
            "High Stability": "高",
            "Medium Stability": "中",
            "Low Stability": "低",
            "Insufficient Data": "資料不足",
        }
        overall_stability_map = {
            "High Stability": "高穩定",
            "Medium Stability": "需確認",
            "Low Stability": "需確認",
            "Insufficient Data": "資料不足",
        }

        top10_stability_df = (
            supplier_df
            .sort_values(
                ["Total_Solvent_kg", "Historical_Batches", "Historical_Records"],
                ascending=[False, False, False],
            )
            .head(10)
            .copy()
            .reset_index(drop=True)
        )
        top10_stability_df.insert(0, "Rank", np.arange(1, len(top10_stability_df) + 1))

        def format_metric_with_level(value, level, value_format):
            if pd.isna(value):
                return "資料不足"
            level_zh = stability_class_map.get(level, str(level))
            return f"{value_format(value)}（{level_zh}）"

        top10_stability_df["Ratio_Consistency_Display"] = top10_stability_df.apply(
            lambda row: format_metric_with_level(
                row["Ratio_Consistency"],
                row["Ratio_Stability_Level"],
                lambda value: f"{value * 100:.1f}%",
            ),
            axis=1,
        )
        top10_stability_df["Efficiency_Variation_Display"] = top10_stability_df.apply(
            lambda row: format_metric_with_level(
                row["Efficiency_Relative_Variation"],
                row["Efficiency_Stability_Level"],
                lambda value: f"{value:.4f}",
            ),
            axis=1,
        )
        top10_stability_df["Ratio_Trend_Display"] = top10_stability_df.apply(
            lambda row: format_metric_with_level(
                row["Ratio_Trend_Per_10_Records"],
                row["Trend_Stability_Level"],
                lambda value: f"{value:+.2f}",
            ),
            axis=1,
        )
        top10_stability_df["High_Stability_Display"] = top10_stability_df[
            "High_Stability_Count"
        ].apply(lambda value: "資料不足" if pd.isna(value) else f"{int(value)}/3")
        top10_stability_df["Overall_Stability_Display"] = (
            top10_stability_df["Stability_Level"]
            .map(overall_stability_map)
            .fillna(top10_stability_df["Stability_Level"])
        )

        top10_stability_display = top10_stability_df[[
            "Rank",
            "Paint_Code",
            "Total_Solvent_kg",
            "Ratio_Consistency_Display",
            "Efficiency_Variation_Display",
            "Ratio_Trend_Display",
            "High_Stability_Display",
            "Overall_Stability_Display",
        ]].rename(columns={
            "Rank": "順位",
            "Paint_Code": "色號",
            "Total_Solvent_kg": "累積稀釋劑添加量（kg）",
            "Ratio_Consistency_Display": "添加比例一致率",
            "Efficiency_Variation_Display": "稀釋效率相對變異",
            "Ratio_Trend_Display": "添加比例趨勢（每10筆）",
            "High_Stability_Display": "高穩定指標數",
            "Overall_Stability_Display": "整體判定",
        })

        st.dataframe(
            top10_stability_display,
            column_config={
                "順位": st.column_config.NumberColumn("順位", format="%d", width="small"),
                "色號": st.column_config.TextColumn("色號", width="medium"),
                "累積稀釋劑添加量（kg）": st.column_config.NumberColumn(
                    "累積稀釋劑添加量（kg）", format="%.0f"
                ),
                "添加比例一致率": st.column_config.TextColumn(
                    "添加比例一致率", help="數值越高越穩定"
                ),
                "稀釋效率相對變異": st.column_config.TextColumn(
                    "稀釋效率相對變異", help="數值越低越穩定"
                ),
                "添加比例趨勢（每10筆）": st.column_config.TextColumn(
                    "添加比例趨勢（每10筆）", help="絕對值越低越穩定"
                ),
                "高穩定指標數": st.column_config.TextColumn(
                    "高穩定指標數", width="small"
                ),
                "整體判定": st.column_config.TextColumn("整體判定", width="small"),
            },
            use_container_width=True,
            hide_index=True,
        )
        st.caption(
            "註：添加比例一致率數值越高越佳；稀釋效率相對變異及"
            "添加比例趨勢絕對值越低越佳。僅3/3判定為整體高穩定。"
        )

        plot_df = supplier_df.dropna(subset=["High_Stability_Count"]).copy()
        if not plot_df.empty:
            plot_df = plot_df.sort_values(
                ["High_Stability_Count", "Total_Solvent_kg", "Historical_Batches"],
                ascending=[True, False, False],
            ).reset_index(drop=True)
            jitter_pattern = [0.00, -0.14, 0.14, -0.22, 0.22, -0.30, 0.30]
            plot_df["Display_Order_In_X"] = plot_df.groupby("High_Stability_Count").cumcount()
            plot_df["Matrix_X"] = (
                plot_df["High_Stability_Count"].astype(float)
                + plot_df["Display_Order_In_X"].map(lambda i: jitter_pattern[int(i) % len(jitter_pattern)])
            )

            fig_matrix = px.scatter(
                plot_df,
                x="Matrix_X",
                y="Total_Solvent_kg",
                size="Historical_Batches",
                size_max=38,
                color="Supplier_Action",
                color_discrete_map=color_map,
                hover_name="Paint_Code",
                custom_data=[
                    "High_Stability_Count", "Historical_Batches", "Historical_Records",
                    "Median_Incoming_Viscosity", "Median_Required_Viscosity",
                    "Median_Ratio_Percent", "Ratio_Consistency", "Stability_Level",
                    "Total_Solvent_kg",
                ],
                title=None,
                height=760,
            )
            fig_matrix.update_traces(
                mode="markers",
                marker=dict(opacity=0.86, line=dict(width=1.3, color="white")),
                hovertemplate=(
                    "<b>%{hovertext}</b><br>──────────────────<br>"
                    "High-Stability Checks: %{customdata[0]:.0f} of 3<br>"
                    "Historical Batches: %{customdata[1]:,.0f}<br>"
                    "Historical Records: %{customdata[2]:,.0f}<br>"
                    "Typical Incoming Viscosity: %{customdata[3]:.1f} s<br>"
                    "Typical Required Viscosity: %{customdata[4]:.1f} s<br>"
                    "Typical Adjustment Ratio: %{customdata[5]:.2f}%<br>"
                    "Adjustment Consistency: %{customdata[6]:.1%}<br>"
                    "Overall Stability: %{customdata[7]}<br>"
                    "Total On-site Solvent Adjustment: %{customdata[8]:,.1f} kg"
                    "<extra></extra>"
                ),
            )

            apply_professional_layout(
                fig_matrix,
                title_text=None,
                subtitle_text=None,
                height=780,
            )
            fig_matrix.update_xaxes(
                title="High-Stability Checks Passed (0-3)",
                tickmode="array", tickvals=[0, 1, 2, 3], ticktext=["0", "1", "2", "3"],
                range=[-0.45, 3.45], mirror=True, showline=True,
                linecolor="#111827", linewidth=1.2,
                title_font=dict(color="#000000", size=14),
                tickfont=dict(color="#000000", size=12),
            )
            max_y = max(float(plot_df["Total_Solvent_kg"].max()), 1.0)
            fig_matrix.update_yaxes(
                title="Total On-site Solvent Adjustment (kg)",
                range=[-max_y * 0.05, max_y * 1.15], mirror=True, showline=True,
                linecolor="#111827", linewidth=1.2,
                title_font=dict(color="#000000", size=14),
                tickfont=dict(color="#000000", size=12),
            )
            fig_matrix.update_layout(
                title=dict(
                    text=(
                        "<b>Figure 1. Supplier Incoming Viscosity Improvement Priority Matrix</b>"
                        "<br><sup>X-axis = High-Stability Checks Passed; "
                        "Y-axis = Total On-site Solvent Adjustment; "
                        "Bubble Size = Historical Batches</sup>"
                    ),
                    x=0.01, xanchor="left", y=0.975, yanchor="top",
                    font=dict(size=20, color="#000000"),
                    pad=dict(b=18),
                ),
                margin=dict(l=80, r=255, t=205, b=90),
                legend=dict(
                    title_text="Action Category",
                    orientation="v",
                    yanchor="top", y=1.00,
                    xanchor="left", x=1.015,
                    font=dict(color="#000000", size=12),
                    title_font=dict(color="#000000", size=12),
                    bgcolor="rgba(255,255,255,0.98)",
                    bordercolor="#D1D5DB",
                    borderwidth=1,
                ),
            )
            fig_matrix.add_vline(x=2.5, line_dash="dash", line_color="#D62728", line_width=1.6, opacity=0.85)
            fig_matrix.add_hline(y=target_solvent_limit, line_dash="dash", line_color="#D62728", line_width=1.6, opacity=0.85)
            fig_matrix.add_annotation(
                x=2.52, y=max_y * 1.08, text="Strong Evidence Zone", showarrow=False,
                font=dict(size=12, color="#000000"), bgcolor="rgba(255,255,255,0.95)",
                bordercolor="#D62728", borderwidth=1, borderpad=3,
            )
            # Put the threshold note outside the right chart frame so it never covers data labels.
            fig_matrix.add_annotation(
                x=1.015, xref="paper", y=target_solvent_limit, yref="y",
                text=f"Threshold: {target_solvent_limit:,.0f} kg",
                showarrow=False, xanchor="left", yanchor="middle",
                font=dict(size=11, color="#000000"), bgcolor="rgba(255,255,255,0.97)",
                bordercolor="#D62728", borderwidth=1, borderpad=3,
            )

            label_df = plot_df[plot_df["Supplier_Action"].isin(["High Supplier Priority", "Validate with Supplier"])].copy()
            label_df = label_df.sort_values(["Action_Order", "Total_Solvent_kg"], ascending=[True, False]).head(10)
            label_df["Label_Rank_In_X"] = label_df.groupby("High_Stability_Count").cumcount()
            normal_offsets = [(0, -34), (-48, -58), (48, -58)]
            bottom_offsets = [(-52, -36), (0, -72), (52, -36)]
            for _, row in label_df.iterrows():
                rank = int(row["Label_Rank_In_X"]) % 3
                point_x = float(row["Matrix_X"])
                point_y = float(row["Total_Solvent_kg"])
                offsets = bottom_offsets if point_y <= target_solvent_limit * 1.5 else normal_offsets
                ax_shift, ay_shift = offsets[rank]
                if point_x >= 2.85 and ax_shift > 0:
                    ax_shift = -58
                if point_x <= 0.10 and ax_shift < 0:
                    ax_shift = 48
                fig_matrix.add_annotation(
                    x=point_x, y=point_y, text=str(row["Paint_Code"]),
                    showarrow=True, arrowhead=0, arrowcolor="rgba(0,0,0,0.35)", arrowwidth=1,
                    ax=ax_shift, ay=ay_shift,
                    font=dict(size=11, color="#000000"), bgcolor="rgba(255,255,255,0.97)",
                    bordercolor="rgba(0,0,0,0.25)", borderwidth=1, borderpad=3,
                )

            matrix_export_plot_df = plot_df.copy()
            st.plotly_chart(fig_matrix, use_container_width=True)
            exported_figs["9. Supplier Priority Matrix"] = fig_matrix
        else:
            st.warning("⚠️ Stability data are insufficient to generate the supplier priority matrix.")

        st.markdown("---")
        st.markdown("### Supplier Decision Summary")
        st.caption("Read from left to right: recommended action → solvent adjustment burden → incoming viscosity → required viscosity → repeatability.")

        summary_cols = [
            "Paint_Code", "Supplier_Action", "Total_Solvent_kg", "Historical_Batches",
            "Median_Incoming_Viscosity", "Median_Required_Viscosity",
            "Median_Ratio_Percent", "Stable_Ratio_Range", "Ratio_Consistency",
        ]
        display_df = supplier_df.sort_values(
            ["Action_Order", "Total_Solvent_kg", "Historical_Batches"],
            ascending=[True, False, False],
        )[summary_cols].copy()
        display_df["Ratio_Consistency"] = display_df["Ratio_Consistency"] * 100

        st.dataframe(
            display_df,
            column_config={
                "Paint_Code": "Paint Code",
                "Supplier_Action": "Recommended Action",
                "Total_Solvent_kg": st.column_config.NumberColumn("Total Solvent Adjustment (kg)", format="%.1f"),
                "Historical_Batches": "Historical Batches",
                "Median_Incoming_Viscosity": st.column_config.NumberColumn("Typical Incoming Viscosity (s)", format="%.1f"),
                "Median_Required_Viscosity": st.column_config.NumberColumn("Typical Required Viscosity (s)", format="%.1f"),
                "Median_Ratio_Percent": st.column_config.NumberColumn("Typical Adjustment Ratio (%)", format="%.2f"),
                "Stable_Ratio_Range": st.column_config.TextColumn(
                    "Stable Ratio Range (%)",
                    help="Data-driven range used to calculate Adjustment Consistency. Records inside this range are counted as consistent.",
                ),
                "Ratio_Consistency": st.column_config.NumberColumn(
                    "Adjustment Consistency (%)",
                    format="%.1f",
                    help="Percentage of historical records whose solvent ratio falls inside the Stable Ratio Range.",
                ),
            },
            use_container_width=True,
            hide_index=True,
        )

        high_priority = supplier_df[supplier_df["Supplier_Action"] == "High Supplier Priority"].sort_values(
            "Total_Solvent_kg", ascending=False
        )
        if not high_priority.empty:
            top_code = high_priority.iloc[0]
            st.success(
                f"✅ Highest supplier priority: {top_code['Paint_Code']}. "
                f"This paint code required {top_code['Total_Solvent_kg']:.1f} kg of on-site solvent adjustment across "
                f"{int(top_code['Historical_Batches'])} historical batches. Typical viscosity was adjusted from "
                f"{top_code['Median_Incoming_Viscosity']:.1f} s to {top_code['Median_Required_Viscosity']:.1f} s."
            )
        else:
            st.info("No paint code currently combines a high on-site solvent adjustment burden with High Stability evidence.")

        with st.expander("View Technical Stability Details"):
            technical_cols = [
                "Paint_Code", "Stability_Level", "High_Stability_Count", "Stability_Benchmark_Source",
                "Peer_Code_Count", "Median_Ratio_Percent", "Stable_Ratio_Range",
                "Ratio_Consistency", "Ratio_Stability_Level",
                "Efficiency_Relative_Variation", "Efficiency_Stability_Level",
                "Ratio_Trend_Per_10_Records", "Trend_Stability_Level",
                "Historical_Records", "Historical_Batches",
            ]
            technical_df = supplier_df.sort_values(
                ["Action_Order", "Total_Solvent_kg"], ascending=[True, False]
            )[technical_cols].copy()
            technical_df["Ratio_Consistency"] = technical_df["Ratio_Consistency"] * 100
            technical_df["High_Stability_Count_Display"] = technical_df["High_Stability_Count"].apply(
                lambda x: "Insufficient Data" if pd.isna(x) else f"{int(x)} of 3"
            )
            technical_df = technical_df.drop(columns=["High_Stability_Count"])
            st.dataframe(
                technical_df,
                column_config={
                    "Paint_Code": "Paint Code",
                    "Stability_Level": "Overall Stability",
                    "High_Stability_Count_Display": "High-Stability Checks Passed",
                    "Stability_Benchmark_Source": "Comparison Basis",
                    "Peer_Code_Count": "Comparable Paint Codes",
                    "Median_Ratio_Percent": st.column_config.NumberColumn("Typical Adjustment Ratio (%)", format="%.2f"),
                    "Stable_Ratio_Range": st.column_config.TextColumn(
                        "Stable Ratio Range (%)",
                        help="Median ratio ± the data-driven tolerance used for the consistency calculation.",
                    ),
                    "Ratio_Consistency": st.column_config.NumberColumn(
                        "Ratio Consistency (%)",
                        format="%.1f",
                        help="Share of records located within the Stable Ratio Range.",
                    ),
                    "Ratio_Stability_Level": "Ratio Stability",
                    "Efficiency_Relative_Variation": st.column_config.NumberColumn("Efficiency Relative Variation", format="%.3f"),
                    "Efficiency_Stability_Level": "Efficiency Stability",
                    "Ratio_Trend_Per_10_Records": st.column_config.NumberColumn("Ratio Trend per 10 Records", format="%+.2f"),
                    "Trend_Stability_Level": "Trend Stability",
                    "Historical_Records": "Historical Records",
                    "Historical_Batches": "Historical Batches",
                },
                use_container_width=True,
                hide_index=True,
            )

        with st.expander("View P33/P67 Thresholds Used in This Analysis"):
            threshold_rows = []
            for metric, label, direction in [
                ("Ratio_Consistency", "Ratio Consistency", "Higher is better"),
                ("Efficiency_Relative_Variation", "Efficiency Relative Variation", "Lower is better"),
                ("Abs_Ratio_Trend_Per_10_Records", "Absolute Time Trend", "Lower is better"),
            ]:
                threshold_rows.append({
                    "Indicator": label,
                    "Direction": direction,
                    "Filtered Dataset P33": global_thresholds[metric]["P33"],
                    "Filtered Dataset P67": global_thresholds[metric]["P67"],
                })
            threshold_df = pd.DataFrame(threshold_rows)
            st.dataframe(
                threshold_df,
                column_config={
                    "Indicator": "Indicator",
                    "Direction": "Direction",
                    "Filtered Dataset P33": st.column_config.NumberColumn("Filtered Dataset P33", format="%.2f"),
                    "Filtered Dataset P67": st.column_config.NumberColumn("Filtered Dataset P67", format="%.2f"),
                },
                use_container_width=True,
                hide_index=True,
            )
    else:
        display_df = pd.DataFrame()
        st.warning("⚠️ Historical records are insufficient to generate the supplier priority analysis.")


# ==========================================
# 7. EXPORT INTERACTIVE HTML REPORT
# ==========================================
st.markdown("---")
st.subheader("📄 Export Report")

export_col1, export_col2 = st.columns(2)

with export_col1:
    st.info("💡 Word export uses Matplotlib to generate chart images. Kaleido is not used.")
    if HAS_DOCX and st.button("📥 Generate & Download Word Report", type="primary"):
        with st.spinner("⏳ Generating Word report..."):
            try:
                doc = Document()

                # A4 landscape gives the decision table enough width and prevents column clipping.
                section = doc.sections[0]
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width, section.page_height = section.page_height, section.page_width
                section.top_margin = Inches(0.45)
                section.bottom_margin = Inches(0.45)
                section.left_margin = Inches(0.45)
                section.right_margin = Inches(0.45)

                title_p = doc.add_paragraph()
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_p.add_run("Supplier Incoming Viscosity Improvement Priority Report")
                title_run.bold = True
                title_run.font.size = Pt(18)

                p = doc.add_paragraph()
                p.add_run(f"Analysis Date: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                p.add_run(f"Filters Applied: {filter_details}\n")
                p.add_run("Note: Charts in this report are exported with Matplotlib. Kaleido is not used.")

                if 'summary_df' in locals() and not summary_df.empty:
                    doc.add_heading("1. Overall Solvent Usage", level=1)

                    figure_title = doc.add_paragraph()
                    figure_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    figure_title_run = figure_title.add_run(
                        "Figure 1. Paint & Solvent Usage vs. Solvent Ratio"
                    )
                    figure_title_run.bold = True
                    figure_title_run.font.size = Pt(11)

                    figure_filter = doc.add_paragraph()
                    figure_filter.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    figure_filter_run = figure_filter.add_run(
                        f"Filters Applied: {filter_details}"
                    )
                    figure_filter_run.italic = True
                    figure_filter_run.font.size = Pt(8.5)

                    top10_chart_buffer = create_top10_usage_ratio_png(summary_df, filter_details)
                    picture_p = doc.add_paragraph()
                    picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    picture_p.paragraph_format.space_before = Pt(2)
                    picture_p.paragraph_format.space_after = Pt(2)
                    picture_p.add_run().add_picture(top10_chart_buffer, width=Inches(10.2))
                else:
                    doc.add_paragraph("No Top 10 usage data is available for export.")

                if 'chart_df' in locals() and chart_df is not None and not chart_df.empty:
                    doc.add_page_break()
                    doc.add_heading("2. Paint Code Adjustment History", level=1)

                    history_context = doc.add_paragraph()
                    history_context.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    context_run = history_context.add_run(
                        f"Paint Code: {selected_code} | {chart_range_text} | Filters: {filter_details}"
                    )
                    context_run.italic = True
                    context_run.font.size = Pt(8.5)

                    figure_title = doc.add_paragraph()
                    figure_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = figure_title.add_run(
                        "Figure 2. Viscosity Adjustment by Historical Record"
                    )
                    title_run.bold = True
                    title_run.font.size = Pt(11)

                    viscosity_buffer = create_viscosity_history_png(chart_df)
                    viscosity_p = doc.add_paragraph()
                    viscosity_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    viscosity_p.paragraph_format.space_before = Pt(2)
                    viscosity_p.paragraph_format.space_after = Pt(4)
                    viscosity_p.add_run().add_picture(viscosity_buffer, width=Inches(10.2))

                    doc.add_page_break()
                    figure_title2 = doc.add_paragraph()
                    figure_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run2 = figure_title2.add_run(
                        "Figure 3. Solvent Ratio and Temperature by Historical Record"
                    )
                    title_run2.bold = True
                    title_run2.font.size = Pt(11)

                    condition_buffer = create_ratio_temperature_png(chart_df)
                    condition_p = doc.add_paragraph()
                    condition_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    condition_p.paragraph_format.space_before = Pt(2)
                    condition_p.paragraph_format.space_after = Pt(2)
                    condition_p.add_run().add_picture(condition_buffer, width=Inches(10.2))

                    note = doc.add_paragraph(
                        "Note: Each record represents one independent viscosity-adjustment event. "
                        "The vertical segment in Figure 2 represents the viscosity drop (Before - After)."
                    )
                    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in note.runs:
                        run.font.size = Pt(8.5)
                else:
                    doc.add_paragraph("No selected paint-code history is available for export.")

                if 'matrix_export_plot_df' in locals() and not matrix_export_plot_df.empty:
                    doc.add_page_break()
                    doc.add_heading("3. Supplier Priority Matrix", level=1)

                    chart_title_p = doc.add_paragraph()
                    chart_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    chart_title_run = chart_title_p.add_run(
                        "Figure 4. Paint Code Stability and Improvement Priority Matrix"
                    )
                    chart_title_run.bold = True
                    chart_title_run.font.size = Pt(13)

                    chart_subtitle_p = doc.add_paragraph()
                    chart_subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    chart_subtitle_run = chart_subtitle_p.add_run(
                        "X-axis = High-Stability Checks Passed; "
                        "Y-axis = Total On-site Solvent Adjustment; "
                        "Bubble Size = Historical Batches"
                    )
                    chart_subtitle_run.font.size = Pt(9)
                    chart_subtitle_run.italic = True

                    chart_buffer = create_supplier_priority_png(
                        matrix_export_plot_df,
                        target_solvent_limit,
                    )
                    picture_p = doc.add_paragraph()
                    picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    picture_p.paragraph_format.space_before = Pt(3)
                    picture_p.add_run().add_picture(chart_buffer, width=Inches(9.8))

                    caption = doc.add_paragraph(
                        "Note: Labels are shown only for key paint codes to avoid overlap."
                    )
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.runs[0].italic = True
                    caption.runs[0].font.size = Pt(8.5)
                else:
                    doc.add_paragraph("No decision matrix data is available for export.")

                if 'display_df' in locals() and not display_df.empty:
                    doc.add_page_break()
                    doc.add_heading("4. Supplier Decision Summary", level=1)
                    doc.add_paragraph(
                        "Decision focus: Recommended Action → Total Solvent Adjustment → "
                        "Incoming/Required Viscosity → Stable Ratio Range → Adjustment Consistency."
                    )
                    # Export every row shown in the app, not only the first 15 rows.
                    export_table_df = display_df.copy()
                    export_table_df = export_table_df[[
                        "Paint_Code", "Supplier_Action", "Total_Solvent_kg", "Historical_Batches",
                        "Median_Incoming_Viscosity", "Median_Required_Viscosity",
                        "Median_Ratio_Percent", "Stable_Ratio_Range", "Ratio_Consistency"
                    ]]
                    export_table_df = export_table_df.rename(columns={
                        "Paint_Code": "Paint Code",
                        "Supplier_Action": "Recommended Action",
                        "Total_Solvent_kg": "Total Solvent Adjustment (kg)",
                        "Historical_Batches": "Historical Batches",
                        "Median_Incoming_Viscosity": "Typical Incoming Viscosity (s)",
                        "Median_Required_Viscosity": "Typical Required Viscosity (s)",
                        "Median_Ratio_Percent": "Typical Adjustment Ratio (%)",
                        "Stable_Ratio_Range": "Stable Ratio Range (%)",
                        "Ratio_Consistency": "Adjustment Consistency (%)",
                    })
                    table = doc.add_table(rows=1, cols=len(export_table_df.columns))
                    table.style = "Table Grid"
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.autofit = True

                    # Repeat the header row when the table spans multiple pages.
                    header_tr = table.rows[0]._tr
                    tr_pr = header_tr.get_or_add_trPr()
                    tbl_header = OxmlElement("w:tblHeader")
                    tbl_header.set(qn("w:val"), "true")
                    tr_pr.append(tbl_header)

                    hdr_cells = table.rows[0].cells
                    for j, col in enumerate(export_table_df.columns):
                        hdr_cells[j].text = str(col)
                        hdr_cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for paragraph in hdr_cells[j].paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(7)

                    for _, row in export_table_df.iterrows():
                        cells = table.add_row().cells
                        for j, col in enumerate(export_table_df.columns):
                            val = row[col]
                            if pd.isna(val):
                                text = ""
                            elif isinstance(val, (float, np.floating)):
                                if "Consistency" in str(col):
                                    text = f"{val:.1f}"
                                elif "Ratio" in str(col):
                                    text = f"{val:.2f}"
                                else:
                                    text = f"{val:.1f}"
                            else:
                                text = str(val)
                            cells[j].text = text
                            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            for paragraph in cells[j].paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.size = Pt(7)

                    table_caption = doc.add_paragraph("Table 1. Supplier decision summary (all filtered paint codes).")
                    table_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table_caption.runs[0].italic = True
                    table_caption.runs[0].font.size = Pt(9)

                word_buffer = io.BytesIO()
                doc.save(word_buffer)
                word_buffer.seek(0)
                st.download_button(
                    label="💾 Download Word Report (.docx)",
                    data=word_buffer,
                    file_name=f"Solvent_Report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.success("✅ Word report generated successfully!")
            except Exception as e:
                st.error(f"❌ Error generating Word report: {e}")
    elif not HAS_DOCX:
        st.warning("python-docx is not installed, so Word export is unavailable.")

with export_col2:
    st.info("💡 Export an HTML report to preserve interactive chart functionality.")
    if st.button("📥 Generate & Download HTML Report"):
        with st.spinner("⏳ Generating HTML report..."):
            try:
                pilot_table_html = display_df.to_html(index=False, classes="summary-table") if 'display_df' in locals() else "<p>No data available.</p>"
                html_content = f"""
                <html>
                <head>
                    <meta charset="utf-8">
                    <title>Supplier Incoming Viscosity Improvement Priority Report</title>
                    <style>
                        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 40px; background-color: #f0f2f6; }}
                        h1 {{ color: #1f77b4; text-align: center; font-size: 32px; }}
                        h2 {{ color: #2c3e50; border-bottom: 2px solid #bdc3c7; padding-bottom: 8px; margin-top: 50px; font-size: 20px; }}
                        .info-box {{ background-color: #ffffff; padding: 20px; border-radius: 10px; margin-bottom: 30px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); border-left: 5px solid #1f77b4; }}
                        .info-box p {{ margin: 8px 0; font-size: 16px; color: #333; }}
                        .chart-container {{ background-color: white; padding: 20px; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); margin-bottom: 30px; width: 100%; overflow: hidden; }}
                        .table-container {{ background-color: white; padding: 20px; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); margin-bottom: 30px; overflow-x: auto; }}
                        .summary-table {{ width: 100%; border-collapse: collapse; font-size: 14px; }}
                        .summary-table th {{ background-color: #2F6B6D; color: white; padding: 10px 8px; border: 1px solid #d9e1e8; text-align: center; }}
                        .summary-table td {{ padding: 9px 8px; border: 1px solid #d9e1e8; text-align: center; }}
                        .summary-table tr:nth-child(even) {{ background-color: #f7f9fb; }}
                    </style>
                </head>
                <body>
                    <h1>📊 Supplier Incoming Viscosity Improvement Priority Report</h1>
                    <div class="info-box">
                        <p><strong>🕒 Analysis Date:</strong> {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
                        <p><strong>🔍 Filters Applied:</strong> {filter_details}</p>
                    </div>
                """
                for i, (fig_title, fig) in enumerate(exported_figs.items()):
                    inc_js = 'cdn' if i == 0 else False
                    fig_html = fig.to_html(full_html=False, include_plotlyjs=inc_js, default_width="100%", default_height="600px")
                    html_content += f"""
                    <h2>{fig_title}</h2>
                    <div class="chart-container">
                        {fig_html}
                    </div>
                    """
                html_content += f"""
                    <h2>Supplier Decision Summary</h2>
                    <div class="table-container">
                        {pilot_table_html}
                    </div>
                """
                html_content += """
                </body>
                </html>
                """
                html_buffer = io.BytesIO(html_content.encode('utf-8'))
                st.download_button(
                    label="💾 Download Interactive Report (.html)",
                    data=html_buffer,
                    file_name=f"Solvent_Report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M')}.html",
                    mime="text/html"
                )
                st.success("✅ HTML report generated successfully!")
            except Exception as e:
                st.error(f"❌ Error generating HTML report: {e}")
