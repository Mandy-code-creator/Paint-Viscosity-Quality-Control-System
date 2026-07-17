import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.cm as cm
import io

# ==========================================
# THƯ VIỆN XUẤT WORD (Kiểm tra xem máy đã cài chưa)
# ==========================================
try:
    from docx import Document
    from docx.shared import Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ==========================================
# 1. PAGE CONFIGURATION & DATA LOAD
# ==========================================
st.set_page_config(page_title="Solvent Reduction Analysis", page_icon="🎨", layout="wide")
st.title("🎨 稀釋劑減量機會分析 (Solvent Reduction Opportunity)")

if not st.session_state.get("raw_data_loaded", False) or st.session_state.get("group_a_data") is None:
    st.warning("⚠️ Please return to the Main App page and upload the raw data first. (尚未載入資料，請先返回首頁載入)")
    st.stop()

df = st.session_state["group_a_data"].copy()

# ==========================================
# 2. DATA CLEANING & PREPARATION
# ==========================================
text_cols = ["Vendor", "Resin", "Solvent_Type", "塗料批號", "線別", "塗裝位置"]
for col in text_cols:
    df[col] = df.get(col, "Unknown").fillna("Unknown").astype(str).str.strip()

df["Paint_Code"] = df.get("塗料編號", pd.Series("Unknown", index=df.index)).fillna("Unknown").astype(str).str.strip().str.upper()
df["Solvent_Type"] = df["Solvent_Type"].str.upper()

invalid_vals = {"", "NAN", "NONE", "NULL", "N/A", "NA", "-", "--"}
df.replace(list(invalid_vals), "Unknown", inplace=True)

df["Batch_ID"] = df["塗料批號"]
pos_map = {"TP": "Primer", "正底漆": "Primer", "BP": "Primer", "背底漆": "Primer", "TF": "Top Finish", "正面漆": "Top Finish", "BF": "Back Finish", "背面漆": "Back Finish"}
df["Position_UI"] = df["塗裝位置"].map(pos_map).fillna(df["塗裝位置"])

num_cols = ["塗料重量", "添加重量", "黏度(秒)", "黏度(秒)_1", "溫度"]
for col in num_cols:
    df[col] = pd.to_numeric(df.get(col, np.nan), errors="coerce")

df["Delta_V"] = df["黏度(秒)"] - df["黏度(秒)_1"]
df["Solvent_Ratio_Percent"] = np.where(df["塗料重量"] > 0, df["添加重量"] / df["塗料重量"] * 100, np.nan)
df["Viscosity_Sensitivity"] = np.where(df["Solvent_Ratio_Percent"] > 0, df["Delta_V"] / df["Solvent_Ratio_Percent"], np.nan)

df = df[(df["塗料重量"]>0) & (df["添加重量"]>0) & (df["黏度(秒)"]>0) & (df["黏度(秒)_1"]>0) & (df["Delta_V"]>0)].copy()

if df.empty:
    st.warning("⚠️ No valid dilution records remain after data cleaning.")
    st.stop()

date_col = next((c for c in ["攪拌日期", "調整日期", "生產日期", "Date"] if c in df.columns), None)
if date_col:
    df["_Analysis_Date"] = pd.to_datetime(df[date_col], errors="coerce")
else:
    df["_Analysis_Date"] = pd.NaT

# ==========================================
# 3. CORE LOGIC HELPER
# ==========================================
def build_summary(source_df, group_cols):
    if source_df.empty: return pd.DataFrame()
    
    agg_dict = {
        "Adjustment_Records": ("Paint_Code", "size"),
        "Historical_Batches": ("Batch_ID", "nunique"),
        "Total_Paint_kg": ("塗料重量", "sum"),
        "Total_Solvent_kg": ("添加重量", "sum"),
        "Median_Paint_kg": ("塗料重量", "median"),
        "Median_Solvent_kg": ("添加重量", "median"),
        "Median_Ratio_Percent": ("Solvent_Ratio_Percent", "median"),
        "Median_Before_Viscosity": ("黏度(秒)", "median"),
        "Median_After_Viscosity": ("黏度(秒)_1", "median"),
        "Median_Viscosity_Drop": ("Delta_V", "median"),
        "Median_Dilution_Efficiency": ("Viscosity_Sensitivity", "median"),
    }
    if "線別" not in group_cols:
        agg_dict["Production_Lines"] = ("線別", lambda x: x[x != "Unknown"].nunique())

    summary = source_df.groupby(group_cols, dropna=False).agg(**agg_dict).reset_index()
    summary["Weighted_Ratio_Percent"] = np.where(summary["Total_Paint_kg"] > 0, summary["Total_Solvent_kg"] / summary["Total_Paint_kg"] * 100, np.nan)
    return summary


def mpl_fig_to_stream(fig, dpi=200):
    """Chuyển 1 figure Matplotlib thành ảnh PNG trong bộ nhớ (không cần Chrome/kaleido)."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight", facecolor="white")
    buf.seek(0)
    plt.close(fig)
    return buf

# ==========================================
# 4. GLOBAL FILTERS
# ==========================================
st.markdown("---")
st.subheader("🔍 全局篩選條件 (Global Filters)")

filter_df = df.copy()
col1, col2, col3, col4, col5 = st.columns(5)

vendor_opts = ["All"] + sorted([str(x) for x in filter_df["Vendor"].unique() if x != "Unknown"])
selected_vendor = col1.selectbox("Vendor (供應商)", vendor_opts)
if selected_vendor != "All": filter_df = filter_df[filter_df["Vendor"] == selected_vendor]

resin_opts = ["All"] + sorted([str(x) for x in filter_df["Resin"].unique() if x != "Unknown"])
selected_resin = col2.selectbox("Resin Type (樹脂種類)", resin_opts)
if selected_resin != "All": filter_df = filter_df[filter_df["Resin"] == selected_resin]

pos_opts = ["All"] + sorted([str(x) for x in filter_df["Position_UI"].unique() if x != "Unknown"])
selected_position = col3.selectbox("Coating Position (塗裝位置)", pos_opts)
if selected_position != "All": filter_df = filter_df[filter_df["Position_UI"] == selected_position]

solvent_opts = ["All"] + sorted([str(x) for x in filter_df["Solvent_Type"].unique() if x != "Unknown"])
selected_solvent = col4.selectbox("Solvent Type (稀釋劑種類)", solvent_opts)
if selected_solvent != "All": filter_df = filter_df[filter_df["Solvent_Type"] == selected_solvent]

line_opts = sorted([str(x) for x in filter_df["線別"].unique() if x != "Unknown"])
selected_lines = col5.multiselect("Production Line (產線)", line_opts, default=line_opts)

if selected_lines:
    filter_df = filter_df[filter_df["線別"].isin(selected_lines)]
else:
    st.warning("⚠️ Please select at least one production line. (請至少選擇一條產線)")
    st.stop()

if filter_df.empty:
    st.warning("⚠️ No records match the global analysis filters. (無符合篩選條件的資料)")
    st.stop()

st.markdown("---")
if "_Analysis_Date" in filter_df.columns and not filter_df["_Analysis_Date"].isna().all():
    min_date = filter_df["_Analysis_Date"].min().strftime("%Y-%m-%d")
    max_date = filter_df["_Analysis_Date"].max().strftime("%Y-%m-%d")
    period_label = f"{min_date} ➔ {max_date}"
else:
    period_label = "All available data (全部歷史資料)"

st.info(f"📅 **資料期間 (Analysis Period):** {period_label} ｜ 📊 **符合條件紀錄數 (Valid Records):** {len(filter_df):,} 筆")
filter_details = f"Vendor: {selected_vendor} | Resin: {selected_resin} | Position: {selected_position} | Solvent: {selected_solvent}"
st.markdown("<br>", unsafe_allow_html=True)

# Dict lưu các biểu đồ Plotly hiển thị trên web (không đổi)
exported_figs = {}
# Dict lưu các biểu đồ Matplotlib tương ứng, chỉ dùng để xuất Word (không cần Chrome/kaleido)
export_mpl_figs = {}

# ==========================================
# 5. HIERARCHICAL OVERVIEW (TREEMAP)
# ==========================================
st.subheader("🗂️ 塗料階層總覽 (Hierarchical Overview)")
st.markdown("Hierarchy: **Vendor ➔ Resin ➔ Position ➔ Solvent Type ➔ Paint Code**. Box size represents total solvent usage (kg).")

tree_df = filter_df.groupby(["Vendor", "Resin", "Position_UI", "Solvent_Type", "Paint_Code"]).agg(
    添加重量=("添加重量", "sum"),
    Delta_V=("Delta_V", "median"), 
    Solvent_Ratio_Percent=("Solvent_Ratio_Percent", "median")
).reset_index()
tree_df = tree_df[tree_df["添加重量"] > 0]

fig_tree = px.treemap(
    tree_df,
    path=[px.Constant("Total"), "Vendor", "Resin", "Position_UI", "Solvent_Type", "Paint_Code"],
    values="添加重量", 
    color="Resin",  
    color_discrete_sequence=px.colors.qualitative.Pastel,
    custom_data=["Delta_V", "Solvent_Ratio_Percent"],
    title=f"Hierarchical Breakdown of Solvent Usage (kg)<br><sup>Filters: {filter_details}</sup>",
    height=700
)

fig_tree.update_traces(
    texttemplate="<b>%{label}</b><br>%{value:,.0f} kg",
    hovertemplate="<b>%{label}</b><br>Solvent Usage: %{value:,.1f} kg<br>Visc Drop (Biến động): ~%{customdata[0]:.1f} s<br>Solvent Added (Tỷ lệ thêm): ~%{customdata[1]:.1f}%",
    root_color="#f8f9fa"
)
fig_tree.update_layout(margin=dict(t=90, l=10, r=10, b=10)) 

st.plotly_chart(fig_tree, use_container_width=True)
exported_figs["1. Hierarchical Treemap"] = fig_tree

# ---- Bản Matplotlib tương ứng cho Word: Top 15 Paint Code theo lượng dung môi ----
_tree_bar = (
    tree_df.groupby(["Resin", "Paint_Code"])["添加重量"].sum()
    .reset_index()
    .sort_values("添加重量", ascending=False)
    .head(15)
)
if not _tree_bar.empty:
    _resins = sorted(_tree_bar["Resin"].unique())
    _cmap = cm.get_cmap("Pastel1", max(len(_resins), 3))
    _color_map = {r: _cmap(i) for i, r in enumerate(_resins)}

    fig_tree_mpl, ax = plt.subplots(figsize=(9, max(4, 0.4 * len(_tree_bar))), dpi=150)
    _plot_df = _tree_bar.sort_values("添加重量", ascending=True)
    bar_colors = [_color_map[r] for r in _plot_df["Resin"]]
    ax.barh(_plot_df["Paint_Code"], _plot_df["添加重量"], color=bar_colors)
    ax.set_xlabel("Solvent Usage (kg)")
    ax.set_title(f"Top 15 Paint Codes by Solvent Usage (kg)\nFilters: {filter_details}", fontsize=10, loc="left")
    for i, (val, resin) in enumerate(zip(_plot_df["添加重量"], _plot_df["Resin"])):
        ax.text(val, i, f" {val:,.0f} kg", va="center", fontsize=8)
    handles = [plt.Rectangle((0,0),1,1, color=_color_map[r]) for r in _resins]
    ax.legend(handles, _resins, title="Resin", loc="lower right", fontsize=8)
    fig_tree_mpl.tight_layout()
    export_mpl_figs["1. Hierarchical Overview (Top 15 by Solvent Usage)"] = fig_tree_mpl

st.markdown("<br>", unsafe_allow_html=True)

# ==========================================
# 6. TABS & VISUALIZATION
# ==========================================
tab_ranking, tab_detail, tab_line = st.tabs([
    "1️⃣ 塗料消耗排名 (Paint Code Ranking)", 
    "2️⃣ 塗料詳細分析 (Paint Code Details)", 
    "3️⃣ 產線黏度比較 (Line Comparison)"
])

# ----- TAB 1: RANKING -----
with tab_ranking:
    st.subheader("1. Paint Code Solvent Consumption (Top 20)")
    
    full_summary_df = build_summary(filter_df, ["Vendor", "Resin", "Position_UI", "Paint_Code", "Solvent_Type"])
    summary_df = full_summary_df.sort_values("Total_Solvent_kg", ascending=False).head(20).reset_index(drop=True)
    summary_df.insert(0, "Rank", np.arange(1, len(summary_df) + 1))

    chart_height = max(450, len(summary_df) * 32)
    
    ch1, ch2 = st.columns(2)
    with ch1:
        df_melt = summary_df.melt(id_vars="Paint_Code", value_vars=["Total_Paint_kg", "Total_Solvent_kg"])
        df_melt["variable"] = df_melt["variable"].map({"Total_Paint_kg": "塗料 (Paint)", "Total_Solvent_kg": "稀釋劑 (Solvent)"})
        
        fig1 = px.bar(
            df_melt, x="value", y="Paint_Code", color="variable", barmode="group", orientation='h', 
            height=chart_height, color_discrete_map={"塗料 (Paint)": "#5B8FF9", "稀釋劑 (Solvent)": "#F6BD16"}
        )
        fig1.update_yaxes(dtick=1, title="", categoryorder="total ascending")
        fig1.update_xaxes(title="Weight (kg)")
        fig1.update_layout(title="Paint vs Solvent Usage", legend_title_text="", legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1))
        st.plotly_chart(fig1, use_container_width=True)
        exported_figs["2. Paint vs Solvent (kg)"] = fig1

        # ---- Matplotlib tương ứng ----
        _s1 = summary_df.sort_values("Total_Solvent_kg", ascending=True)
        fig1_mpl, ax = plt.subplots(figsize=(8, max(4, 0.4 * len(_s1))), dpi=150)
        y_pos = np.arange(len(_s1))
        bar_h = 0.35
        ax.barh(y_pos + bar_h/2, _s1["Total_Paint_kg"], height=bar_h, color="#5B8FF9", label="塗料 (Paint)")
        ax.barh(y_pos - bar_h/2, _s1["Total_Solvent_kg"], height=bar_h, color="#F6BD16", label="稀釋劑 (Solvent)")
        ax.set_yticks(y_pos)
        ax.set_yticklabels(_s1["Paint_Code"], fontsize=8)
        ax.set_xlabel("Weight (kg)")
        ax.set_title("Paint vs Solvent Usage", fontsize=10, loc="left")
        ax.legend(fontsize=8, loc="lower right")
        fig1_mpl.tight_layout()
        export_mpl_figs["2. Paint vs Solvent (kg)"] = fig1_mpl

    with ch2:
        sorted_df = summary_df.sort_values("Weighted_Ratio_Percent", ascending=True)
        fig2 = px.bar(
            sorted_df, x="Weighted_Ratio_Percent", y="Paint_Code", orientation='h', text_auto='.2f', 
            height=chart_height, color_discrete_sequence=["#5B8FF9"]
        )
        fig2.update_traces(textposition="outside", cliponaxis=False)
        fig2.update_yaxes(dtick=1, title="")
        fig2.update_xaxes(title="Ratio (%)")
        fig2.update_layout(title="Weighted Solvent Ratio (%)")
        st.plotly_chart(fig2, use_container_width=True)
        exported_figs["3. Weighted Solvent Ratio"] = fig2

        # ---- Matplotlib tương ứng ----
        fig2_mpl, ax = plt.subplots(figsize=(8, max(4, 0.4 * len(sorted_df))), dpi=150)
        ax.barh(sorted_df["Paint_Code"], sorted_df["Weighted_Ratio_Percent"], color="#5B8FF9")
        for i, val in enumerate(sorted_df["Weighted_Ratio_Percent"]):
            ax.text(val, i, f" {val:.2f}", va="center", fontsize=8)
        ax.set_xlabel("Ratio (%)")
        ax.set_title("Weighted Solvent Ratio (%)", fontsize=10, loc="left")
        ax.tick_params(axis='y', labelsize=8)
        fig2_mpl.tight_layout()
        export_mpl_figs["3. Weighted Solvent Ratio"] = fig2_mpl

    st.markdown("---")
    
    # Dual Axis Chart
    fig_dual = go.Figure()
    fig_dual.add_trace(go.Bar(x=summary_df["Paint_Code"], y=summary_df["Total_Paint_kg"], name="Paint (kg)", marker_color="#5B8FF9", yaxis="y1"))
    fig_dual.add_trace(go.Bar(x=summary_df["Paint_Code"], y=summary_df["Total_Solvent_kg"], name="Solvent (kg)", marker_color="#F6BD16", yaxis="y1"))
    fig_dual.add_trace(go.Scatter(x=summary_df["Paint_Code"], y=summary_df["Weighted_Ratio_Percent"], name="Solvent Ratio (%)", mode="lines+markers", line=dict(color="#5AD8A6", width=3), marker=dict(size=8), yaxis="y2"))
    
    for i, row in summary_df.iterrows():
        fig_dual.add_annotation(
            x=row["Paint_Code"], y=row["Weighted_Ratio_Percent"],
            text=f"<b>{row['Weighted_Ratio_Percent']:.2f}%</b>", 
            xref="x", yref="y2", showarrow=False, yshift=18, 
            font=dict(color="black", size=12), bgcolor="rgba(255, 255, 255, 0.85)", borderpad=2
        )

    fig_dual.update_layout(
        title=f"Paint & Solvent Usage vs. Solvent Ratio<br><sup>Filters Applied: {filter_details}</sup>",
        xaxis=dict(title="Paint Code"),
        yaxis=dict(title="Weight (kg)", side="left", showgrid=False),
        yaxis2=dict(title="Solvent Ratio (%)", overlaying="y", side="right", showgrid=False, range=[0, summary_df["Weighted_Ratio_Percent"].max() * 1.25]),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        height=600
    )
    st.plotly_chart(fig_dual, use_container_width=True)
    exported_figs["4. Dual Axis Usage vs Ratio"] = fig_dual

    # ---- Matplotlib tương ứng ----
    fig_dual_mpl, ax1 = plt.subplots(figsize=(11, 6), dpi=150)
    x = np.arange(len(summary_df))
    w = 0.35
    ax1.bar(x - w/2, summary_df["Total_Paint_kg"], width=w, color="#5B8FF9", label="Paint (kg)")
    ax1.bar(x + w/2, summary_df["Total_Solvent_kg"], width=w, color="#F6BD16", label="Solvent (kg)")
    ax1.set_xticks(x)
    ax1.set_xticklabels(summary_df["Paint_Code"], rotation=45, ha="right", fontsize=7)
    ax1.set_ylabel("Weight (kg)")
    ax1.set_title(f"Paint & Solvent Usage vs. Solvent Ratio\nFilters: {filter_details}", fontsize=10, loc="left")

    ax2 = ax1.twinx()
    ax2.plot(x, summary_df["Weighted_Ratio_Percent"], color="#5AD8A6", marker="o", linewidth=2)
    ax2.set_ylabel("Solvent Ratio (%)")
    ax2.set_ylim(0, summary_df["Weighted_Ratio_Percent"].max() * 1.25 if summary_df["Weighted_Ratio_Percent"].max() > 0 else 1)
    for xi, val in zip(x, summary_df["Weighted_Ratio_Percent"]):
        ax2.annotate(f"{val:.2f}%", (xi, val), textcoords="offset points", xytext=(0, 8), ha="center", fontsize=7)

    lines_1, labels_1 = ax1.get_legend_handles_labels()
    lines_2, labels_2 = ax2.get_legend_handles_labels()
    ax1.legend(lines_1 + lines_2, labels_1 + labels_2, loc="upper right", fontsize=8)
    fig_dual_mpl.tight_layout()
    export_mpl_figs["4. Dual Axis Usage vs Ratio"] = fig_dual_mpl

# ----- TAB 2: DETAILS -----
with tab_detail:
    st.subheader("2. Paint Code Details")
    selected_code = st.selectbox("Select Paint Code", summary_df["Paint_Code"].unique())
    detail_df = filter_df[filter_df["Paint_Code"] == selected_code].copy()
    detail_title_filter = f"{filter_details} | Paint Code: {selected_code}"

    ch3, ch4 = st.columns(2)
    with ch3:
        detail_df = detail_df.reset_index(drop=True)
        detail_df["Record_Index"] = detail_df.index + 1
        
        fig3 = go.Figure()
        fig3.add_trace(go.Scatter(x=detail_df["Record_Index"], y=detail_df["黏度(秒)"], mode="lines+markers", name="Before Viscosity", marker=dict(color="#5B8FF9", size=8), yaxis="y1"))
        fig3.add_trace(go.Scatter(x=detail_df["Record_Index"], y=detail_df["黏度(秒)_1"], mode="lines+markers", name="After Viscosity", marker=dict(color="#5AD8A6", size=8), yaxis="y1"))
        
        has_temp = "溫度" in detail_df.columns and not detail_df["溫度"].isna().all()
        if has_temp:
            fig3.add_trace(go.Scatter(x=detail_df["Record_Index"], y=detail_df["溫度"], mode="lines+markers", name="Temperature (°C)", marker=dict(color="#F6BD16", size=8, symbol="diamond"), line=dict(color="#F6BD16", width=2, dash="dot"), yaxis="y2"))
            chart_title = "Viscosity & Temperature Variation (Before vs After)"
        else:
            chart_title = "Viscosity Variation (Before vs After)"
        
        fig3.update_layout(
            title=f"{chart_title}<br><sup>Filters: {detail_title_filter}</sup>",
            xaxis_title="Record Output Sequence", yaxis=dict(title="Viscosity (s)", side="left"),
            yaxis2=dict(title="Temperature (°C)", overlaying="y", side="right", showgrid=False),
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
        )
        st.plotly_chart(fig3, use_container_width=True)
        exported_figs["5. Viscosity & Temperature Variation"] = fig3

        # ---- Matplotlib tương ứng ----
        fig3_mpl, ax1 = plt.subplots(figsize=(9, 5), dpi=150)
        ax1.plot(detail_df["Record_Index"], detail_df["黏度(秒)"], marker="o", color="#5B8FF9", label="Before Viscosity")
        ax1.plot(detail_df["Record_Index"], detail_df["黏度(秒)_1"], marker="o", color="#5AD8A6", label="After Viscosity")
        ax1.set_xlabel("Record Output Sequence")
        ax1.set_ylabel("Viscosity (s)")
        ax1.set_title(f"{chart_title}\nFilters: {detail_title_filter}", fontsize=10, loc="left")
        lines_1, labels_1 = ax1.get_legend_handles_labels()
        if has_temp:
            ax2 = ax1.twinx()
            ax2.plot(detail_df["Record_Index"], detail_df["溫度"], marker="D", linestyle=":", color="#F6BD16", label="Temperature (°C)")
            ax2.set_ylabel("Temperature (°C)")
            lines_2, labels_2 = ax2.get_legend_handles_labels()
            ax1.legend(lines_1 + lines_2, labels_1 + labels_2, loc="upper right", fontsize=8)
        else:
            ax1.legend(lines_1, labels_1, loc="upper right", fontsize=8)
        fig3_mpl.tight_layout()
        export_mpl_figs["5. Viscosity & Temperature Variation"] = fig3_mpl

    with ch4:
        line_usage = build_summary(detail_df, ["線別"]).sort_values("Total_Solvent_kg")
        fig4 = px.bar(line_usage, x="Total_Solvent_kg", y="線別", text="Weighted_Ratio_Percent", orientation='h', title=f"Solvent Usage by Line<br><sup>Filters: {detail_title_filter}</sup>", color_discrete_sequence=["#5B8FF9"])
        fig4.update_traces(texttemplate='%{text:.2f}%', textposition='outside')
        fig4.update_yaxes(title="")
        fig4.update_xaxes(title="Total Solvent (kg)")
        st.plotly_chart(fig4, use_container_width=True)
        exported_figs["6. Solvent Usage by Line"] = fig4

        # ---- Matplotlib tương ứng ----
        fig4_mpl, ax = plt.subplots(figsize=(8, max(3, 0.5 * len(line_usage))), dpi=150)
        ax.barh(line_usage["線別"], line_usage["Total_Solvent_kg"], color="#5B8FF9")
        for i, (val, pct) in enumerate(zip(line_usage["Total_Solvent_kg"], line_usage["Weighted_Ratio_Percent"])):
            ax.text(val, i, f" {pct:.2f}%", va="center", fontsize=8)
        ax.set_xlabel("Total Solvent (kg)")
        ax.set_title(f"Solvent Usage by Line\nFilters: {detail_title_filter}", fontsize=10, loc="left")
        fig4_mpl.tight_layout()
        export_mpl_figs["6. Solvent Usage by Line"] = fig4_mpl

# ----- TAB 3: LINE COMPARISON -----
with tab_line:
    st.subheader("3. Production Line Comparison")
    comp_code = st.selectbox("Select Paint Code for Line Comparison", summary_df["Paint_Code"].unique(), key="line_comp")
    comp_df = filter_df[filter_df["Paint_Code"] == comp_code]
    line_summary = build_summary(comp_df, ["線別"]).sort_values("線別")

    if len(line_summary) >= 2:
        ch5, ch6 = st.columns(2)
        with ch5:
            fig5 = go.Figure()
            for i, row in line_summary.iterrows():
                fig5.add_trace(go.Scatter(x=[row["Median_After_Viscosity"], row["Median_Before_Viscosity"]], y=[row["線別"], row["線別"]], mode="lines+markers", marker=dict(size=12), name=row["線別"]))
            fig5.update_layout(title="Viscosity Drop (Before vs After)", xaxis_title="Viscosity (s)", yaxis_title="")
            st.plotly_chart(fig5, use_container_width=True)
            exported_figs["7. Line Comparison - Viscosity Drop"] = fig5

            # ---- Matplotlib tương ứng ----
            fig5_mpl, ax = plt.subplots(figsize=(8, max(3, 0.6 * len(line_summary))), dpi=150)
            for i, row in line_summary.iterrows():
                ax.plot([row["Median_After_Viscosity"], row["Median_Before_Viscosity"]], [row["線別"], row["線別"]], marker="o", markersize=10, linewidth=2)
            ax.set_xlabel("Viscosity (s)")
            ax.set_title("Viscosity Drop (Before vs After)", fontsize=10, loc="left")
            fig5_mpl.tight_layout()
            export_mpl_figs["7. Line Comparison - Viscosity Drop"] = fig5_mpl

        with ch6:
            fig6 = px.bar(line_summary.sort_values("Weighted_Ratio_Percent"), x="Weighted_Ratio_Percent", y="線別", orientation='h', text_auto='.2f', title="Weighted Solvent Ratio", color_discrete_sequence=["#5AD8A6"])
            fig6.update_yaxes(title="")
            st.plotly_chart(fig6, use_container_width=True)
            exported_figs["8. Line Comparison - Solvent Ratio"] = fig6

            # ---- Matplotlib tương ứng ----
            _s6 = line_summary.sort_values("Weighted_Ratio_Percent")
            fig6_mpl, ax = plt.subplots(figsize=(8, max(3, 0.6 * len(_s6))), dpi=150)
            ax.barh(_s6["線別"], _s6["Weighted_Ratio_Percent"], color="#5AD8A6")
            for i, val in enumerate(_s6["Weighted_Ratio_Percent"]):
                ax.text(val, i, f" {val:.2f}", va="center", fontsize=8)
            ax.set_xlabel("Ratio (%)")
            ax.set_title("Weighted Solvent Ratio", fontsize=10, loc="left")
            fig6_mpl.tight_layout()
            export_mpl_figs["8. Line Comparison - Solvent Ratio"] = fig6_mpl


# ==========================================
# 7. XUẤT FILE WORD (EXPORT BÁO CÁO TỰ ĐỘNG)
# Đã bỏ hoàn toàn phụ thuộc vào kaleido/Chrome.
# Toàn bộ ảnh trong file Word được vẽ bằng Matplotlib,
# hoạt động trên mọi môi trường (kể cả Streamlit Cloud/GitHub)
# mà không cần cài thêm Chrome.
# ==========================================
st.markdown("---")
st.subheader("📄 Xuất báo cáo (Export Report)")

if not HAS_DOCX:
    st.warning("⚠️ Máy chủ chưa được cài đặt công cụ tạo file Word. Để sử dụng tính năng này, bạn hãy thêm vào file requirements.txt:\n\n`python-docx`")
else:
    if st.button("🚀 Chụp toàn bộ Biểu đồ và Lưu thành file Word"):
        with st.spinner("Đang xử lý hình ảnh và tạo file Word... (Quá trình này mất khoảng vài giây)"):
            try:
                doc = Document()
                doc.add_heading('BÁO CÁO PHÂN TÍCH TIÊU THỤ DUNG MÔI SƠN', 0)
                doc.add_paragraph(f"Ngày phân tích: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph(f"Bộ lọc dữ liệu: {filter_details}")

                # Hàm chèn ảnh Matplotlib (PNG) vào file Word
                def add_chart_to_word(doc, title, mpl_fig, note=None):
                    doc.add_heading(title, level=2)
                    image_stream = mpl_fig_to_stream(mpl_fig)
                    doc.add_picture(image_stream, width=Inches(6.0))
                    if note:
                        note_p = doc.add_paragraph(note)
                        note_p.italic = True
                    doc.add_paragraph("")

                # Duyệt qua các biểu đồ Matplotlib đã được lưu trong dict
                for chart_title, mpl_fig in export_mpl_figs.items():
                    note = None
                    if "Hierarchical" in chart_title:
                        note = ("Ghi chú: Do giới hạn kỹ thuật khi xuất file Word, biểu đồ Treemap đa cấp trên "
                                "giao diện web được trình bày dưới dạng cột ngang Top 15 mã sơn theo lượng dung môi "
                                "sử dụng (cùng số liệu, khác cách hiển thị).")
                    add_chart_to_word(doc, chart_title, mpl_fig, note)

                # Lưu file vào bộ nhớ đệm Streamlit
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

                st.success("✅ Đã tạo file Word thành công! Nhấn nút bên dưới để tải về máy.")
                st.download_button(
                    label="⬇️ TẢI FILE BÁO CÁO (Word .docx)",
                    data=doc_io,
                    file_name="Bao_Cao_Dung_Moi_Son.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"❌ Lỗi khi tạo file Word. Chi tiết lỗi: {e}")
